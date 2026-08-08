"""سجلّ الانحدار الموحّد — one guard per real incident, one meta-test for coverage.

> **الغرض (أمر المُشرِف).** لكل حادثة إنتاجية حقيقية في هذا المستودع — صفوف
> `docs/LESSONS.md` **وفخاخ** `silk-operations` §2 (THE TRAPS) — حارسٌ واحد
> يُفشِل على عودة نفس العائلة، و**اختبار تغطية شامل** (meta) يثبت أنّ كل حادثة
> مُسجَّلة هنا فعلاً — فلا تسقط حادثة من الشبكة بصمت. يشمل ذلك **الحوادث الثلاث
> لعطل 501 في تصدير docx** (صفوف LESSONS ٣/١١/١٣) بحُرّاس سلوكيين فعليين.
>
> **Why a registry (not just per-file lock-tests).** Lock-tests live scattered
> across `tests/`; this file is the single index that maps EVERY known incident
> to a live guard and then proves — mechanically — that the index is complete
> against both incident ledgers. A new incident that lands in either ledger
> without a registry entry fails `test_meta_registry_covers_every_known_incident`.

هرمتي بالكامل (قراءة مصدر + سلوك محلي، بلا شبكة). الحُرّاس السلوكية (٣/١١/١٣)
تبني/تنقّي فعلياً من المدوّنة القانونية الحقيقية الشكل — لا نماذج مثالية.

Run: python3 -m pytest tests/test_regression_registry.py -q
"""
from __future__ import annotations

import os
import re
import sys
import tempfile

import pytest

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)
sys.path.insert(0, os.path.join(_ROOT, "tools"))


def _read(rel: str) -> str:
    with open(os.path.join(_ROOT, rel), encoding="utf-8") as f:
        return f.read()


def _exists(rel: str) -> bool:
    return os.path.exists(os.path.join(_ROOT, rel))


def _needles(rel: str, *needles: str):
    """حارس وجود: كل إبرة حاضرة في الملف — يعيد callable للتسجيل."""
    def check():
        assert _exists(rel), f"ملف الإنفاذ مفقود: {rel}"
        src = _read(rel)
        missing = [n for n in needles if n not in src]
        assert not missing, f"{rel}: رموز/علامات إنفاذ مفقودة {missing}"
    return check


def _absent(rel: str, *forbidden: str):
    def check():
        src = _read(rel)
        present = [n for n in forbidden if n in src]
        assert not present, f"{rel}: رموز يجب أن تكون قد أُزيلت لا تزال {present}"
    return check


# ── حُرّاس سلوكية للحوادث الثلاث لعطل docx-501 (LESSONS ٣/١١/١٣) ──────────────

def _guard_docx501_row3():
    """LESSONS ٣ — 501 شُحن لأن الاختبارات نماذج مموّهة. الحارس: تصدير docx
    العميل يُنتِج ملفاً حقيقياً قابلاً للفتح من **المدوّنة القانونية الحقيقية
    الشكل** (لا نموذج)، بلا 501."""
    import silk_render
    from silk_reports import render_client_docx
    from canonical_netherlands import netherlands_research_blob
    view = silk_render.build_view(netherlands_research_blob())
    path = render_client_docx(view, os.path.join(tempfile.mkdtemp(), "c.docx"))
    assert os.path.exists(path)
    from docx import Document
    doc = Document(path)
    assert any(p.text.strip() for p in doc.paragraphs), "docx فارغ"


def _guard_docx501_row11():
    """LESSONS ١١ — 501 تكرّر لأن محفّزات الحارس العربية بلا استبدال مقابل.
    الحارس: مصطلح حكم عربي ممنوع («درجة الثقة») يُحيَّد فعلياً بمُطهِّر/منقِّي
    العميل فلا يبقى في مخرَج التصدير."""
    from silk_reports import _client_redact_text, _client_forbidden_hits
    leaked = "التقييم يعتمد درجة الثقة العالية على مصدر البيانات"
    cleaned = _client_redact_text(leaked)
    assert not _client_forbidden_hits(cleaned), (
        f"محفّز حارس عربي بقي بعد التنقية: {_client_forbidden_hits(cleaned)}")


def _guard_docx501_row13():
    """LESSONS ١٣ — docx يفشل حياً (501) بينما الهرمتي أخضر؛ الحارس كان يرفض
    على تسرّب واحد. القاعدة «نقِّ لا ترفض»: مصطلح إنجليزي عارٍ متسرّب يُستبدَل
    بمحايد ويُسلَّم المستند — لا 501. الحارس: `_client_redact_text` ينقّي
    مصطلحاً إنجليزياً تشغيلياً بدل رفعه."""
    from silk_reports import _client_redact_text, _client_forbidden_hits
    leaked = "mission status: successful run"
    assert _client_forbidden_hits(leaked), "التهيئة خاطئة — النص يجب أن يتسرّب أولاً"
    cleaned = _client_redact_text(leaked)
    assert not _client_forbidden_hits(cleaned), (
        f"مصطلح إنجليزي تشغيلي بقي بعد التنقية: {_client_forbidden_hits(cleaned)}")


# ── حُرّاس سلوكية لفخّي المسار (markets:[] + view بعد التخزين) ────────────────

def _guard_trap_markets_empty():
    """TRAP «markets:[] misroutes exporters» — نتيجة /research دوماً markets:[]؛
    أي مسار عرض يثق بـ`markets[0]` يحصل على {} فيُنتِج صفحة /analyze فارغة.
    الحارس: build_view للمدوّنة القانونية يُنتِج فرع deep_research (لا قالب
    فارغ) رغم markets:[]."""
    import silk_render
    from canonical_netherlands import netherlands_research_blob
    blob = netherlands_research_blob()
    assert blob["markets"] == [], "المدوّنة القانونية يجب أن تحمل markets:[]"
    view = silk_render.build_view(blob)
    assert view.get("deep_research"), "build_view لم يُنتِج فرع deep_research"


def _guard_trap_view_after_persist():
    """TRAP «view attached AFTER persist on one path» — البلوب المخزَّن قد لا
    يحمل view مشتقّاً؛ مسار القراءة يجب أن يعيد بناءه. الحارس: `GET /analyses/{id}`
    يبني view عند غيابه (`found["view"] = _view(found)`)."""
    src = _read("api.py")
    assert 'found["view"] = _view(found)' in src or 'found["view"]=_view(found)' in src, (
        "مسار /analyses/{id} لا يعيد بناء view الغائب")
    assert 'found.setdefault("analysis_id"' in src, (
        "مسار /analyses/{id} لا يضمن analysis_id للبلوبات الأقدم")


# ── حُرّاس تراب باقية (وجود رمز الإصلاح، file:line-anchored) ──────────────────

def _guard_trap_two_sanitizers():
    _needles("silk_reports.py", "_client_assert_clean", "_client_sanitize")()
    _needles("silk_render.py", "_strip_internal_plumbing")()


def _guard_trap_redaction_mangling():
    # الإصلاح الالتفافي: مرحلة التصعيد سُمّيت `..._escalate{attempt}` لتفادي
    # الجزء القصير الذي كان يُشوَّه؛ والمنقِّح `_redact` باقٍ. (بنيوياً مفتوح
    # — لا حارس طول أدنى بعد؛ مُتتبَّع في silk-operations §2.)
    _needles("silk_ai_judge.py", "_escalate{attempt}")()
    _absent("silk_ai_judge.py", "maxtok_retry")()
    _needles("silk_diagnostics.py", "def _redact")()


def _guard_trap_strip_plumbing_three_leaks():
    """TRAP «_strip_internal_plumbing leaked three raw forms» — الحارس يبني
    السلاسل الإنتاجية الحرفية الثلاث ويؤكّد تحييد كلٍّ منها (silk-operations §4:
    القفل بالسلسلة الحرفية). الثلاث: ريبر DataPoint(...) يُحيَّد كاملاً بلا
    نصف-ترجمة؛ JSON مضمَّن بمفاتيح score/summary يُحيَّد؛ رمز حكم بأي حالة أحرف."""
    from silk_render import _strip_internal_plumbing
    # (١) ريبر DataPoint(...) — كامل التحييد، لا نصف-ترجمة، تُستخرَج القيمة
    dp = ("مبنيّ على DataPoint(value='واردات 120 مليون دولار', source='UN "
          "Comtrade', confidence=0.9, note='n', retrieved_at='2026-07-01', "
          "status='')")
    o1 = _strip_internal_plumbing(dp)
    assert "DataPoint(" not in o1 and "confidence=" not in o1, o1
    assert "درجة الثقة=" not in o1, f"نصف-ترجمة: {o1}"
    assert "واردات 120 مليون دولار" in o1, o1
    # (٢) JSON مضمَّن بمفاتيح score/summary
    o2 = _strip_internal_plumbing('التوصية: {"score": 0.72, "summary": "سوق واعد"}')
    assert "{" not in o2 and '"summary"' not in o2, o2
    assert "سوق واعد" in o2, o2
    # (٣) رمز حكم بأي حالة أحرف
    o3 = _strip_internal_plumbing("الحكم go — مراقبة قبل الدخول")
    assert not re.search(r"\bgo\b", o3, re.I), f"رمز حكم خام بقي: {o3}"


def _guard_trap_parallel_cache_window():
    # فخّ معروف غير مُصلَح بعد (نافذة الذاكرة المؤقتة عبر ١٢ بعثة متوازية).
    # الحارس يُبقيه مُتتبَّعاً: آلية التوازي (ThreadPoolExecutor) لا تزال في
    # مُشغّل البعثات، والفخّ موسوم صراحةً «Known, not yet fixed» في المهارة.
    _needles("silk_missions.py", "ThreadPoolExecutor")()
    _needles(".claude/skills/silk-operations/SKILL.md", "not yet fixed")()


# كل مدخلة: Incident(key, source, match, check)
#   source: "LESSONS" (key=رقم الصفّ int) أو "trap" (key=slug، match=جزء من
#   اسم الفخّ العريض في §2). check: callable يُفشِل على عودة الانحدار.


def _guard_datapoint_repr_flexible():
    """LESSONS ١٧ — ريبر DataPoint المختصر/الشاذ كان يمرّ نصف مترجم (هجوم
    المشرف الحي). الحارس: النمط المرن + شبكة الأمان يمسكان كل العائلة."""
    import silk_render as _r
    cases = [
        "DataPoint(value=None, confidence=0.0)",
        "DataPoint(value='12.5', source='comtrade', confidence=0.9, "
        "note='ok (x)', retrieved_at='2026', status='ok')",
        "DataPoint(confidence=0.5, value=None)",
        "قبل DataPoint(value=None, confidence=0.0) بعد",
    ]
    for c in cases:
        out = _r._strip_internal_plumbing(c)
        assert "DataPoint" not in out and "confidence" not in out and \
               "درجة الثقة=" not in out, f"leak: {c!r} -> {out!r}"
    assert _r._strip_internal_plumbing(cases[1]).strip().startswith("12.5")


def _guard_vendor_name_leak():
    """LESSONS ١٨ — اسم مزوّد داخلي (Volza/Explee/…) تسرّب لسطح العميل (بلاغ
    UK الحي). الحارس السلوكي: الأسطر الحرفية المسرّبة (سطر «الخطوة التالية»
    القديم + ترجمة `silk_narrative` التي تُسمّي المزوّد) تُحيَّد فعلياً بمنقِّي
    العميل فلا يبقى اسم مزوّد في مخرَج التصدير؛ وسطر next_step المُولَّد لم
    يعُد يحمل اسم مزوّد أصلاً."""
    from silk_reports import (_client_redact_text, _client_forbidden_hits,
                              _client_sanitize)
    # (١) الأسطر الحرفية المسرّبة من البلاغ الحي — كلٌّ يتسرّب أولاً ثم يُحيَّد.
    leaked = [
        "فعّل خدمة التعميق المدفوعة للتحقق من المستوردين وجهات الاتصال (Volza/Explee)",
        "إكسبلي غير متاح حالياً",
        "فولزا: لا مستوردون بالاسم مرصودون لرمز 0804 في GBR",
        "buyers via Serper and SerpApi, priced by LocalPrice",
        "seasonality from pytrends; risk news from GDELT",
    ]
    for line in leaked:
        assert _client_forbidden_hits(line), (
            f"التهيئة خاطئة — يجب أن يتسرّب أولاً: {line!r}")
        cleaned = _client_redact_text(_client_sanitize(line))
        assert not _client_forbidden_hits(cleaned), (
            f"اسم مزوّد بقي بعد التنقية: {_client_forbidden_hits(cleaned)}")
    # (٢) الحارس الصارم يملك أسماء المزوّدين لاتينيةً وعربيةً معاً — لا يعتمد
    # على المُطهِّر وحده (متغيّر مستقبلي يفلت المُطهِّر يبقى يُرفَع بصوت عالٍ).
    for v in ("Volza", "Explee", "إكسبلي", "فولزا", "LocalPrice", "Serper",
              "SerpApi", "pytrends", "GDELT"):
        hits = _client_forbidden_hits(f"مبنيّ على {v} التجارية")
        assert any(h.startswith("vendor_name") for h in hits), (
            f"اسم مزوّد ليس في قائمة الرفض الصارم (_client_assert_clean): {v}")
    # (٣) سطر next_step المُولَّد لا يحمل اسم مزوّد إطلاقاً.
    import silk_render
    from canonical_netherlands import netherlands_research_blob
    blob = netherlands_research_blob()
    blob["deep_research"]["verdict"] = {"verdict": "GO", "confidence": 0.7,
                                        "ai": {"verdict": "GO"}}
    view = silk_render.build_view(blob)
    nxt = (view.get("deep_research") or {}).get("next_step") or ""
    assert nxt and not _client_forbidden_hits(nxt), (
        f"سطر الخطوة التالية يحمل اسم مزوّد: {nxt!r}")


def _guard_export_format_contract():
    """LESSONS ١٩ — عائلة export-format-contract (بلاغ المُشرِف عند السطر): زرّ
    «تصدير التقرير» الأساسي كان موصولاً بـ`dlReport("docx")` فينزّل Word بينما
    المُسلَّم النهائي للعميل PDF غير قابل للتحرير (§3، اتفاق المالك). الحارس:
    (١) زرّ PDF موصول بـ`dlReport("pdf")` لا docx؛ (٢) `dlReport` يملك فرع pdf
    بامتداد `.pdf` ورسالة 503 عربية صريحة؛ (٣) بطاقة الدردشة المصغّرة تُصدِّر
    PDF؛ (٤) الخادم يخدم report.pdf بنوع application/pdf؛ (٥) صورة النشر
    (Dockerfile) + وظيفة e2e تُثبّتان محرّك التحويل فلا يموت الزرّ حياً."""
    html = _read("web/index.html")
    # (١) الوصلة الأساسية: PDF لا docx (السطر المعطوب الأصلي غائب).
    assert '$("#pdfBtn").addEventListener("click",function(){dlReport("pdf")})' \
        in html, "زرّ PDF غير موصول بـdlReport(\"pdf\")"
    assert '$("#pdfBtn").addEventListener("click",function(){dlReport("docx")})' \
        not in html, "زرّ PDF لا يزال موصولاً بتنزيل docx (البلاغ الأصلي)"
    # زرّ Word ثانوي حاضر (النسخة القابلة للتحرير للمشغّل، لا العميل).
    assert 'id="wordBtn"' in html and \
        '$("#wordBtn").addEventListener("click",function(){dlReport("docx")})' \
        in html, "زرّ Word الثانوي غائب أو غير موصول"
    # (٢) فرع pdf في dlReport: امتداد .pdf + رسالة 503 العربية الصريحة.
    assert 'kind==="pdf"' in html, "dlReport بلا فرع pdf"
    assert '"سِلك_تقرير_"+id+".pdf"' in html, "اسم/امتداد ملف الـPDF خاطئ"
    assert "محرّك التحويل غير متاح — جرّب Word مؤقتاً" in html, \
        "رسالة 503 العربية الصريحة غائبة"
    assert "r.status===503" in html, "فرع pdf لا يعالج 503 صراحةً"
    # (٣) بطاقة الدردشة المصغّرة تُصدِّر PDF لا docx.
    assert 'data-act="pdf"' in html, "بطاقة الدردشة المصغّرة لا تُصدِّر PDF"
    assert 'this.dataset.act==="board"?nav("board"):dlReport("pdf")' in html, \
        "معالج بطاقة الدردشة لا يستدعي dlReport(\"pdf\")"
    # (٤) الخادم يخدم report.pdf بنوع application/pdf.
    api = _read("api.py")
    assert "/analyses/{analysis_id}/report.pdf" in api and \
        'media_type="application/pdf"' in api, "نقطة نهاية report.pdf غائبة/خاطئة"
    # (٥) محرّك التحويل مثبَّت على النشر وفي وظيفة e2e — كي يعمل الزرّ حيّاً لا
    # في CI فقط (البند ٦). صورةُ النشر صارت apps/api/Dockerfile في هذا المستودع
    # الأحادي (حُذِف Dockerfile المحرّك القديم، تدقيق C6) — الحارس يتبع المحرّك
    # إلى صورته الحيّة الجديدة.
    dockerfile = _read("../../../apps/api/Dockerfile")
    assert "libreoffice-writer" in dockerfile, \
        "محرّك تحويل PDF غير مثبَّت في صورة النشر — الزرّ سيموت حياً بـ503"
    e2e = _read(".github/workflows/e2e-live-shape.yml")
    assert "libreoffice-writer" in e2e, \
        "وظيفة e2e لا تثبّت محرّك التحويل — تأكيد %PDF سيفشل"
    # التدفّق يؤكّد توقيع %PDF على المسار الأساسي.
    flow = _read("tests/e2e/live_shape_flow.cjs")
    assert 'pdfBuf[0] === 0x25 && pdfBuf[1] === 0x50' in flow, \
        "تدفّق e2e لا يؤكّد توقيع %PDF لزرّ PDF"


def _guard_world_tier2_no_fabrication():
    """LESSONS ٢٠ — عائلة tier2-fabrication (تصميم الميزة أ، قفل استباقي): توسيع
    الترتيب لكل دول العالم يجب ألّا يختلق قيمة فئة-٢ ولا يفجّر ميزانية كومتريد.
    الحارس (قراءة مصدر + سلوك حيّ): (١) وحدة الترتيب لا تقرأ أيّ CSV محلّي؛
    (٢) الفئة-٢ تحمل الوسم التعاقدي + فجوتَي موقع السعودية/المنافسة معلنتين؛
    (٣) نداء العالم الواحد + التدهور عند نفاد الميزانية موجودان؛ (٤) ملف القفل
    قائم."""
    src = _read("silk_market_ranker.py")
    # (١) لا CSV محلّي في وحدة الترتيب إطلاقاً.
    for forbidden in ("agreements_l1", "demographics_l1", "market_locale",
                      "muslim_share", "requirements_l1"):
        assert forbidden not in src, f"الترتيب يقرأ CSV محلّياً: {forbidden}"
    # (٢) الوسم التعاقدي + الفجوة المعلنة + المسجّل + الصمّام.
    for needle in ('TIER2_LABEL = "تغطية أساسية — بيانات محلية محدودة"',
                   'def _tier2_gather_row', 'status="tier2_gap"',
                   'def _world_markets_enabled', 'def world_import_totals',
                   'def _comtrade_budget_left'):
        assert needle in src, f"علامة إنفاذ الفئة-٢ مفقودة: {needle}"
    # (٣) نداء العالم الواحد (partner=0) مشترك للفئتين + تدهور الميزانية.
    assert 'flow="M", partner=0' in src, "نداء العالم الواحد (partner=0) غائب"
    assert '_comtrade_budget_left()' in src and '_WORLD_BUDGET_RESERVE' in src, \
        "فرع التدهور عند نفاد الميزانية غائب"
    # (٤) ملف القفل قائم بأقفاله السبعة.
    assert _exists("tests/test_world_coverage_tierA.py"), "ملف قفل الميزة أ مفقود"
    lock = _read("tests/test_world_coverage_tierA.py")
    for fn in ("test_tier_separation_and_labels",
               "test_tier2_never_carries_a_local_csv_value",
               "test_tier2_gather_makes_zero_comtrade_calls",
               "test_budget_exhausted_degrades_to_tier1_only",
               "test_ranking_is_deterministic_on_fixture"):
        assert f"def {fn}" in lock, f"قفل الميزة أ مفقود: {fn}"


def _guard_out_of_coverage_thin_study():
    """LESSONS ٢٢ — عائلة out-of-coverage-thin-study (مواصفة المالك، الميزة أ):
    سوقٌ خارج التغطية يجب ألّا يشغّل دراسةً هزيلة بل يُعاد برسالةٍ صادقة ويُسجَّل
    إشارةَ طلب. الحارس (قراءة مصدر): البوّابة + الرسالة الحرفية + التسجيل +
    تسطيح الواجهة + ملف القفل."""
    api = _read("api.py")
    assert "def _market_in_coverage" in api, "دالّة فحص التغطية غائبة"
    assert '"error": "out_of_coverage"' in api, "بوّابة خارج التغطية غائبة"
    assert "هذه السوق خارج التغطية الحالية" in api and \
        "تواصل معنا لإضافتها" in api, "الرسالة الصادقة الحرفية غائبة"
    assert '"out_of_coverage_demand"' in api, "تسجيل إشارة الطلب غائب"
    assert "_world_markets_enabled()" in api, "البوّابة غير مقيّدة بالصمّام"
    html = _read("web/index.html")
    assert "x.message||x.reason||x.error" in html, \
        "الواجهة لا تُسطّح رسالة detail (لن تظهر رسالة خارج التغطية)"
    assert _exists("tests/test_out_of_coverage_guard.py"), "ملف قفل البوّابة مفقود"
    lock = _read("tests/test_out_of_coverage_guard.py")
    for fn in ("test_out_of_coverage_market_returns_honest_message_and_logs_demand",
               "test_tier1_curated_market_is_always_covered",
               "test_flag_off_no_coverage_guard_any_country_works_todays_way"):
        assert f"def {fn}" in lock, f"قفل البوّابة مفقود: {fn}"


def _guard_intake_no_silent_guess():
    """LESSONS ٢١ — عائلة intake-silent-guess (تصميم الميزة ب، قفل استباقي):
    استقبال المنتج من صورة يجب ألّا يختلق اسماً ولا يبدأ تحليلاً قبل تأكيد
    المستخدم، والمحوّل أماميّ معزول عن طبقات التحليل. الحارس (قراءة مصدر):
    (١) عقد عدم الاختلاق (فرع readable/العتبة => تعذّر قراءة صادق)؛ (٢) حدود
    الصورة + التقييس + العزل؛ (٣) القياس (حجز واحد) في نقطة النهاية؛ (٤) المحوّل
    لا يستورد/يستدعي طبقات التحليل؛ (٥) ملف القفل قائم."""
    import ast as _ast
    src = _read("silk_product_intake.py")
    # (١) عقد عدم الاختلاق + الرسالة الموحّدة + العتبة.
    for needle in ('READ_FAILED_MSG = "تعذّرت القراءة — اكتب الاسم يدوياً"',
                   'def _read_failed', 'def intake_image', 'readable',
                   '_MIN_CONFIDENCE', 'def enabled'):
        assert needle in src, f"علامة إنفاذ الاستقبال مفقودة: {needle}"
    # (٢) حدود الصورة + التقييس + العزل.
    for needle in ('MAX_IMAGE_BYTES', 'ALLOWED_MEDIA_TYPES', 'def _decode_and_check',
                   'def _sanitize', 'def _isolate', '_MAGIC'):
        assert needle in src, f"علامة سلامة الصورة مفقودة: {needle}"
    # (٣) القياس — نقطة النهاية تحجز تفعيلة واحدة كأيّ نداء مدفوع.
    api = _read("api.py")
    assert 'def _intake_vision_allowed' in api and \
        'try_reserve_paid_calls(1)' in api, "قياس نداء الرؤية غائب"
    assert '@app.post("/products/intake")' in api, "نقطة نهاية الاستقبال غائبة"
    assert 'intake.enabled()' in api, "صمّام SILK_IMAGE_INTAKE غير مفحوص"
    # (٤) المحوّل أماميّ معزول — لا يستورد أيّ طبقة تحليل، ولا يستدعيها نصّاً.
    tree = _ast.parse(src)
    imported = {n.names[0].name.split(".")[0] for n in _ast.walk(tree)
                if isinstance(n, _ast.Import)}
    imported |= {(n.module or "").split(".")[0] for n in _ast.walk(tree)
                 if isinstance(n, _ast.ImportFrom)}
    forbidden = {"silk_engine", "silk_missions", "silk_market_analyst",
                 "silk_ai_judge", "silk_market_ranker", "correlation",
                 "silk_synthesis", "silk_llm_runtime"}
    assert imported.isdisjoint(forbidden), imported & forbidden
    for banned in ("analyze(", "deep_research(", "write_reviewed_report",
                   "ResearchManager", "rank_markets("):
        assert banned not in src, f"الاستقبال يمسّ مسار التحليل: {banned}"
    # (٥) ملف القفل قائم بأقفاله المركزية.
    assert _exists("tests/test_product_intake_featureB.py"), "ملف قفل الميزة ب مفقود"
    lock = _read("tests/test_product_intake_featureB.py")
    for fn in ("test_low_confidence_or_unreadable_never_fabricates",
               "test_intake_module_imports_no_pipeline_code",
               "test_endpoint_image_call_is_metered_from_the_cap",
               "test_image_validation_rejects_bad_inputs"):
        assert f"def {fn}" in lock, f"قفل الميزة ب مفقود: {fn}"


def _guard_unresolved_hs_silent_spend():
    """LESSONS ٢٣ — حادثة الفيتوتشيني: دراسةٌ مدفوعةٌ بدأت برمز HS غير محسوم.
    الحارس السلوكي: (١) المُصنِّف لا يختلق (منتجٌ مجهول => منتقٍ يدوي، hs6=None،
    ثقة 0.0)؛ (٢) `_validate` يرفض فصلًا مستبعَدًا وما ليس رمزًا (عقد عدم اختلاق)؛
    (٣) بوّابة `unresolved_hs` موجودةٌ وتسبق حجز الدولار في `/research`."""
    import silk_hs_classifier as hsc
    out = hsc.classify("qwxzptvbmzzz منتج لا وجود له", allow_claude=False)
    assert out["hs6"] is None and out["status"] == "manual" and \
        out["confidence"] == 0.0, out
    assert hsc._validate({"hs6": "270900", "confidence": 0.9}) is None  # فصل ٢٧
    assert hsc._validate({"hs6": "زائف", "confidence": 0.9}) is None    # ليس رمزًا
    api = _read("api.py")
    assert "def _require_hs6" in api and '"error": "unresolved_hs"' in api, \
        "بوّابة hs6 الصلبة غائبة"
    assert "def classify_hs" in api and "def _classify_general_allow_claude" in api, \
        "نقطة/حارس التصنيف غائبة"
    gate = api.index('"error": "unresolved_hs"')
    reserve = api.index("try_reserve_usd(_expected_usd)")
    assert gate < reserve, "بوّابة hs6 يجب أن تسبق حجز الدولار (لا إنفاق على رمز مجهول)"


def _guard_hardcoded_product_rule():
    """LESSONS ٢٤ — الحارسان (مُصنِّف HS + استشارة بلد المنشأ) قاعدتان مبنيّتان
    على البيانات لا حالتا منتج (نفس عائلة «التمور السعودية»). الحارس: (١) منطقهما
    يخلو من أيّ منتج/ISO/HS من العيّنات، والعتبة config-driven؛ (٢) سلوكيًا القاعدة
    تُعمَّم من ترتيب البيانات — عيّنةٌ مُرقَّعةٌ صناعيّةٌ تُطلق/تصمت بالعتبة."""
    import inspect
    import unittest.mock as _mock
    import silk_hs_classifier as hsc
    import silk_market_ranker as ranker
    blob = inspect.getsource(hsc)
    for fn in (ranker.world_export_totals, ranker.top_world_exporters,
               ranker.is_top_world_exporter, ranker._producer_advisory_topn):
        blob += "\n" + inspect.getsource(fn)
    for tok in ("معكرونة", "pasta", "fettuccine", "تمور", "dates", "olive",
                "عسل", "honey", "ITA", "ESP", "GBR", "ARE",
                "190219", "150910", "080410", "040900"):
        if tok.isascii():
            assert not re.search(r"(?<![A-Za-z0-9])" + re.escape(tok)
                                 + r"(?![A-Za-z0-9])", blob), \
                f"ترميزٌ صلبٌ في منطق الحارس: {tok}"
        else:
            assert tok not in blob, f"ترميزٌ صلبٌ في منطق الحارس: {tok}"
    assert "SILK_PRODUCER_ADVISORY_TOPN" in blob, "العتبة ليست config-driven"

    # سلوكي: القاعدة من البيانات — رموزٌ صناعيّةٌ بحتة (لا اسم حقيقي).
    def _fake(hs_code, year):
        return [{"iso3": c, "m49": "0", "total_usd": 9 - i}
                for i, c in enumerate(["XXA", "XXB", "XXC"])]
    with _mock.patch.object(ranker, "world_export_totals", side_effect=_fake):
        top, _l = ranker.is_top_world_exporter("AAAAAA", "XXA", 2023, 2)
        bot, _l2 = ranker.is_top_world_exporter("AAAAAA", "XXC", 2023, 2)
    assert top is True and bot is False, "القاعدة لا تتبع ترتيب البيانات"


def _guard_g41_domestic_production():
    """LESSONS ٦٤ — حارسُ المعقولية يقرأ الإنتاجَ المحليّ من البروفايل
    (DEF-1/G4.1). الحارس: (١) سوقٌ مُنتِجة (نيجيريا) لا تُوسَم؛ (٢) قطر (لا
    إنتاجٍ محلّيّ) تبقى مضبوطة (لا انحدار)؛ (٣) الإعفاءُ مرئيٌّ في المانيفست
    («guard_relaxed_domestic_producer») لا صامت."""
    import silk_plausibility as P

    def _blob(iso3, market_usd, imports_usd):
        return {"market": {"iso3": iso3}, "hs_code": "200811",
                "deep_research": {"missions": {"m": {"findings": [
                    {"value": imports_usd, "source": "UN Comtrade",
                     "note": f"إجمالي استيراد {iso3} من العالم"},
                    {"value": market_usd, "source": "ويب",
                     "note": "حجم السوق الكامل"}]}}}}

    nga = _blob("NGA", "497 مليون دولار", "7,000,000 دولار")
    assert P.check_magnitudes(nga) == [], "سوقٌ مُنتِجة (نيجيريا) وُسِمت زوراً"
    P.annotate(nga)
    exempt = (nga.get("deep_research") or {}).get("plausibility_exemptions")
    assert exempt and exempt[0].get("kind") == "guard_relaxed_domestic_producer", \
        "الإعفاءُ يجب أن يُسجَّل في المانيفست (لا صامت)"
    qat = _blob("QAT", "497 مليون دولار", "7,000,000 دولار")
    assert P.check_magnitudes(qat), "قطر (لا إنتاج) يجب أن تبقى مضبوطة — انحدار!"


def _guard_bloc_list_single_source():
    """LESSONS ٦٣ — عضويةُ الكتلة التجارية من مصدرٍ واحدٍ لا تتشعّب (DEF-2).
    الحارس: (١) `silk_blocs.EU27` كاملةٌ ٢٧؛ (٢) كلُّ مستهلكٍ هو الكائنُ نفسُه
    بالهُويّة (`is`) فلا نسخةَ قد تسقط أعضاءً؛ (٣) لا مستهلكٍ يُعيد تعريفَ مجموعةٍ
    خامّ (يستورد المصدرَ الواحد)؛ (٤) سلوكيًا عضوٌ كان غائباً (المجر) ينال الطبقةَ
    الكاملة ويطابق بندَ EU."""
    import inspect
    import silk_blocs
    import silk_requirements_agent as reqs
    import silk_tariffs_agent as tariffs
    import silk_eurostat_agent as euro

    assert len(silk_blocs.EU27) == 27, "EU27 ليست ٢٧ عضواً"
    for iso in ("HUN", "ROU", "BGR", "HRV", "CYP", "EST",
                "LVA", "LTU", "LUX", "MLT", "SVK", "SVN"):
        assert iso in silk_blocs.EU27, f"عضوٌ غائبٌ عن EU27: {iso}"

    assert reqs._EU is silk_blocs.EU27, "الاشتراطات لا تشير للمصدر الواحد"
    assert tariffs._EU_ISO3 is silk_blocs.EU27, "التعريفة لا تشير للمصدر الواحد"
    assert reqs._GCC is silk_blocs.GCC and tariffs._GCC_MEMBERS is silk_blocs.GCC
    assert euro.EU_EFTA_MARKETS == silk_blocs.EU27 | silk_blocs.EFTA

    for mod in (reqs, tariffs, euro):
        assert "silk_blocs" in inspect.getsource(mod), \
            f"{mod.__name__} لا يستورد المصدر الواحد"

    # سلوكي: المجر (كانت غائبةً) تنال «مقنّن بالكامل» وتطابق بندَ EU.
    tier, _n = reqs.codification_tier("HUN")
    assert tier == "مقنّن بالكامل", "المجر سقطت للطبقة الجزئية"
    eu_row = {"market": "EU", "category": "all", "direction": "import"}
    assert reqs._matches(eu_row, "HUN", "all", "import", animal=False)


def _guard_wrong_direction_study():
    """LESSONS ٢٥ — عائلة wrong-direction-study (Wave 1.5، A): استشارةُ بلد
    المنشأ تُعمَّم لأشقّائها. الحارس السلوكي: (١) تصدير إلى بلد المنشأ نفسه =>
    self_origin (config-driven عبر env)؛ (٢) فصلٌ مقيَّد من مرجع المالك؛
    (٣) البوّابة في api؛ (٤) صفر ISO/HS مكتوب صلبًا في منطق المطابقة."""
    import silk_prerun as sp
    import os as _os
    old = _os.environ.get("SILK_ORIGIN_ISO3")
    _os.environ["SILK_ORIGIN_ISO3"] = "SAU"
    try:
        assert any(a["kind"] == "self_origin"
                   for a in sp.sibling_advisories("080410", "SAU"))
        assert not any(a["kind"] == "self_origin"
                       for a in sp.sibling_advisories("080410", "ITA"))
    finally:
        if old is None:
            _os.environ.pop("SILK_ORIGIN_ISO3", None)
        else:
            _os.environ["SILK_ORIGIN_ISO3"] = old
    # فصلٌ مقيَّد من المرجع (خنزير في سوقٍ خليجية) — عضوٌ من العائلة.
    assert any(a["kind"] == "restricted_chapter"
               for a in sp.sibling_advisories("020329", "SAU"))
    api = _read("api.py")
    assert '"error": "prerun_advisory"' in api and "advisories_ack" in api, \
        "بوّابة أشقّاء الاستشارة غائبة"
    # صفر رمز HS/دولة مكتوب صلبًا في منطق المطابقة.
    import inspect
    blob = "\n".join(inspect.getsource(fn) for fn in (
        sp.sibling_advisories, sp._restricted_hits))
    assert not re.search(r"(?<!\d)\d{4,6}(?!\d)", blob), "رمز HS صلب في المطابقة"
    assert not re.search(r'"[A-Z]{3}"', blob), "رمز دولة صلب في المطابقة"


def _guard_silent_external_failure():
    """LESSONS ٢٦ — عائلة silent-external-failure (Wave 1.5، C): فشلُ خدمةٍ
    خارجية مُهيَّأة يُعلَن للمشغّل. الحارس السلوكي: (١) record_service_failure
    يكتب صفَّ service_failure؛ (٢) المكشطة تُعلِن فشلها؛ (٣) جدول التدقيق قائم."""
    import silk_ops_log
    import tempfile as _tf
    import unittest.mock as _mock
    with _tf.TemporaryDirectory() as td:
        path = os.path.join(td, "ops.db")
        with _mock.patch.object(silk_ops_log, "_db_path", lambda: path):
            silk_ops_log.record_service_failure("comtrade", "429 rate limited")
            rows = silk_ops_log.last_errors(5, path)
    assert rows and rows[0]["kind"] == "service_failure" and \
        rows[0]["context"]["service"] == "comtrade"
    assert "record_service_failure" in _read("silk_gmaps.py"), \
        "المكشطة لا تُعلِن فشلها للمشغّل"
    assert _exists("docs/EXTERNAL_SERVICES_FAILURE_AUDIT.md"), "جدول التدقيق مفقود"


def _guard_readiness_before_spend():
    """LESSONS ٢٧ — عائلة spend-before-knowing (Wave 1.5، D): لوحةُ الجاهزية
    تعرض كلَّ تدهورٍ قبل الحجز. الحارس: نقطة `/research/readiness` + المُركِّب
    `_readiness_checks` (مع can_run/blocking) قائمان، والصمّام config-driven."""
    api = _read("api.py")
    assert "def _readiness_checks" in api and "def research_readiness" in api, \
        "لوحة الجاهزية (نقطة/مُركِّب) غائبة"
    assert '"/research/readiness"' in api and '"can_run"' in api and \
        '"blocking"' in api, "عقد لوحة الجاهزية غير مكتمل"
    import silk_prerun
    assert hasattr(silk_prerun, "advisories_enabled")


def _guard_leads_table_hygiene():
    """LESSONS ٢٨ — عنقود أوّل PDF: جدولُ روابط العميل نُقِّي عند الحدّ. الحارس
    السلوكي على المدوّنة القانونية (فيتوتشيني): جغرافيا خاطئة/نثر/حشو تُسقَط،
    الصالح يبقى، وسطر الإخلاء بارامتري بالمنتج (لا «التمور السعودية»)."""
    import silk_render
    import silk_reports
    from canonical_fettuccine import fettuccine_research_blob
    md = silk_reports.render_markdown(
        silk_render.build_view(fettuccine_research_blob()))
    seg = md[md.find("قائمة مستوردين"):]
    assert "Pastificio Milano" in seg          # صالح — يبقى
    assert "NutsWorld" not in seg              # جغرافيا أمريكية — يُسقَط
    assert "Italy imports a significant" not in seg   # نثر — يُسقَط
    assert "Anonimo Distribuzione" not in seg  # حشو — يُسقَط
    assert "فيتوتشيني" in seg and "التمور السعودية" not in md


def _guard_report_arabic_shape_a4():
    """LESSONS ٢٩ — العلامة «سِلك» كُسِرت «ِس لك» + الصفحة Letter لا A4. الحارس
    السلوكي: docx يحوي «سلك» متّصلة بلا كسرة، بمقاس A4 (210×297مم)."""
    import silk_render
    import silk_reports
    from canonical_fettuccine import fettuccine_research_blob
    import tempfile
    from docx import Document
    view = silk_render.build_view(fettuccine_research_blob())
    path = silk_reports.render_client_docx(
        view, os.path.join(tempfile.mkdtemp(), "r.docx"))
    doc = Document(path)
    txt = "\n".join(p.text for p in doc.paragraphs)
    for s in doc.sections:
        for hf in (s.header, s.footer):
            txt += "\n" + "\n".join(p.text for p in hf.paragraphs)
    assert "سلك" in txt and "سِلك" not in txt, "العلامة غير آمنة التشكيل"
    sec = doc.sections[0]
    assert abs(sec.page_width.mm - 210) < 1 and abs(sec.page_height.mm - 297) < 1, \
        "الصفحة ليست A4"


def _guard_client_template_no_hardcoded_product():
    """LESSONS ٣٠ — «التمور السعودية» كانت مثبَّتةً في تقرير أيّ منتج (عائلة
    hardcoded-product-rule موسَّعة للقوالب). الحارس: سطر الإخلاء بارامتري بالمنتج
    ولا يحمل اسم منتجٍ مثبَّت."""
    import inspect
    from silk_gmaps import maps_disclaimer, MAPS_DISCLAIMER
    src = inspect.getsource(maps_disclaimer)
    for tok in ("التمور", "dates", "معكرونة", "pasta"):
        assert tok not in src, f"اسم منتجٍ مثبَّت في سطر الإخلاء: {tok}"
    assert "التمور" not in MAPS_DISCLAIMER
    assert "عسل" in maps_disclaimer("عسل")   # يُشتَقّ من المنتج فعلًا


def _guard_analyze_persist_canonical_db():
    """LESSONS ٣١ — نتائج /analyze لم تكن محفوظةً في القاعدة القانونية: المحرّك
    ثبّت `db_path="data/silk.db"` النسبيّ فكتب لقرصٍ لا يقرأ منه أحد (المعرّف «1»
    ثم 404). الحارس السلوكي: مع SILK_DATA_DIR مضبوطًا، `analyze(persist=True)`
    يكتب لقاعدة `_db_path()` نفسها التي يقرأ منها `get_analysis` (بمسار افتراضي)."""
    import importlib
    tmp = tempfile.mkdtemp()
    saved = {k: os.environ.get(k) for k in ("SILK_DATA_DIR", "SILK_DB")}
    try:
        os.environ["SILK_DATA_DIR"] = tmp
        os.environ.pop("SILK_DB", None)
        import silk_engine
        import silk_storage
        importlib.reload(silk_storage)
        importlib.reload(silk_engine)
        # لا شبكة: المحرّك يتدهور لفجوات معلنة لكن الصفّ يُحفَظ ويُقرَأ.
        import unittest.mock as M
        with M.patch("requests.get", side_effect=OSError("blocked")), \
             M.patch("requests.post", side_effect=OSError("blocked")):
            result = silk_engine.analyze("شاي أخضر", persist=True)
        aid = result.get("analysis_id")
        assert aid is not None, "لم يُرفَق analysis_id رغم persist=True"
        assert silk_storage._db_path() == os.path.join(tmp, "silk.db")
        found = silk_storage.get_analysis(aid)   # path=None → _db_path()
        assert found is not None and found.get("product") == "شاي أخضر", (
            "الصفّ غير موجود في القاعدة القانونية — الجذر: كُتب لقرصٍ نسبيّ آخر")
    finally:
        for k, v in saved.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
        import silk_storage
        importlib.reload(silk_storage)


def _guard_new_source_contracts():
    """LESSONS ٣٢ — مصدرٌ جديد = نفس العقود (فجوة معلنة/ops/مخزَّن/محكوم/نظيف
    الشروط). الحارس السلوكي: (أ) IMF/WTO دون الشبكة => فجوة معلنة None/0.0 لا
    اختلاق؛ (ب) WTO بلا مفتاح => فجوة معلنة بصفر نداء شبكة؛ (ج) سلسلة التراجع
    كلا-الفشلين تُبقي مصدر WITS؛ (د) البوّابة العربية للبنك الدولي تطابق تامّ
    (لا تُحوِّل WITS)؛ (هـ) كل نطاق مُفضَّل بعثته تملك web_search (لا إعداد ميت)."""
    from unittest.mock import patch
    import silk_imf_agent as imf
    import silk_wto_tariff as wto
    import silk_tariffs_agent as tar
    import silk_missions as M
    from silk_data_layer import DataPoint, public_source_url, WORLD_BANK_AR_PORTAL

    # (أ) لا اختلاق دون الشبكة
    with patch("silk_cache.cached_get", return_value=None):
        assert imf.imf_indicator("NLD", "gdp_growth").value is None
    # (ب) WTO بلا مفتاح => صفر نداء شبكة
    saved = {k: os.environ.get(k) for k in ("WTO_TTD_API_KEY", "WTO_API_KEY")}
    try:
        os.environ.pop("WTO_TTD_API_KEY", None)
        os.environ.pop("WTO_API_KEY", None)
        with patch("silk_cache.cached_get") as cg:
            dp = wto.wto_applied_tariff("080410", "NLD")
        cg.assert_not_called()
        assert dp.value is None
    finally:
        for k, v in saved.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
    # (ج) سلسلة التراجع: كلا الفشلين => مصدر WITS يبقى
    with patch("silk_wto_tariff.wto_applied_tariff",
               return_value=DataPoint(None, "WTO TTD", 0.0, "x")), \
         patch("silk_tariffs_agent.applied_tariff",
               return_value=DataPoint(None, "World Bank WITS", 0.0, "y")):
        dp = tar.tariff_with_fallback("080410", "NLD")
    assert dp.value is None and dp.source == "World Bank WITS"
    # (د) البوّابة العربية للعميل: تطابق تامّ، WITS لا يُحوَّل
    assert public_source_url("World Bank", arabic=True) == WORLD_BANK_AR_PORTAL
    assert public_source_url("World Bank WITS", arabic=True) != WORLD_BANK_AR_PORTAL
    assert public_source_url("World Bank") == "https://data.worldbank.org/"
    # (هـ) لا نطاق مُفضَّل بعثته بلا web_search
    for key in M.PREFERRED_DOMAINS:
        assert "web_search" in M.MISSIONS[key]["allowed_tools"]


def _guard_report_quality_upgrade():
    """LESSONS ٣٢ — إصلاحُ المحرّك لا تحرير التقرير (تدقيق زبدة الفول السوداني/
    اليمن): كل عائلة عيبٍ تحريريّ صارت إنفاذًا حتميًّا. الحارس السلوكي على
    مدوّنة اليمن الإنتاجية الشكل: (١) عقد التأكيد يُعلِّم الرمز الخاطئ ولا
    يُعلِّم الصحيح؛ (٢) الرمز المُعلَّم يُعيد التأطير بملاحظةٍ واحدة + يسقف
    الثقة؛ (٣) شرطا قلب الحكم حقلان مهيكلان."""
    import silk_render as R
    from silk_hs_confirm import confirm_hs, is_flagged, CONTEXTUAL_TAG
    from tools.canonical_yemen import yemen_research_blob
    # (١) عقد التأكيد: الصفة المميّزة لا تخسر أمام كلمة ثانوية عارية.
    assert is_flagged(confirm_hs("زبدة الفول السوداني", "040510"))
    assert confirm_hs("تمور", "080410")["confirmed"] is not False
    # (٢) التأطير + سقف الثقة على المدوّنة، بملاحظةٍ واحدة (لا تكرار).
    dr = R.build_view(yemen_research_blob())["deep_research"]
    assert dr["hs_flagged"] is True
    assert dr["verdict"]["confidence"] <= 0.5
    assert sum(1 for l in dr["limits"] if CONTEXTUAL_TAG in l) == 1
    # (٣) شرطا قلب الحكم المهيكلان (حكم مراقبة).
    assert len(dr["flip_conditions"]) == 2
    assert all(c.get("closes_via") for c in dr["flip_conditions"])


def _guard_parse_provenance_not_prose():
    """LESSONS ٣٣ — حلِّل المصدر لا النثر: قاعدةُ إفصاح التقادُم تُرسى إلى
    بياناتٍ بنيوية. الحارس السلوكي: (١) `fact_year` يقرأ الوسم البنيويّ
    `year=YYYY`/`retrieved_at`؛ (٢) حقيقةٌ متقادِمة تُوسَم بأيّ صياغة؛
    (٣) رمز HS 2008 بلا حقيقة خلفه لا يُوسَم؛ (٤) «الطعام 2013» بلا حقيقة لا
    يُوسَم (لا false-positive نثريّ)."""
    import silk_render as R
    from silk_staleness import fact_year, stale_fact_years, is_stale_fact
    # (١) المصدر البنيويّ.
    assert fact_year({"value": 1, "note": "x year=2013", "retrieved_at": "2026"}) == 2013
    assert fact_year({"value": 1, "retrieved_at": "2018-12-31"}) == 2018
    assert not is_stale_fact({"value": 1, "retrieved_at": "2026-01-01"})
    # (٢) الوسم مستقلّ عن الصياغة.
    for s in ["في 2013 بلغ الدخل.", "عام 2013م.", "الدخل 2013 منخفض."]:
        assert R._STALE_TAG in R._tag_stale_years(s, {2013}), s
    # (٣) رمز HS 2008 لا يُوسَم (ليس سنة حقيقة، وليس في القائمة).
    assert R._STALE_TAG not in R._tag_stale_years("البند 2008 للمحضرات.", {2013})
    # (٤) «الطعام 2013» بلا حقيقة متقادِمة => بلا وسم (لا مطابقة داخل كلمة).
    assert R._STALE_TAG not in R._tag_stale_years("استهلاك الطعام 2013.", set())
    # (٥) القائمة تُشتَقّ من حقائق اليمن (2013/2018).
    from tools.canonical_yemen import yemen_research_blob
    ms = yemen_research_blob()["deep_research"]["missions"]
    allf = [f for v in ms.values() for f in v["findings"]]
    assert stale_fact_years(allf) == {2013, 2018}


def _guard_hs_gate_shared_choke_point_fail_safe():
    """LESSONS ٣٥ — تقرير الكويت الحيّ (زبدة الفول السوداني، 2026-07-21):
    بوّابة تأكيد HS كانت موصولة بـ/research وحده خلف صمّامٍ مُطفأ افتراضياً.
    الحارس السلوكي: (١) `gate_enabled` فشل-آمن — مفعّلة بلا أيّ متغيّر env؛
    (٢) `preflight_block` نقطة اختناق واحدة تحجب رمزاً غير مؤكَّد؛ (٣) كلا
    معالجَي `/analyze` و`/research` في api.py يستدعيانها فعلياً (لا نسخة
    مكرَّرة/مسار واحد فقط)."""
    import silk_hs_confirm as C
    saved = os.environ.pop("SILK_HS_CONFIRM_GATE", None)
    try:
        # (١) فشل-آمن: بلا أيّ ضبط => مفعّلة.
        assert C.gate_enabled() is True
        # إطفاءٌ صريح فقط يُعطّلها.
        os.environ["SILK_HS_CONFIRM_GATE"] = "0"
        assert C.gate_enabled() is False
        os.environ["SILK_HS_CONFIRM_GATE"] = "1"
        assert C.gate_enabled() is True
        del os.environ["SILK_HS_CONFIRM_GATE"]
        # (٢) نقطة الاختناق تحجب فعلياً — نفس عيّنة الحادثة الحية.
        blocked = C.preflight_block("زبدة الفول السوداني", "040510")
        assert blocked is not None and blocked["error"] == "hs_confirmation_needed"
        assert C.preflight_block("زبدة الفول السوداني", "040510",
                                 hs_confirmed=True) is None
    finally:
        if saved is None:
            os.environ.pop("SILK_HS_CONFIRM_GATE", None)
        else:
            os.environ["SILK_HS_CONFIRM_GATE"] = saved
    # (٣) كلا المعالجَين يستدعيان نقطةَ الاختناق — لا مسارٌ واحد فقط.
    # يُحتسَب الغلافُ `preflight_resolve` (البوّابة + قياسُ السمة الرقمية،
    # LESSONS ٦٥) لأنه **يستدعي `preflight_block` نفسها** لا يستبدلها —
    # وهذا مُتحقَّقٌ منه بنيوياً أدناه كي لا ينحرف الغلافُ لبوّابةٍ موازية.
    api_src = _read("api.py")
    calls = (api_src.count("preflight_block(")
             + api_src.count("preflight_resolve("))
    assert calls >= 2, (
        "نقطةُ اختناق البوّابة يجب أن تُستدعى من كلا /analyze و/research")
    import inspect
    wrapper = inspect.getsource(C.preflight_resolve)
    assert "preflight_block(" in wrapper, (
        "preflight_resolve لا يستدعي preflight_block — بوّابةٌ موازية")


def _guard_cross_market_checkpoint_leak():
    """LESSONS ٣٦ — تسرّب اليمن↔الكويت: نقاط تفتيش بعثات `/research` كانت
    تُقرأ بمفتاح analysis_id فقط بلا عمود سوق، واستئنافٌ بسوقٍ مختلف يُعيد
    استهلاكها بصمت. الحارس السلوكي: (١) نقطة تفتيش مختومة بسوقٍ (اليمن) لا
    تُعاد لطلبٍ بسوقٍ آخر (الكويت)؛ (٢) صفوفٌ قديمة بلا ختم لا تُحجَب؛
    (٣) بوّابة `/research`'s resume_market_mismatch (٤٠٩) موجودة في api.py
    **قبل** فرع «مكتملة => أعِدها كما هي» (لا إرجاعٌ صامتٌ يتجاهل الطلب)."""
    import silk_storage
    from silk_agents import AgentReport
    import tempfile as _tf
    db = os.path.join(_tf.mkdtemp(), "silk.db")
    yemen_report = AgentReport(agent_name="x", findings=[], failed=False,
                               summary="سوق عدن المركزي / ربوع")
    silk_storage.save_mission_checkpoint(1, "consumer_culture", yemen_report,
                                         path=db, market_iso3="YEM")
    # (١) طلبٌ بسوق آخر لا يستلم الصفّ.
    assert "consumer_culture" not in silk_storage.load_mission_checkpoints(
        1, path=db, market_iso3="KWT")
    assert "consumer_culture" in silk_storage.load_mission_checkpoints(
        1, path=db, market_iso3="YEM")
    # (٢) صفٌّ قديم بلا ختم (market_iso3=None) لا يُحجَب.
    old_report = AgentReport(agent_name="y", findings=[], failed=False, summary="s")
    silk_storage.save_mission_checkpoint(2, "tradeflow", old_report, path=db)
    assert "tradeflow" in silk_storage.load_mission_checkpoints(
        2, path=db, market_iso3="KWT")
    # (٣) بوّابة API تسبق فرع الإعادة الصامتة لتشغيلةٍ مكتملة.
    api_src = _read("api.py")
    assert "resume_market_mismatch" in api_src
    gate_idx = api_src.index("resume_market_mismatch")
    completed_shortcut_idx = api_src.index(
        'if run_row.get("status") == "completed"')
    assert gate_idx < completed_shortcut_idx, (
        "بوّابة تعارض السوق يجب أن تسبق فرع «مكتملة => أعِدها كما هي»")


def _guard_golden_contract_test_exists_and_covers_both_paths():
    """LESSONS ٣٧ — الاختبار الذهبي موجودٌ فعلياً ويفحص كِلا مسارَي الدخول
    على نفس سيناريو الحادثة (زبدة الفول السوداني/الكويت)، لا مساراً واحداً."""
    assert _exists("tools/canonical_kuwait_peanut_butter.py")
    assert _exists("tests/test_golden_deep_research_contract.py")
    golden_src = _read("tests/test_golden_deep_research_contract.py")
    assert '"/analyze"' in golden_src and '"/research"' in golden_src
    assert "resume_market_mismatch" in golden_src
    smoke_src = _read("tools/post_deploy_smoke.py")
    assert "hs_confirmation_needed" in smoke_src, (
        "فحص الدخان بعد النشر يجب أن يثبت بوّابة HS حياً (Wave 3.2)")


def _guard_general_hs_classifier_no_lookup_table_ceiling():
    """LESSONS ٣٩ — عائلة `lookup-table-ceiling`: بذرة CSV تلميحٌ ابتدائي لا
    الحاكم النهائي. الحارس السلوكي: (١) بوّابة سلامة الفصل ترفض رمزاً خارج
    بنية WCO الحقيقية بمعزلٍ عن ادّعاء أيّ نموذج؛ (٢) منتجٌ محسومٌ جيداً
    («تمور») تلقائيٌّ بلا أيّ نداء كلود؛ (٣) منتجٌ مُعلَّم (زبدة الفول
    السوداني) لا يمرّ تلقائياً بلا كلود؛ (٤) نقطة الاختناق `preflight_block`
    تُلحِق `candidates` فعلياً بردّ الحجب — لا رفضٌ عارٍ بلا توجيه."""
    import silk_hs_classifier as hsc
    from silk_hs_resolver import chapter_valid
    # (١) سلامة الفصل بنيويةٌ بمعزلٍ عن مصدر الادّعاء.
    assert chapter_valid("999999") is False
    assert hsc._validated_candidate("أيّ منتج", "999999") is None
    # (٢) لا هدر — منتجٌ واثقٌ لا يستدعي كلود إطلاقاً.
    from unittest.mock import patch
    with patch("silk_ai_judge._call") as mock_call:
        r = hsc.classify_general("تمور", allow_claude=True)
    assert r["tier"] == "auto" and r["hs6"] == "080410"
    assert mock_call.called is False
    # (٣) منتجٌ مُعلَّم — الرمز اللفظي الخاطئ (040510 زبدة ألبان) **لا يفوز
    # أبداً** (القاعدة الدائمة). بعد إتاحة الرمز الصحيح 200811 في البذرة
    # (طلب المالك 2026-07-23) صار الحسم الحتمي يُنتج **العائلة الصحيحة**
    # تلقائياً — تصحيحٌ يقوّي القاعدة (لا فئةً مجاورةً خاطئةً بثقة).
    r2 = hsc.classify_general("زبدة الفول السوداني", hs_code="040510",
                              allow_claude=False)
    assert r2["hs6"] != "040510"
    assert r2["hs6"] == "200811"
    # (٤) preflight_block يُلحِق مرشّحين فعليّين بردّ الحجب.
    from silk_hs_confirm import preflight_block
    with patch.dict(os.environ, {"SILK_HS_CONFIRM_GATE": "1"}):
        blocked = preflight_block("زبدة الفول السوداني", "040510",
                                  allow_claude=False)
    assert blocked is not None and blocked.get("candidates")


def _guard_watchdog_owner_only_no_client_contamination():
    """LESSONS ٣٨ — الحارس («كاميرا مراقبة»، طلب المُشرِف): مراقبةٌ دائمة
    مملوكة للمالك حصراً بلا أيّ تلوّث لسطح العميل. الحارس السلوكي:
    (١) نقطة استدعاءٍ واحدة مشتركة يُستدعاها كلا `/analyze` و`/research`
    (نفس معيار البند ٣٥: عدّ استدعاءات `_attach_watchdog(` ≥ ٣ — التعريف
    + نداءان)؛ (٢) `silk_render.py`/`silk_reports.py` (طبقتا العرض/التصدير
    التي يراها العميل) لا تستوردان `silk_watchdog` إطلاقاً؛ (٣) `observe()`
    لا يعدّل نتيجة التحليل الممرَّرة إليه؛ (٤) عطلٌ داخلي في الحارس لا يرفع
    استثناءً أبداً — يُعاد سجلٌّ يحمل `self_error` بدل إسقاط التحليل."""
    import inspect
    api_src = _read("api.py")
    assert api_src.count("_attach_watchdog(") >= 3, (
        "_attach_watchdog يجب أن تُستدعى من كلا /analyze و/research")
    import silk_render
    import silk_reports
    assert "silk_watchdog" not in inspect.getsource(silk_render)
    assert "silk_watchdog" not in inspect.getsource(silk_reports)
    import silk_watchdog
    result = {"product": "x", "view": {"deep_research": {}},
             "data_economics": {}, "market": {}}
    before = dict(result)
    silk_watchdog.observe(result, "research", analysis_id=None)
    assert result == before, "الحارس عدَّل نتيجة التحليل — خرق مبدأ عدم التلوّث"
    rec = silk_watchdog.observe(object(), "research", analysis_id=999)
    assert rec is not None and rec.get("self_error")


def _guard_ui_tier_consumption_single_choke_point():
    """LESSONS ٤٠ — بلاغ «UI-ONLY FIX» (المُشرِف): نقطة اختناق التصنيف
    (`res.tier` من `/classify_hs`) لها موقعُ استهلاكٍ واحدٌ في الواجهة، لا
    مسارٌ ثانٍ يثق بـhs6 خامًا. الحارس السلوكي: (١) شارة «✓ صُنّف تلقائياً»
    نصٌّ حرفيٌّ ظهورهُ الوحيد داخل `ensureHs` مشروطًا بـ`tier==="auto"`؛
    (٢) معالجا نقر صفّ الفهرس (`#pDrop`) وتأكيد استخلاص الصورة (`#intakeGo`)
    يمرّان عبر `ensureHs` بدل ضبط الحسم مباشرةً؛ (٣) نصّ الشارة المشترك
    (`resolvedAs`) لم يعد يحمل ادّعاء «صُنّف تلقائياً» بذاته — وإلا يظهر على
    أيّ تأكيدٍ يدويّ (اختيار مرشّح، إدخال يدويّ) رغم أنه ليس تلقائيًا فعلاً."""
    html = _read("web/index.html")
    badge = "✓ صُنّف تلقائياً"
    assert html.count(badge) == 1, (
        f"شارة «{badge}» ظهرت {html.count(badge)} مرّة — يجب أن تكون نقطة "
        "انطلاقٍ واحدة فقط داخل ensureHs")
    ensure_hs_start = html.index("function ensureHs(")
    ensure_hs_body = html[ensure_hs_start:html.index("function _pct(", ensure_hs_start)]
    assert badge in ensure_hs_body and 'res.tier==="auto"' in ensure_hs_body
    pdrop_start = html.index('$("#pDrop").addEventListener("click"')
    assert "ensureHs(function(){})" in html[pdrop_start:pdrop_start + 1000]
    intake_go_start = html.index('$("#intakeGo").addEventListener("click"')
    assert "ensureHs(function(){})" in html[intake_go_start:intake_go_start + 1000]
    # نصّ الشارة المشتركة نفسه بلا ادّعاء «تلقائي» — وإلا تظهر على أيّ تأكيدٍ
    # يدويّ (اختيار مرشّح/إدخال يدويّ) عبر إعادة استعمال t("resolvedAs").
    resolved_as_start = html.index("resolvedAs:{")
    resolved_as_line = html[resolved_as_start:resolved_as_start + 120]
    assert "صُنّف تلقائياً" not in resolved_as_line, (
        "resolvedAs المشتركة تحمل ادّعاء «صُنّف تلقائياً» — تُعيد ظهور الشارة "
        "على مساراتٍ يدويةٍ غير محسومة (نفس عائلة الحادثة)")


def _guard_active_resolution_beats_rejected_and_short_root_collision():
    """LESSONS ٤١ — «ONE FIX» (المُشرِف): رفضٌ بلا بديلٍ صحيحٍ مأزقٌ لا حَل؛
    التاجر لا يعرف رموز HS ولا يجوز أن يُطلَب منه ذلك. الحارس السلوكي:
    (١) مرشّح كلود المصادَق يتصدّر على مرشّحٍ حتميٍّ مرفوضٍ (تداخلٌ دون
    العتبة) رغم بقاء الأخير «مُتحقَّقاً» لمجرّد وجوده في بذرتنا الجزئية —
    نفس بلاغ «زبدة الفول السوداني»، منتجٌ مختلف؛ (٢) احتواء جذرٍ من حرفين
    («بن» داخل «بنكهة»/«جبن») لا يُحتسَب تداخلاً حقيقياً — نواة المطابقة
    ترفض التصادف اللفظي القصير بمعزلٍ عن تدفّق المصنِّف بأكمله."""
    import silk_hs_classifier as hsc
    from silk_hs_confirm import _covered
    # (١) الترتيب: مرشّحٌ حتميٌّ مرفوضٌ (متحقَّق، تداخلٌ ضعيف) لا يتصدّر على
    # مرشّح كلود (غير متحقَّق لكنه عابرٌ للعتبة وأعلى تداخلاً) — نفس السيناريو
    # الذي كان يُبقي الرمز المرفوض معروضاً كخيارٍ أساسيّ بلا بديل.
    rejected = {"hs6": "040510", "overlap": 0.33, "verified": True,
               "model_confidence": 0.5, "source": "deterministic"}
    resolved = {"hs6": "200811", "overlap": 0.6, "verified": False,
               "model_confidence": 0.9, "source": "llm"}
    ordered = sorted([rejected, resolved], key=hsc._rank_key, reverse=True)
    assert ordered[0]["hs6"] == "200811", (
        "المرشّح المرفوض تصدّر على المرشّح المصادَق من كلود — التاجر يبقى "
        "بين تأكيد رمزٍ خاطئ وإدخال رمزٍ يجهله")
    # (٢) نواة التداخل: احتواء جذرٍ قصيرٍ (حرفان) لا يُعَدّ تطابقاً — التطابق
    # التامّ يبقى بلا قيدٍ على الطول.
    assert _covered("بنكهه", ["بن", "غير", "محمص"]) is False
    assert _covered("زبده", ["زبده"]) is True


def _guard_dza_quality_gate_six_findings():
    """LESSONS ٤٢ — «تحليل #1» (زبدة الفول السوداني/الجزائر DZA، 2026-07-21):
    ستّ نتائج فشل بوّابة الجودة معاً على تشغيلة واحدة (Markdown/تنسيق شارد،
    ثقة خام، تكرار رقم مفتاحي ×٢، عمود سعر مضلِّل، سقف الملحق التقني). رمز
    HS الخاطئ خارج نطاق هذا الحارس عمداً (يُصلَح عبر مسار مصنِّف HS العام).
    الحارس السلوكي: يعيد بناء المدوّنة الحقيقية الشكل (tools/canonical_
    dza_peanut_butter.py) ويؤكّد أن الحكم لم يعد FAIL بعد الإصلاح، وأن
    حارسي الانحدار الحقيقيين (raw_confidence/currency_label_mismatch) صفر."""
    from tools.canonical_dza_peanut_butter import dza_research_blob
    import silk_render
    import silk_quality_gate as QG
    view = silk_render.build_view(dza_research_blob())
    out = QG.run_quality_gate(view)
    assert out["verdict"] != "FAIL", f"لا يزال FAIL: {out['findings']}"
    checks = {f["check"] for f in out["findings"]}
    assert "raw_confidence" not in checks
    assert "currency_label_mismatch" not in checks
    fired = checks & QG._REGRESSION_GUARD_FIRED
    assert fired == set(), f"حارس انحدار أُطلِق رغم الإصلاح: {fired}"


def _guard_hs_classifier_valve_fail_safe_default():
    """LESSONS ٤٣ — بلاغ حيّ متكرّر (المالك): المُصنِّف العام مُصلَحٌ ومُختبَرٌ
    ليتعرَّف على الرمز الصحيح لمنتجٍ متعدِّد الصفات، لكنه كان خلف صمّامٍ
    `SILK_HS_CLASSIFIER` مُطفأٍ افتراضياً — فلا يعمل أبداً في الإنتاج ما لم
    يُضبَط صراحةً. الحارس السلوكي: (١) الصمّام مفعَّلٌ حين المتغيّر غير
    مضبوط إطلاقاً؛ (٢) يُطفَأ فقط بقيمةٍ صريحة (0/false/no/off)؛ (٣) مع
    الصمّام الافتراضي وحده (بلا أيّ ضبطٍ إضافي)، محاكاة تصنيف منتجٍ متعدِّد
    الصفات تحسم تلقائياً للفصل الصحيح لا الصفة الثانوية العارضة."""
    import os
    from unittest.mock import patch
    import silk_hs_classifier as hsc
    os.environ.pop("SILK_HS_CLASSIFIER", None)
    assert hsc.enabled() is True, "الصمّام يجب أن يكون فشلاً-آمناً (مفعَّلاً) دون ضبط"
    with patch.dict(os.environ, {"SILK_HS_CLASSIFIER": "0"}):
        assert hsc.enabled() is False
    fake = ('{"candidates":[{"hs6":"200811","description_ar":'
           '"فول سوداني محضّر أو محفوظ","reason_ar":'
           '"زبدة الفول السوداني محضّرةٌ من الفول السوداني","confidence":0.92}]}')
    with patch("silk_ai_judge.available", return_value=True), \
         patch("silk_ai_judge._call", return_value=fake), \
         patch("silk_usage.try_reserve_paid_calls", return_value=True), \
         patch("silk_usage.try_reserve_usd", return_value=True):
        r = hsc.classify_general("زبدة الفول السوداني", hs_code="040510",
                                 allow_claude=True)
    assert r["tier"] == "auto", f"لم يُحسَم تلقائياً بالإعدادات الافتراضية: {r}"
    assert r["hs6"] != "040510"
    # (٤) الصمّام مرئيٌّ عن بُعد من /health (نفس نمط persist_guard) — لا
    # اعتماد على قراءة الشيفرة لمعرفة حالته الفعلية على النشر الحيّ.
    pytest.importorskip("fastapi")
    from fastapi.testclient import TestClient
    import api
    with patch.dict(os.environ, {"SILK_API_KEY": "", "ANTHROPIC_API_KEY": "k"}):
        health = TestClient(api.create_app()).get("/health").json()
    assert health["hs_classifier"]["enabled"] is True
    with patch.dict(os.environ, {"SILK_HS_CLASSIFIER": "0",
                                 "ANTHROPIC_API_KEY": "k"}):
        health_off = TestClient(api.create_app()).get("/health").json()
    assert health_off["hs_classifier"]["enabled"] is False
    assert any("SILK_HS_CLASSIFIER" in w
              for w in (health_off.get("warnings") or [])), (
        "تعطيلٌ صريحٌ للصمّام مع مفتاح كلود متاح يجب أن يظهر تحذيراً في /health")


def _guard_verdict_tone_recognizes_arabic_labels():
    """LESSONS ٤٤ — Master Prompt Part 2 §B: بوابة اتساق الحكم عند التسليم
    كشفت أنّ `silk_render._verdict_tone` كانت تتعرّف على الرموز الإنجليزية
    فقط (GO/WATCH/CONDITIONAL/NO-GO)، فأيّ مسارٍ يضع التسمية العربية مباشرةً
    (`"دخول مشروط"` لا `"CONDITIONAL-GO"`) كان ينهار إلى tone="unknown"
    فتعرض الشارة «تعذّر إصدار توصية» بينما جدول/متن التقرير يذكران التسمية
    الصحيحة — تناقضٌ شارة/متن. الحارس السلوكي: التسمية العربية والرمز
    الإنجليزي المطابق يُنتِجان نفس الـtone؛ وبوابة اتساق التسليم (شارة/جدول/
    سطر القرار) تمرّ فعلياً على مدوّنة الكويت القانونية بلا رفعٍ."""
    from silk_render import _verdict_tone
    assert _verdict_tone("دخول مشروط") == _verdict_tone("CONDITIONAL-GO") == "conditional"
    assert _verdict_tone("مراقبة السوق") == _verdict_tone("WATCH") == "watch"
    assert _verdict_tone("عدم الدخول حالياً") == _verdict_tone("NO-GO") == "nogo"
    assert _verdict_tone("التوصية بالدخول") == _verdict_tone("GO") == "go"

    from tools.canonical_kuwait_peanut_butter import kuwait_research_blob
    from silk_render import build_view
    from silk_reports import render_docx, render_client_docx
    import os
    import tempfile
    os.environ["SILK_HERMETIC"] = "1"
    view = build_view(kuwait_research_blob())
    tmp = tempfile.mkdtemp()
    render_docx(view, os.path.join(tmp, "r.docx"))
    render_client_docx(view, os.path.join(tmp, "c.docx"))


def _guard_price_fix_scoped_to_table_window():
    """LESSONS ٤٥ — دالة الإصلاح `silk_render._fix_price_column_currency_
    label` تقتصر على نافذة الجدول نفسه (لا كامل المستند) عند البحث عن
    عملةٍ أخرى، مطابقةً لدالة الفحص الشقيقة (اللائحة ٤٢). حارسٌ مضاد: تناقضٌ
    حقيقي داخل نفس الجدول يبقى مُصلَحاً بالعملة الصحيحة."""
    from silk_render import _fix_price_column_currency_label
    unrelated_euro_elsewhere = (
        "| المنتج | السعر/كجم بالدولار |\n| --- | --- |\n| صنف | 6.0$ |\n\n"
        "## قسمٌ آخر\nخطر صرف العملة: اليورو هو عملة السوق نفسها.")
    out = _fix_price_column_currency_label(unrelated_euro_elsewhere)
    assert "السعر/كجم بالدولار" in out and "السعر/كجم باليورو" not in out

    same_table_mismatch = (
        "| المنتج | السعر/كجم بالدولار |\n| --- | --- |\n| صنف | 9.14€ |")
    out2 = _fix_price_column_currency_label(same_table_mismatch)
    assert "السعر/كجم باليورو" in out2


def _guard_quality_gate_is_client_export_delivery_condition():
    """LESSONS ٤٦ — حزمة الفكس v2.1: بوابة الجودة شرط تسليم للعميل (FAIL =>
    409) + عائلة فحوصات كاتب/عرض بنيوية تُفشِل على golden-bad. حارسٌ سلوكي:
    الفحوصات الجديدة تُطلِق فعلياً على مدخلات تُعيد إنتاج العطل، والدوال
    الحاجبة موجودة في api.py."""
    import silk_quality_gate as qg

    sections = "\n".join(
        f"## {i}. {s}\nنصّ القسم بجملة تنتهي بنقطة."
        for i, s in enumerate((
            "الخلاصة التنفيذية", "منهجية البحث ونطاقه",
            "نظرة عامة على السوق وحجمه", "ديناميكيات السوق",
            "تحليل المستهلك والطلب", "المشهد التنافسي",
            "التنظيم والوصول للسوق", "اللوجستيات وسلسلة الإمداد",
            "تقييم المخاطر", "التوصيات الاستراتيجية", "الملاحق"), 1))

    def _checks(text):
        return {f["check"] for f in qg.run_quality_gate(
            {"deep_research": {"report": {"text": text},
                              "missions": {}, "analyst": {},
                              "verdict": {"verdict": "WATCH"}}})["findings"]}

    # عيّنات golden-bad تُعيد إنتاج العطل الموصوف — كل فحص جديد يُطلِق.
    assert "hhi_false_precision" in _checks(
        sections + "\n\nمؤشر التركّز HHI = 2184.7 هنا.")
    assert "near_duplicate_figure" in _checks(
        sections + "\n\nالواردات 6,733,369 دولاراً وفي جدول 6,733,376 دولاراً.")
    assert "supplier_rank_gap" in _checks(
        sections + "\n\n#1 تونس، #2 الجزائر، #5 إيران، #6 المغرب.")
    assert "lpi_invalid_edition_year" in _checks(
        sections + "\n\nمؤشر LPI 3.2 لعام 2022 مرتفع.")
    # الدوال الحاجبة موجودة في مسار التصدير + الحارس + المُصدِّر.
    _needles("api.py", "def _block_client_export_if_gate_failed",
             "def _gate_verdict_for_client_export")()
    _needles("silk_watchdog.py", "def record_blocked_export")()
    _needles("silk_reports.py", "def _client_references_section")()


# ── حُرّاس برنامج إصلاح جودة التقارير (WP-1…WP-7، صفوف 47-53) ────────────────

def _guard_wp1_verdict_determinism():
    """صفّ ٤٧ — الحكم الحتمي هو المعروض الوحيد + temperature=0 + سُلَّم ثقة واحد."""
    _needles("silk_narrative.py", "def authoritative_verdict")()
    _needles("silk_llm_provider.py", "def _supports_sampling_params",
             "def _scrub_sampling_params", '["temperature"] = 0')()
    _needles("tests/test_wp1_verdict_determinism.py",
             "test_sampling_params_present_for_a_still_supported_model",
             "test_sampling_params_absent_for_the_repo_default_model")()
    _needles("silk_style_contract.py", "def confidence_band_label")()
    _needles("tests/test_wp1_verdict_determinism.py",
             "test_three_consecutive_renders_are_byte_identical")()


def _guard_wp2_no_raw_internal_output():
    """صفّ ٤٨ — لا نائب/سقالة/بتر يصل العميل؛ الحجب لا التسليم المشوَّه."""
    _needles("silk_reports.py", "def _client_prose",
             "def _client_missing_narrative_heads")()
    _needles("silk_ai_judge.py", "def rephrase_client_sections")()
    _needles("silk_quality_gate.py", "_check_client_scaffold_leak",
             "_check_placeholder_leak")()
    _needles("tests/test_wp2_client_output_hygiene.py",
             "test_gate_fails_on_literal_so_what_in_client_text")()


def _guard_wp3_evidence_integrity():
    """صفّ ٤٩ — شارة واعية بالمنشأ + مصالحة رقمية + تفريد مصادر مُطبَّع."""
    _needles("silk_narrative.py", "def evidence_badge_for",
             "RECONCILED_OUT_TAG")()
    _needles("silk_render.py", "def _reconcile_numeric_conflicts")()
    _needles("tests/test_wp3_evidence_integrity.py",
             "test_near_duplicate_values_reconcile_to_one_canonical")()


def _guard_wp4_gaps_consistency():
    """صفّ ٥٠ — مصدر واحد لمدخلات الفجوات الأربعة + حارس تناقض الختام."""
    _needles("silk_reports.py", "def _client_gap_inputs")()
    _needles("silk_quality_gate.py", "_check_gaps_closing_contradiction")()
    _needles("tests/test_wp4_gaps_consistency.py",
             "test_gate_fails_on_closing_contradiction")()


def _guard_wp5_rtl_bracket_isolation():
    """صفّ ٥١ — عزل RLM قبل _finalize_rtl + فحص أقواس آلي على الـPDF."""
    _needles("silk_reports.py", "def _bidi_isolate_brackets",
             "def count_suspicious_brackets", "def _pdf_bracket_check")()
    _needles("tools/rtl_calibration.py", "def build_bracket_fixture")()
    _needles("tests/test_wp5_rtl_brackets.py",
             "test_pdf_bracket_check_fails_export_above_threshold")()


def _guard_wp6_injector_adversarial_locks():
    """صفّ ٥٢ — حاقنا §D-1/§D-2 مقفولان بجُمل التقارير المُسلَّمة."""
    _needles("silk_render.py", "def _already_explained_nearby",
             "def _year_in_growth_span")()
    _needles("tests/test_wp6_injector_hardening.py",
             "test_delivered_sentence_growth_span_year_not_tagged_stale",
             "test_delivered_sentence_dash_explained_cagr_not_redefined")()


def _guard_wp7_delivery_gate_hardening():
    """صفّ ٥٣ — تجاوز بسلطة مالك منفصلة + بوابة نصّ المُنتَج النهائي."""
    _needles("api.py", "owner_override_required", "X-Owner-Key")()
    _needles("silk_watchdog.py", "def record_override",
             "def override_records_for")()
    _needles("silk_quality_gate.py", "def run_client_artifact_text_gate")()
    _needles("tests/test_wp7_delivery_gate_hardening.py",
             "test_artifact_text_gate_catches_all_leak_classes")()


def _guard_zero_confidence_finding_declared_gap():
    """LESSONS ٥٤ — بند بعثة قيمته غير فارغة بثقة 0.0 (خرق حارس المراقبة الحي
    على demand_trends): ادعاء بثقة صفرية — مصرَّحاً بها أو موروثة من نقطة فجوة
    مستشهَد بها — يُعلَن فجوة في gaps لا يُشحَن بنداً أبداً."""
    import json as _json

    import silk_llm_runtime as _rt
    from silk_data_layer import DataPoint as _DP
    reg = {"gap1": _DP(None, "FAOSTAT", 0.0, "401 — فجوة معلنة", "2026-07-23")}
    text = _json.dumps({"findings": [
        {"claim": "ادعاء صفري الثقة", "datapoint_ids": ["gap1"],
         "confidence": 0.0}], "gaps": [], "summary": ""}, ensure_ascii=False)
    out = _rt._parse_output(text, reg)
    assert out["findings"] == [], "بند بثقة 0.0 شُحن بدل إعلانه فجوة"
    assert any("ادعاء صفري الثقة" in g for g in out["gaps"]), \
        "الادعاء الصفري لم يُعلَن فجوة"


def _guard_coverage_gate_year_fallback():
    """LESSON ٥٦ — بوّابة «خارج التغطية» كانت تفشل مفتوحةً دوماً (استطلاع سنة
    اليوم-١ بلا سُلَّم fallback، وكومتريد متأخّر). الحارس (قراءة مصدر + سلوك):
    السُّلَّم + المُحلِّل + السنة المشتركة موجودة، والبوّابة تستعملها، والأقفال قائمة."""
    src = _read("silk_market_ranker.py")
    for n in ("DEFAULT_STUDY_YEAR", "def coverage_year_ladder",
              "def world_import_totals_resolved"):
        assert n in src, f"علامة إنفاذ سُلَّم التغطية مفقودة: {n}"
    api = _read("api.py")
    assert "world_import_totals_resolved" in api, "البوّابة لا تستعمل السُّلَّم"
    # سلوك: السُّلَّم يبدأ من سنة اليوم-١ ويضمن سنة الدراسة في الذيل.
    import datetime as _dt
    import silk_market_ranker as _R
    ladder = _R.coverage_year_ladder()
    assert ladder[0] == _dt.date.today().year - 1, ladder
    assert _R.DEFAULT_STUDY_YEAR in ladder, ladder
    lock = _read("tests/test_out_of_coverage_guard.py")
    for fn in ("test_coverage_gate_closes_when_current_year_empty_but_study_year_full",
               "test_world_import_totals_resolved_ladders_to_first_nonempty_year"):
        assert f"def {fn}" in lock, f"قفل سُلَّم التغطية مفقود: {fn}"


def _guard_sanitizer_obfuscation_variants():
    """LESSON ٥٧ — سبع صيغ تشويش أكّد المشرف نفاذها بالتنفيذ المباشر. الحارس
    السلوكي يبني السلاسل السبع الحرفية ويؤكّد تحييد كلٍّ (المسار العام أو
    مسار العميل) — القفل بالسلسلة الحرفية (silk-operations §4). الصيغة السابعة
    (عدّ نداءات الأدوات العربية) أُعيدت في متابعة #7 المستقلّة (قرار المالك
    2026-07-23): تِلِمتري مشروع يُقرأ من الملخّص الخام، لكنه على أسطح العميل
    سباكةٌ تُجرَّد — الفكس يوفّق بينهما (تتبّع من الخام + تجريد للعرض)."""
    import silk_render as _SR
    from silk_reports import (_client_forbidden_hits, _client_redact_text,
                              _client_sanitize)

    def _gen(s):
        return _SR._strip_internal_plumbing(s)

    def _client_clean(s):
        return not _client_forbidden_hits(_client_redact_text(_client_sanitize(s)))

    # (١) stop_reason مباعَد/عارٍ بلا قيمة — المسار العام.
    assert "stop_reason" not in _gen("التوليد stop_reason =  انتهى")
    # (٢) اسم مزوّد لاتيني مباعَد «S e r p A p i» — مسار العميل.
    assert _client_clean("مبنيّ على S e r p A p i التجارية")
    # (٣) درجة ثقة بأرقام عربية-هندية «ثقة=٠٫٦٤» — المسار العام.
    o3 = _gen("التقييم ثقة=٠٫٦٤ للمصدر")
    assert "٠٫٦٤" not in o3 and "ثقة=" not in o3, o3
    # (٤) اسم مزوّد عربي مُشكَّل «إكْسبِلي» — مسار العميل.
    assert _client_clean("المصدر إكْسبِلي غير متاح")
    # (٥) «سجلات الخادم» بلا شدّة — المسار العام.
    assert "سجلات الخادم" not in _gen("خطأ داخلي راجع سجلات الخادم الآن")
    # (٦) بادئة مفتاح بعثة مرقّمة «m3_» — المسار العام.
    assert "m3_" not in _gen("أنتجت m3_pricing_scout النتيجة")
    # (٧) عدّ نداءات أدوات بأرقام عربية «نداءات أدوات: ٢» — المسار العام
    # (أُعيدت في متابعة #7 المستقلّة، قرار المالك 2026-07-23).
    assert "نداءات أدوات" not in _gen("الملخّص | نداءات أدوات: ٢")


def _guard_wave2_med_hardening():
    """LESSON ٥٩ — خمسة إصلاحات MED من تدقيق v2 (الموجة ٢). الحارس يؤكّد نقاط
    الإنفاذ الخمس (قراءة مصدر) + وجود ملف الأقفال السلوكية."""
    st = _read("silk_storage.py")
    assert "def _reconcile_leaked_usd" in st, "#3 مُصالِح الحجز المتسرّب غائب"
    assert "def reconcile_failed_run_usd" in st, "#3 مصالحة الفشل الرشيق غائبة"
    assert "status = 'failed'" in st, "#3 المكنَس لا يمسح صفوف 'failed'"
    api = _read("api.py")
    assert api.count("reconcile_failed_run_usd(analysis_id)") >= 2, (
        "#3 أحد مساري فشل /research لا يُصالِح")
    assert "begin_data_counter()" in api and "record_usd" in api, "#6 قياس الرؤية غائب"
    assert "البند #5" in api and "أداة اختبار المفاتيح قبل ضبط" in api, "#5 غير موثَّق"
    assert 'SILK_GMAPS_ENRICH_GRACE_S", "25"' in api, "#4 المهلة الآمنة غائبة"
    assert '"processing": processing' in api, "#4 علم processing غائب"
    html = _read("web/index.html")
    assert "function _expBusy(" in html and "if(S.exportBusy){" in html, (
        "#7 صمّام تعطيل التصدير/حارس النقر المزدوج غائب")
    assert _exists("tests/test_wave2_med_fixes.py"), "ملف أقفال الموجة ٢ مفقود"


def _guard_composite_source_id_attribution():
    """LESSONS ٦٠ — إسنادٌ مركّب (بلاغ قطر): معرّفُ المصدر ذرّيّ، الإسنادُ
    المتعدّد قائمةٌ لا سلسلةٌ مدموجة؛ المراجع تُسطّح فيُسنِد كلٌّ لرابطه."""
    from silk_data_layer import (atomic_source_ids, is_atomic_source_id,
                                 public_source_url)
    assert not is_atomic_source_id("IMF WEO، World Bank")   # فاصلُ دمجٍ مرفوض
    assert is_atomic_source_id("WITS/WTO Tariff")           # «/» مشروع
    assert atomic_source_ids("A", ("A", "B")) == ["A", "B"]
    # أمانةُ GAFTA لها رابطٌ ذرّيّ مستقلّ (كانت تعيش داخل مركّبٍ بلا رابط).
    assert public_source_url("GAFTA secretariat") == "https://www.lasportal.org"
    assert public_source_url("GCC secretariat") != public_source_url(
        "GAFTA secretariat")
    _needles("silk_llm_runtime.py", "source_ids=tuple(pub_sources)")()
    _needles("silk_evals.py", "listed-but-unused", "معرّفُ مصدرٍ **مركّب**")()
    _needles("tests/test_hf_attribution_truncation_plausibility.py",
             "def test_three_source_finding_yields_three_atomic_references")()


def _guard_renderer_truncation_and_empty_parens():
    """LESSONS ٦١ — بترٌ داخل رقمٍ + قوسٌ فارغ (بلاغ قطر): القصُّ لا ينتهي داخل
    رقم، وحذفُ الاستشهاد لا يترك «()»."""
    import re as _re
    from silk_reports import _trim_sentence, _client_sanitize
    from silk_render import _strip_internal_plumbing
    out = _trim_sentence("تعافٍ جزئيّ إلى 7.12 مليون دولار مؤكَّد", 22)
    assert not _re.search(r"[0-9٠-٩][.،]\s*$", out), out
    assert _client_sanitize("قيمة (/) مؤكَّدة") == "قيمة مؤكَّدة"
    assert "()" not in _strip_internal_plumbing("المتاجر (dp3) كارفور")
    _needles("silk_render.py", "_DP_GROUP_RE", "_EMPTY_CITATION_GROUP_RE")()
    _needles("tests/test_hf_attribution_truncation_plausibility.py",
             "def test_trim_sentence_never_ends_inside_a_number")()


def _guard_cross_source_plausibility():
    """LESSONS ٦٢ — مقدارٌ غيرُ مُصالَح (بلاغ قطر): حارسُ معقوليةٍ يقارن المقاديرَ
    بمرتكزات التشغيلة ويُوسَم/يُتحفَّظ عليه، ويُسجَّل في المانيفست."""
    import silk_plausibility as P
    result = {"deep_research": {"missions": {
        "trade_flow": {"findings": [{"value": 7_000_000.0, "source": "UN Comtrade",
                                    "note": "إجمالي استيراد قطر من العالم USD"}]},
        "consumer_culture": {"findings": [{"value": "497 مليون دولار",
            "source": "ويب", "note": "حجم سوق الفول السوداني الكامل"}]}}}}
    flags = P.check_magnitudes(result)
    assert flags and flags[0]["detail"]["import_ratio"] > 20
    assert P.check_magnitudes({"deep_research": {"missions": {}}}) == []  # فشلٌ آمن
    _needles("silk_render.py", "silk_plausibility.annotate")()
    _needles("silk_evals.py", "def _plausibility_reconciled")()
    _needles("tests/test_hf_attribution_truncation_plausibility.py",
             "def test_plausibility_flags_implausible_market_size")()


def _guard_ask_what_the_product_answers():
    """LESSONS ٦٥ — الحوارُ كان يسأل عن رقمٍ يُجيب عنه المنتجُ نفسُه (نسبةُ
    دهن). الحارس **سلوكيّ**: يقيس فعلاً على الوصف الرسميّ الحقيقي، ويتحقّق
    أنّ نقطةَ الاختناق موصولةٌ بمسارَي الدخول معاً (لا إصلاحَ نصفيّ)."""
    import silk_hs_attributes as A
    # الترويسةُ التي أنتجت البلاغ — بنودُها من مرجعنا الرسميّ، بلا نموذج.
    codes = [c for c in sorted(
        __import__("silk_hs_resolver").load_hs_reference())
        if c.startswith("0401") and A.band_of(c)]
    assert len(codes) >= 3, "لم تُقرأ نطاقاتُ الترويسة من الوصف الرسميّ"
    disc = A.discriminator([{"hs6": c} for c in codes])
    assert disc and disc["dimension"] == "fat" and disc["unit"] == "%"
    # الصمّامُ مُطفأٌ افتراضياً (D1/اللائحة ٧٠) — يُفعَّل هنا صراحةً لأنّ هذا
    # الحارسَ يختبر **سلوكَ الميزة**، لا افتراضَها (ذاك حارسُ اللائحة ٧٠).
    _saved = os.environ.get("SILK_HS_ATTRIBUTE_RESOLVE")
    os.environ["SILK_HS_ATTRIBUTE_RESOLVE"] = "1"
    try:
        # (أ) قياسٌ من بطاقة العبوة يحسم بلا أيّ سؤال، وموسومٌ بمصدره.
        got = A.resolve_by_attribute(
            "منتجٌ من هذه الترويسة", [{"hs6": c} for c in codes],
            allow_web=False,
            label_attributes=[{"name": disc["label_ar"], "value": 3.5,
                               "unit": "%"}])
        assert got["hs6"] and got["resolved_from"] == "image"
        assert "صورة العبوة" in got["note_ar"]
        # (ب) بلا قياسٍ لا يُحسَم رمزٌ أبداً، والحوارُ يحمل ما نقص وحدودَ البنود.
        gap = A.resolve_by_attribute(
            "منتجٌ من هذه الترويسة", [{"hs6": c} for c in codes],
            allow_web=False)
        assert gap["hs6"] is None and gap["missing_ar"] and gap["bands_ar"]
    finally:
        if _saved is None:
            os.environ.pop("SILK_HS_ATTRIBUTE_RESOLVE", None)
        else:
            os.environ["SILK_HS_ATTRIBUTE_RESOLVE"] = _saved
    # (ج) نقطةُ اختناقٍ واحدة موصولةٌ بكِلا مسارَي الإنفاق + بوّابةِ الالتباس.
    src = _read("api.py")
    assert src.count("preflight_resolve") >= 3, (
        "نقطةُ القياس غير موصولةٍ بمسارَي /analyze و/research معاً")
    _needles("api.py", "resolve_or_probe", "attribute_probe",
             "label_attributes", "hs_provenance")()
    _needles("silk_hs_confirm.py", "def resolve_or_probe",
             "def preflight_resolve")()
    _needles("silk_render.py", "hs_provenance",
             "الرمز محدَّد من صورة العبوة", "الرمز محدَّد من مصدر ويب")()
    _needles("web/index.html", "attribute_probe", "label_attributes",
             'd.error==="hs_ambiguous"')()
    # (د) عائلةُ اللائحة ١٢ (الحدود تناقض المتن): المصالحةُ اللفظية
    # (`revalidate`) لا تعمل على رمزٍ حُسِم بقياس — وإلا ظهر «المُحلِّل يعيد
    # رمزاً آخر» بجوار «الرمز محدَّد من صورة العبوة» في نفس التقرير.
    assert ('if not (isinstance(hs_provenance, dict) '
            'and hs_provenance.get("hs6")):') in src, (
        "المصالحةُ اللفظية تعمل على رمزٍ مقيس — تناقضٌ محتوم")


def _guard_band_boundary_strictness_and_second_axis():
    """LESSONS ٦٦ — عيبان يُصدِران **رمزاً خاطئاً بوسمِ مصدرٍ واثق** (اختلاقٌ
    لا فجوة): (أ) تسطيحُ صرامةِ الحدّ («less than 6» تُعامَل كـ«not exceeding
    6») فتُبتلَع قيمةُ الحدّ؛ (ب) ترويسةٌ تنقسم بمحورين (لونُ الشاي × وزنُ
    التعبئة) يحسمها قياسٌ واحد. الحارس **سلوكيّ** على المرجع الحقيقيّ."""
    import silk_hs_attributes as A
    FAKE = "000000"                       # ليس في المرجع => يُقرأ الوصفُ الحرّ
    assert A.band_of(FAKE, "x") is None or True
    inc = A.band_of(FAKE, "of a fat content, not exceeding 6%")
    strict = A.band_of(FAKE, "of a fat content, less than 6%")
    assert inc and strict, "لم تُقرأ الحدود من العبارة"
    assert inc["hi"] == strict["hi"] == 6.0
    assert inc["hi_inclusive"] is True and strict["hi_inclusive"] is False, (
        "صرامةُ الحدّ مُسطَّحة — «less than» تُعامَل معاملةَ «not exceeding»")
    assert A._contains(inc, 6.0) is True and A._contains(strict, 6.0) is False
    lo_inc = A.band_of(FAKE, "of a fat content, at least 1%")
    assert lo_inc and lo_inc["lo_inclusive"] is True, "حدٌّ أدنى شاملٌ مفقود"
    # (ب) المحورُ الثاني: ترويسةٌ تنقسم بلونٍ ووزنٍ معاً لا تُحسَم بالوزن وحده.
    import collections
    from silk_hs_resolver import load_hs_reference
    heads = collections.defaultdict(list)
    for code in load_hs_reference():
        b = A.band_of(code)
        if b:
            heads[code[:4]].append(b)
    multi = {h: bs for h, bs in heads.items() if len(bs) >= 2}
    assert len(multi) >= 40, "المرجعُ لم يُقرأ — الحارس بلا عيّنة"
    accepted = refused = 0
    for head, bands in multi.items():
        d = A.discriminator([{"hs6": b["hs6"]} for b in bands])
        if d is None:
            refused += 1
            continue
        accepted += 1
        assert len({b["axis"] for b in d["bands"]}) == 1, (
            f"{head}: قُبِلت رغم محورٍ غيرِ رقميٍّ إضافي")
        for prev, nxt in zip(d["bands"], d["bands"][1:]):
            assert prev["hi"] == nxt["lo"], f"{head}: فجوةٌ/تداخل"
            assert bool(prev["hi_inclusive"]) != bool(nxt["lo_inclusive"]), (
                f"{head}: حدٌّ مزدوجُ التغطية أو مكشوف")
    assert accepted >= 5 and refused >= 5, (
        f"توازنُ الحارس مختلّ (مقبولة {accepted}، مرفوضة {refused})")
    _needles("tests/test_hs_attribute_autoresolve.py",
             "def test_bound_strictness_comes_from_the_matched_phrase",
             "def test_property_every_multiband_heading_is_either_clean_or_refused",
             "def test_second_axis_heading_is_refused_not_resolved")()


def _guard_dimension_terms_not_frozen_in_code():
    """LESSONS ٦٧ — عودةُ عائلة الدرس ٣٠ (كلمةُ نطاقٍ حرفيةٌ مجمَّدةٌ في قالبٍ
    قابلٍ لإعادة الاستعمال): مصطلحُ البُعد كان مكتوباً في استعلام الويب،
    فيعمل على الألبان ويُخرِس كلَّ ترويسةٍ أخرى. الحارس: المعجمُ من ملفٍ،
    وصفرُ مصطلحٍ عربيٍّ في المنطق، وصفرُ مصطلحٍ في بناء الاستعلام."""
    import inspect
    import silk_hs_attributes as A
    lex = A.load_dimensions()
    assert len(lex) >= 8, "معجمُ الأبعاد لم يُقرأ من الملفّ"
    assert _exists("data/measurement_dimensions.csv")
    body = _read("silk_hs_attributes.py")
    body = re.sub(r'"""(?:.|\n)*?"""', "", body)
    body = "\n".join(ln.split("#", 1)[0] for ln in body.splitlines())
    arabic = set()
    for row in lex.values():
        if row["label_ar"]:
            arabic.add(row["label_ar"])
        arabic.update(t for t in row["syn"]
                      if any("\u0600" <= ch <= "\u06ff" for ch in t))
    leaked = sorted(t for t in arabic if t and t in body)
    assert not leaked, f"مصطلحُ بُعدٍ عربيٌّ مجمَّدٌ في المنطق: {leaked}"
    qsrc = inspect.getsource(A.probe_web)
    frozen = sorted(t for dim, row in lex.items()
                    for t in ((row["label_ar"],) + tuple(row["syn"]) + (dim,))
                    if len(t) >= 3 and t in qsrc)
    assert not frozen, f"مصطلحٌ مجمَّدٌ في بناء الاستعلام: {frozen}"
    # وبُعدٌ خارج الملفّ يتدهور لمفتاحه — لا مصطلحٌ مختلَق.
    assert A.dimension_terms("zzz_x") == ("zzz_x", ("zzz_x",))


def _guard_multi_axis_heading_confident_wrong_code():
    """LESSONS ٦٨ — «تطابقٌ رقميٌّ على ترويسةٍ متعدّدةِ المحاور = رمزٌ خاطئ
    بثقة». الترويسةُ قد تنقسم بمحورٍ رقميٍّ **وآخرَ غيرِ رقميّ معاً**
    (0902 = لونُ الشاي أخضر/أسود × وزنُ التعبئة ≤٣كجم/>٣كجم). قياسُ الوزن
    وحده لا يُحدِّد بنداً — يختار أحدَ اثنين يختلفان في اللون أيضاً، فيخرج
    رمزٌ **خاطئ** موسومٌ «الرمز محدَّد من صورة العبوة». اختلاقٌ لا فجوة.

    حارسٌ **سلوكيّ** على المرجع الحقيقيّ لا فحصُ وجود: يبني مجموعاتِ مرشّحين
    فعلية ويؤكّد الرفض. (الصفّ ٦٦ يفحص هذه العائلة ضمن فحصٍ مركّب مع صرامةِ
    الحدّ؛ هذا الصفُّ يفردها بحارسها الخاصّ بأمر المُشرِف — العائلةُ اكتُشفت
    خارج قائمة الفجوات المُسمّاة، فتستحقّ قفلاً لا يذوب في غيره.)"""
    import collections
    import itertools
    import silk_hs_attributes as A
    from silk_hs_resolver import load_hs_reference

    # (١) حادثةُ العائلة بعينها: لونٌ مختلف + وزنٌ مختلف => رفضٌ قاطع.
    ref = load_hs_reference()
    tea = [c for c in ("090210", "090220", "090230", "090240") if c in ref]
    assert len(tea) == 4, f"مرجعُ 0902 ناقص: {tea}"
    green_light, green_heavy, black_light, black_heavy = tea
    assert A.discriminator([{"hs6": green_light}, {"hs6": black_heavy}]) is None, (
        "قُبِل مُميِّزٌ لبندين يختلفان في اللون **والوزن** — وزنٌ يحسم رمزاً "
        "يختلف في محورٍ آخر: رمزٌ خاطئ بوسمِ مصدرٍ واثق")
    assert A.discriminator([{"hs6": green_heavy}, {"hs6": black_light}]) is None
    # (٢) ضابطٌ موجب — نفسُ المحور يمرّ، فالحارسُ ليس رفضاً شاملاً.
    same_axis = A.discriminator([{"hs6": green_light}, {"hs6": green_heavy}])
    assert same_axis is not None, (
        "رُفِض بندان يختلفان بالوزن وحده — الحارسُ يرفض كلَّ شيء (لا قيمة له)")
    assert len({b["axis"] for b in same_axis["bands"]}) == 1

    # (٣) كنسٌ شاملٌ على المرجع كلِّه: **كلُّ** زوجٍ عابرِ المحور يُرفَض.
    by_head = collections.defaultdict(list)
    for code in ref:
        band = A.band_of(code)
        if band is not None:
            by_head[code[:4]].append(band)
    multi = {h: bs for h, bs in by_head.items() if len(bs) >= 2}
    assert len(multi) >= 40, f"عيّنةٌ أضعفُ من المتوقَّع: {len(multi)}"
    cross_pairs = 0
    for head, bands in multi.items():
        by_axis = collections.defaultdict(list)
        for b in bands:
            by_axis[b["axis"]].append(b["hs6"])
        if len(by_axis) < 2:
            continue
        for g1, g2 in itertools.combinations(list(by_axis.values()), 2):
            cross_pairs += 1
            assert A.discriminator([{"hs6": g1[0]}, {"hs6": g2[0]}]) is None, (
                f"{head}: قُبِل زوجٌ عابرُ المحور {g1[0]}/{g2[0]}")
    assert cross_pairs >= 100, (
        f"أزواجٌ عابرةُ المحور أقلُّ من المتوقَّع ({cross_pairs}) — "
        "الكنسُ لم يعمل فعلياً")
    # (٤) وأنّ المقبولَ ما يزال موجوداً (لا انهيارَ تغطيةٍ صامت).
    accepted = sum(
        1 for bs in multi.values()
        if A.discriminator([{"hs6": b["hs6"]} for b in bs]) is not None)
    assert accepted >= 5, f"لم يبقَ مقبولٌ يُذكَر ({accepted})"
    _needles("silk_hs_attributes.py", "def _residual_axis", '"axis"')()


def _guard_client_operator_document_divergence():
    """LESSONS ٦٩ — «اختبارُ عرضٍ أخضرُ بجوار مُسلَّمِ عميلٍ خاطئ؛ أكِّدْ على
    الأثر المُصيَّر». سطرُ إفصاحِ مصدر الرمز كان يظهر في مستند **المشغّل**
    ويغيب عن مستند **العميل** — المُسلَّم الحقيقيّ — لأنّ القالبين يبنيان من
    مصدرين مختلفين (`deep_research` مقابل `limits`). اختبارُ الوحدة على
    `view["limits"]` بقي أخضرَ طوال الوقت.

    الحارس **يفتح ملفّ .docx المُصدَّر فعلاً** — لا يقرأ عرضاً ولا يفحص وجودَ
    رمز."""
    pytest.importorskip("docx")
    from docx import Document
    import silk_render
    import silk_reports
    from canonical_netherlands import netherlands_research_blob

    def _doc_text(path: str) -> str:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    url = "https://example-provenance.test/label"
    blob = netherlands_research_blob()
    blob["hs_provenance"] = {
        "hs6": "040120", "resolved_from": "web", "attribute": "fat",
        "label_ar": "نسبة الدهن", "value": 3.5, "unit": "%",
        "source_url": url, "confidence": 0.5}

    saved = os.environ.get("SILK_HERMETIC")
    os.environ["SILK_HERMETIC"] = "1"
    try:
        view = silk_render.build_view(blob)
        out = tempfile.mkdtemp()
        # **كِلا** المُسلَّمين — التباعدُ هو العطل، فلا يكفي فحصُ أحدهما.
        for renderer, label in (("render_client_docx", "العميل"),
                                ("render_docx", "المشغّل")):
            path = getattr(silk_reports, renderer)(
                view, os.path.join(out, f"{renderer}.docx"))
            assert os.path.exists(path), f"{label}: لم يُنتَج ملفّ"
            text = _doc_text(path)
            assert "الرمز محدَّد من مصدر ويب" in text, (
                f"مستند {label}: سطرُ إفصاح مصدر الرمز غائبٌ عن الملفّ "
                "المُصدَّر فعلاً (اختبارُ العرض لا يكشف هذا)")
            assert url in text, f"مستند {label}: الرابطُ المُستشهَد به غائب"
        # والنفيُ المقابل: بلا حسمٍ آليّ لا جملةَ إفصاحٍ مُقحَمة.
        plain = silk_render.build_view(netherlands_research_blob())
        clean_path = silk_reports.render_client_docx(
            plain, os.path.join(out, "plain.docx"))
        assert "الرمز محدَّد" not in _doc_text(clean_path), (
            "جملةُ إفصاحٍ ظهرت على مستندٍ لم يُحسَم رمزُه آلياً")
    finally:
        if saved is None:
            os.environ.pop("SILK_HERMETIC", None)
        else:
            os.environ["SILK_HERMETIC"] = saved
    _needles("silk_reports.py", "def _hs_provenance_sentence")()


def _guard_attribute_resolver_flag_off_by_default():
    """D1 — الصمّامُ **مُطفأٌ افتراضياً**. تفعيلُه عند الدمج يجعل ميزةً لم
    تُجرَّب قطّ بمفتاحٍ حيّ تُصدِر رموزاً جمركية لكلّ مستخدم؛ وتفعيلُه قرارُ
    مالكٍ منفصلٌ بعد الدمج مشروطٌ بإغلاق G8. الحارسُ سلوكيّ: يقرأ الدالّة
    فعلاً في غياب المتغيّر وفي حضوره."""
    import silk_hs_attributes as A
    saved = os.environ.pop("SILK_HS_ATTRIBUTE_RESOLVE", None)
    try:
        assert A.enabled() is False, (
            "الصمّامُ مفعّلٌ افتراضياً — خرقُ D1: ميزةٌ بلا دليلٍ حيّ تُصدِر "
            "رموزاً جمركية لكلّ مستخدم عند الدمج")
        for on in ("1", "true", "yes", "on"):
            os.environ["SILK_HS_ATTRIBUTE_RESOLVE"] = on
            assert A.enabled() is True, f"لم يُفعَّل بـ{on!r}"
        for off in ("0", "false", "no", "off", "", "maybe"):
            os.environ["SILK_HS_ATTRIBUTE_RESOLVE"] = off
            assert A.enabled() is False, f"فُعِّل بقيمةٍ ليست تفعيلاً: {off!r}"
        # ومُطفأً: لا حسمَ إطلاقاً مهما كانت القراءةُ صالحة.
        os.environ.pop("SILK_HS_ATTRIBUTE_RESOLVE", None)
        out = A.resolve_by_attribute(
            "منتجٌ ما", [{"hs6": c} for c in
                        ("040110", "040120", "040140", "040150")],
            label_attributes=[{"name": A.dimension_terms("fat")[0],
                               "value": 3.5, "unit": "%"}],
            allow_web=False)
        assert out["hs6"] is None and out["resolved_from"] is None
    finally:
        if saved is None:
            os.environ.pop("SILK_HS_ATTRIBUTE_RESOLVE", None)
        else:
            os.environ["SILK_HS_ATTRIBUTE_RESOLVE"] = saved


def _guard_cross_basis_edge_refusal():
    """D2 — قراءةٌ بأساسِ نسبةٍ مخالف قربَ حافّة لا تحسم. `g/100ml` كتلة/حجم
    بينما نصّ HS «by weight» كتلة/كتلة؛ الفارقُ ~٣٪ نسبيّاً (~٠٫١٨ عند حدّ
    ٦٫٠) — يكفي لعبور الحدّ، فتُلتَفّ صرامةُ G1 من هذا الباب الواحد."""
    import silk_hs_attributes as A
    saved = os.environ.get("SILK_HS_ATTRIBUTE_RESOLVE")
    os.environ["SILK_HS_ATTRIBUTE_RESOLVE"] = "1"
    try:
        d = A.discriminator([{"hs6": c} for c in
                             ("040110", "040120", "040140", "040150")])
        assert d is not None
        assert A.band_basis(d["bands"][0]) == "mm", "أساسُ النطاق لم يُقرأ"
        for v in (5.9, 6.0, 6.1):          # داخل ٠٫٥ من حدّ ٦٫٠
            assert A.select_by_value(d, v, "g/100ml") is None, (
                f"{v} g/100ml حُسِمت قربَ حافّة رغم مخالفة الأساس")
        assert A.select_by_value(d, 3.5, "g/100ml") == "040120"   # بعيدةٌ
        for same in ("%", "g/100g", "% w/w"):                     # نفسُ الأساس
            assert A.select_by_value(d, 5.9, same) == "040120", same
        # وكلُّ مدخلٍ مخالفِ الأساس في الجدول محروسٌ فعلاً (لا بابَ جديد).
        cross = {u for u in A._UNIT_FAMILY
                 if A.cross_basis_conflict(u, d["bands"])}
        assert cross >= {"g/100ml", "gm/100ml", "mg/100ml", "ml/100ml"}, cross
    finally:
        if saved is None:
            os.environ.pop("SILK_HS_ATTRIBUTE_RESOLVE", None)
        else:
            os.environ["SILK_HS_ATTRIBUTE_RESOLVE"] = saved
    _needles("silk_hs_attributes.py", "def cross_basis_conflict",
             "def near_any_edge", "_UNIT_BASIS", "_EDGE_MARGIN")()


def _dialog_band_numbers(band: dict) -> list[str]:
    """أرقامُ النطاق كما يجب أن تظهر — مشتقّةٌ من **المُحلِّل** لا من المُصيِّر
    المُختبَر، فالتأكيدُ ليس دائرياً."""
    unit = band.get("unit") or ""
    out = []
    for v in (band.get("lo"), band.get("hi")):
        if v is None:
            continue
        out.append(f"{int(v) if float(v).is_integer() else v}{unit}")
    return out


def _guard_dialog_band_text_from_the_official_reference_only():
    """البند ٧٢ — نصُّ الحدّ المعروض مشتقٌّ من المرجع الرسميّ حصراً.

    الحادثة: نثرُ نموذجٍ وقتَ الطلب كان يغلب الوصفَ الرسميّ بالتداخل اللفظيّ،
    فعُرِض على التاجر حدٌّ **يناقض** بندَه (`040110` بوصف «لا تتجاوز 6%»
    وبندُه ≤١٪)، ورمزٌ لا وجودَ له في المرجع (`040190`). الحارسُ سلوكيّ: يمرّ
    نثراً مناقضاً عبر نقطة الاختناق ويقرأ الناتج، ثم يكنس **كلّ** بندٍ ذي
    نطاقٍ في `data/hs_reference.csv` لا عيّنة."""
    import silk_hs_attributes as attrs
    import silk_hs_dialog as dialog
    from silk_hs_resolver import load_hs_reference, official_description

    contradicting = {
        "040110": "نسبة الدهن لا تتجاوز 6%",
        "040120": "نسبة الدهن تتجاوز 6%",
        "040190": "حليب وقشطة أخرى",
    }
    rows = dialog.build_candidates(
        "حليب نادك كامل الدسم", list(contradicting), contradicting)
    shown = {r["hs6"] for r in rows}
    assert "040190" not in shown, "رمزٌ مجهولٌ للمدوّنتين عُرِض على التاجر"
    # والنفيُ المقابل: نقصُ نسختنا من المرجع ليس سقفاً — بندٌ حقيقيٌّ تعرفه
    # البذرةُ وحدها يُعرَض، بلا نصٍّ مُختلَق (اللائحة ٣٩ من طرفها المقابل).
    from silk_hs_resolver import load_hs_codes
    seed_only = sorted({r["hs_code"] for r in load_hs_codes()}
                       - set(load_hs_reference()))
    assert seed_only, "لا فجوةَ بين المدوّنتين — الحارسُ فقد موضوعَه"
    kept = {r["hs6"]: r for r in dialog.build_candidates("—", seed_only)}
    assert set(seed_only) <= set(kept), (
        f"بنودٌ حقيقيةٌ سقطت لنقصِ نسختنا: {set(seed_only) - set(kept)}")
    assert all(kept[c]["description_ar"] == "" for c in seed_only), (
        "نصٌّ غيرُ رسميٍّ سدّ فراغَ الوصف")
    assert {"040110", "040120", "040140", "040150"} <= shown, shown
    by_code = {r["hs6"]: r for r in rows}
    assert "6%" not in by_code["040110"]["band_ar"], (
        "حدُّ 040110 المعروض يناقض بندَه الرسميّ (≤١٪) — عودةُ الحادثة")
    for code, prose in contradicting.items():
        row = by_code.get(code)
        if row is None:
            continue
        assert row["description_ar"] == official_description(code), (
            f"{code}: الوصفُ المعروض ليس الوصفَ الرسميّ حرفياً")
        assert prose not in row["band_ar"], f"{code}: نثرُ نموذجٍ صار حدَّ بند"
        assert prose not in row["description_ar"], f"{code}: نثرٌ صار وصفاً"

    # كنسٌ كامل: كلُّ بندٍ ذي نطاقٍ في المرجع يُصيَّر بأرقام نطاقه هو.
    checked = 0
    for code in load_hs_reference():
        band = attrs.band_of(code)
        if band is None:
            continue
        text = dialog.band_text_ar(code)
        assert text, f"{code}: بندٌ ذو نطاقٍ بلا نصّ حدٍّ معروض"
        for needle in _dialog_band_numbers(band):
            assert needle in text, (
                f"{code}: النصُّ «{text}» لا يحمل حدَّ المرجع {needle}")
        checked += 1
    assert checked >= 380, f"الكنسُ لم يشمل المرجعَ فعلياً: {checked}"
    _needles("silk_hs_dialog.py", "def build_candidates", "def band_text_ar",
             "def official_text", "def in_official_reference")()
    # ولا مُصيِّرَ ثانٍ: كلُّ منتجٍ لقائمة الحوار يمرّ بنقطة الاختناق.
    _needles("silk_hs_classifier.py", "silk_hs_dialog")()
    _needles("silk_hs_confirm.py", "silk_hs_dialog")()


def _guard_dialog_axis_siblings_never_partial():
    """البند ٧٣ — لا تُعرَض مجموعةٌ جزئيةٌ من محورٍ رقميّ.

    الحادثة: سقط `040140`/`040150` من القائمة فلم يجد منتجٌ كامل الدسم خياراً
    صحيحاً أصلاً — فيختار التاجر أقربَ المعروض ويخرج برمزٍ **خاطئ بلا إشارة**.
    السببُ مركّب: فجوةُ بذرةٍ (٨ من ٣٩٣ بنداً فقط قابلةٌ للبلوغ باسمٍ عربيّ)
    مضروبةٌ في اقتطاعٍ صلبٍ إلى ثلاثة في الخادم وفي الواجهة معاً.

    الحارسُ خاصّيّ على **كامل** المرجع: أيُّ عضوٍ من أيّ مجموعةِ محورٍ يُدخَل
    وحدَه يُخرِج المجموعةَ كاملة."""
    import collections
    import silk_hs_attributes as attrs
    import silk_hs_dialog as dialog
    from silk_hs_resolver import load_hs_reference

    groups: dict = collections.defaultdict(list)
    for code in load_hs_reference():
        band = attrs.band_of(code)
        if band is not None:
            groups[(code[:4], band["axis"])].append(code)
    families = {k: sorted(v) for k, v in groups.items() if len(v) >= 2}
    assert len(families) >= 60, f"مجموعاتُ المحاور تبدو مبتورة: {len(families)}"

    covered = 0
    for members in families.values():
        for member in members:
            shown = [r["hs6"] for r in dialog.build_candidates("—", [member])]
            missing = [m for m in members if m not in shown]
            assert not missing, (
                f"{member}: أشقّاءُ المحور غائبون عن الحوار {missing}")
            covered += 1
    assert covered >= 150, f"الكنسُ لم يشمل المجموعاتِ فعلياً: {covered}"
    # ولا اقتطاعَ في الواجهة يُعيد العطلَ بعد إصلاح الخادم.
    _absent("web/index.html", "cands.slice(0,3)", "cands.slice(0, 3)")()
    _needles("web/index.html", "c.band_ar")()
    # وحدُّ النطاق: الإكمالُ يصل الحوارَ ولا يصل المُحلِّلَ الرقميّ — وإلا
    # اتّسعت تغطيةُ الحسم من بابٍ خلفيّ (نهيُ المُشرِف الصريح).
    import silk_hs_attributes as _attrs
    import silk_hs_confirm as confirm
    rows = dialog.build_candidates("—", ["040110"])
    assert {r["hs6"] for r in rows if r["axis_completion"]} == {
        "040120", "040140", "040150"}, rows
    seen: dict = {}
    real = _attrs.resolve_by_attribute
    _attrs.resolve_by_attribute = (
        lambda p, c, **kw: (seen.setdefault("codes",
                                            [x.get("hs6") for x in c]),
                            real(p, c, **kw))[1])
    try:
        confirm.resolve_or_probe("—", rows, allow_web=False)
    finally:
        _attrs.resolve_by_attribute = real
    assert seen["codes"] == ["040110"], (
        f"المُحلِّلُ غُذّي بأشقّاءَ لم يطلبهم المستدعي: {seen['codes']}")


def _guard_dialog_prose_carries_no_product_brand_or_country():
    """البند ٧٤ — نثرُ الحوار وصفٌ رسميٌّ للبند، لا صدىً لِما كتبه التاجر.

    الحادثة: النصُّ المعروض حمل **اسمَ العلامة التجارية** داخل وصفٍ يُقدَّم
    بوصفه رسمياً («… حليب نادك كامل الدسم») — فيبدو الوصفُ مُصادِقاً على
    منتجِ التاجر بينما هو وصفُ بندٍ جمركيّ. عائلةُ اللائحة ٣٠ (لا منتجَ
    مثبَّتٌ في القوالب) من الطرف المقابل: لا منتجَ **مُقحَمٌ** في النثر."""
    import silk_hs_dialog as dialog

    product = "حليب نادك كامل الدسم"
    for prose, banned in (
            (f"وصفٌ عامّ لـ{product} من هولندا", ("نادك", "هولندا")),
            ("عبوة نادك المستوردة من هولندا وألمانيا",
             ("نادك", "هولندا", "ألمانيا")),
            ("Nadec milk from Netherlands", ("Netherlands",)),
    ):
        out = dialog.sanitize_prose(prose, product)
        for token in banned:
            assert token not in out, (
                f"«{token}» نجا في نثر الحوار: {out!r} — عودةُ الحادثة")

    rows = dialog.build_candidates(
        product, ["040110", "040120"],
        {"040110": f"يناسب {product}", "040120": "قشطة من هولندا"})
    assert rows, "نقطةُ الاختناق لم تُخرِج شيئاً"
    for row in rows:
        for token in ("نادك", "هولندا"):
            assert token not in row["reason_ar"], (
                f"{row['hs6']}: «{token}» وصل نثرَ الحوار عبر نقطة الاختناق")
    _needles("silk_hs_dialog.py", "def sanitize_prose", "_AR_CLITICS",
             "def _country_terms")()


def _guard_gate_passes_synthetic_but_silent_on_real():
    """LESSONS ٧٥ — بوّابةُ A3 (TAM أصغر من تدفّق دولةٍ واحدة) شُحنت في PR A
    باختباراتٍ تركيبية خضراء ثم **لم تُطلِق على تحليل ٧ الحيّ** — الحالة التي
    بُنيت لها بالضبط: الكاتبُ كتب TAM بصيغة رمز `$` (`2,090,000$`) بينما اشترط
    التطابقُ لفظَ «دولار». الحارس السلوكي على مدوّنة تحليل ٧ الحقيقية الشكل
    (`tools/canonical_nadec_yemen_dairy.py`، مُمرَّرة عبر `build_view`):
      (١) البوّابات الثلاث تُطلِق فعلاً (A3 صيغة `$` + سردُ المرآة + تصعيدُ
          التقادُم) والحكمُ الكلّي FAIL؛
      (٢) المستخلِص يلتقط صيغة الرمز `$` (سبب الصمت القديم)؛
      (٣) العيّنة النظيفة (الكويت) لا تُطلِق أياً منها (لا إيجابٌ كاذب)."""
    import silk_render as R
    import silk_quality_gate as QG
    from tools.canonical_nadec_yemen_dairy import nadec_yemen_research_blob
    from tools.canonical_kuwait_peanut_butter import kuwait_research_blob
    view = R.build_view(nadec_yemen_research_blob())
    dr = view["deep_research"]
    # (١) البوّابات الأربع تُطلِق + FAIL كلّي (بلاغ المالك: قلبُ إشارة CAGR
    # بسنة الأساس المرصودة أُضيف بعد مِجَسّ /trend الحيّ — أساس 2018 نموّ مقابل
    # أساس 2019 انكماش على نفس سلسلة 040110).
    assert QG._check_tam_below_single_country_flow(dr)
    assert QG._check_mirror_divergence_contraction_narrative(dr)
    assert QG._check_stale_year_driving_conclusion(dr)
    assert QG._check_cagr_sign_flips_under_base_year(dr)
    out = QG.run_quality_gate(view)
    assert out["verdict"] == QG.FAIL
    fired = {f["check"] for f in out["findings"]}
    assert {"tam_below_single_country_flow",
            "mirror_divergence_contraction_narrative",
            "stale_year_driving_conclusion",
            "cagr_sign_flips_under_base_year"} <= fired
    # (٢) صيغةُ الرمز `$` محفوظةٌ بعد التطهير + المستخلِص يلتقطها.
    txt = (dr.get("report") or {}).get("text") or ""
    assert "2,090,000$" in txt
    assert 2_090_000 in [round(v) for _, _, v in QG._iter_usd_amounts(txt)]
    # (٣) العيّنة النظيفة لا تُطلِق أياً من البوّابات الثلاث.
    kdr = R.build_view(kuwait_research_blob())["deep_research"]
    assert not QG._check_tam_below_single_country_flow(kdr)
    assert not QG._check_mirror_divergence_contraction_narrative(kdr)
    assert not QG._check_stale_year_driving_conclusion(kdr)
    # (٤) الثلاث بنودُ فشلٍ حاجبة + قفلُ الانحدار موجود.
    for c in ("tam_below_single_country_flow",
              "mirror_divergence_contraction_narrative",
              "stale_year_driving_conclusion",
              "cagr_sign_flips_under_base_year"):
        assert c in QG.FAIL_TRIGGER_CHECKS
    assert _exists("tests/test_gate_regression_locks_analysis7.py")
    assert _exists("tools/canonical_nadec_yemen_dairy.py")


def _guard_seat_lock_is_load_bearing():
    """LESSONS ٧٦ — اختبارُ تزامنٍ اجتاز شيفرةً **غير ذرّية** فبدا حارساً وهو خامل.

    مسارُ إعادة تنشيط المستخدم (PR-2) كان فحصاً-ثمّ-تحديثاً بلا قفلٍ فوري، فطلبان
    متوازيان يمرّان معاً من بوّابة المقاعد ويكتبان ⇒ تجاوزُ سقفٍ مدفوع. التقطته
    المراجعة الذاتية (§58)، لكن **اختبارَ الحاجز الأوّل اجتاز النسخة المعطوبة**:
    القسمُ الحرج أقصر من ميلي ثانية فتسلسلَ الخيطان بحكم GIL.

    الحارس يحمي **الاثنين معاً** — القفل والمُميِّز:
      (١) `BEGIN IMMEDIATE` قائم في مسارَي كتابة المقعد (الإنشاء والتنشيط)؛
      (٢) نافذةُ الفحص المُوسَّعة (`_widen_seat_check_window`) مربوطةٌ فعلاً
          باختبارَي المقاعد — فلو حُذفت لعاد الاختبار يجتاز كوداً غير ذرّي بصمت.
    إثباتٌ محفوظ: بحذف القفل يخرج الاختبار `[True, True]` (٤ نشطين على سقف ٣).
    """
    # **الشكل التنفيذي وحده يُحتسَب.** الفحص الأوّل كان على النصّ المجرّد
    # «BEGIN IMMEDIATE»، وهو يظهر في **شروح** هذه الوحدات أيضاً — فكان ذكرٌ في
    # docstring يُشبِع الحارس، ويمكن حذف قفلٍ فعليّ دون أن يُحمِّر. نفس ثقب
    # «ذكر العمود ≠ قيد عليه» الذي أُغلق في حارس العزل، تكرّر هنا فأُغلق.
    # Only the executable statement counts — a docstring mention is not a lock.
    _LOCK = 'conn.execute("BEGIN IMMEDIATE")'
    src = _read("silk_platform/users.py")
    assert src.count(_LOCK) >= 2, (
        "silk_platform/users.py: مسارا كتابة المقعد (create_sub_user/set_active) "
        "يجب أن يبدأا معاملة كتابة فورية فعلياً — بلا ذلك يتجاوز تنشيطان "
        f"متزامنان السقف (وُجد {src.count(_LOCK)} من عبارات القفل التنفيذية)")
    # التنشيط تحديداً محروسٌ بالحالة كي تخسر الكتابة الثانية بلا أثر.
    assert "AND is_active = ?" in src, (
        "silk_platform/users.py: تحديث التنشيط بلا شرط `is_active` — الكتابة "
        "الثانية تنجح صمتاً")
    # القاعدة أوسع من المقاعد: **كل** مسار يستهلك حدّاً مدفوعاً يفتح معاملة كتابة
    # فورية قبل فحصه. أُضيف الخصم المقيس (PR-3) لأن فحص الخمول ثمّ الخصم بلا قفل
    # يُنتج خصماً مزدوجاً على مفتاح واحد — والدفتر غير قابل للتعديل.
    bill = _read("silk_platform/billing.py")
    assert _LOCK in bill, (
        "silk_platform/billing.py: فحص الخمول والخصم يجب أن يتشاركا معاملة كتابة "
        "فورية فعلية — بلا ذلك تخصم نقرتان متزامنتان مرّتين على مفتاح واحد")

    tests = _read("tests/test_platform_concurrency.py")
    for widener in ("def _widen_seat_check_window", "def _widen_charge_check_window"):
        assert widener in tests, (
            f"tests/test_platform_concurrency.py: {widener} محذوف — بدونه يجتاز "
            "كودٌ غير ذرّي اختبارَ التزامن (أخضر فارغ)")
    for name, widener in (
            ("test_concurrent_sub_user_creates_never_exceed_seat_cap",
             "_widen_seat_check_window"),
            ("test_concurrent_reactivations_never_exceed_seat_cap",
             "_widen_seat_check_window"),
            ("test_concurrent_metered_charges_on_one_key_charge_exactly_once",
             "_widen_charge_check_window")):
        assert name in tests, f"اختبار قفل حدٍّ مدفوع مفقود: {name}"
        body = tests.split(f"def {name}(")[1].split("\ndef ")[0]
        assert widener in body, (
            f"{name}: لا يُوسِّع نافذة الفحص ⇒ قد يجتاز شيفرةً غير ذرّية")


def _guard_readiness_names_the_offending_variable():
    """LESSONS ٧٧ — أداةُ تشخيصٍ أبلغت بالفشل وحجبت سببَه.

    `readiness()` شُحن في #197 لينهي «الدخول مستحيلٌ بلا تفسير»، ثم ضبط المالك
    البوّابةَ وظلّ الدخول يُرفَض: كلمتُه خالفت السياسة، فرفع التلبيد، فابتلعه
    الإقلاع صواباً، فبقيت القاعدة بصفر مستخدمين — والجهوزيّة قالت `seeded:false`
    مع `seed_gate_set:true` **بلا سبب**، فصمتت عند السؤال الوحيد المهم.

    الحارس يحمي ثلاثة أشياء معاً:
      (١) الحقلُ `seed_error` قائمٌ في الجهوزيّة (لا عودةَ إلى «فشلٌ بلا سبب»)؛
      (٢) الرفضُ يقع **قبل** المحاولة (`seed_problem()` في `maybe_seed`) كي
          تسمّي الرسالةُ المتغيّرَ المخالف بعينه لا الكلمةَ السليمة؛
      (٣) لا قيمةَ كلمةِ مرورٍ في أيّ مخرَج — تُنشَر الأسماءُ والقواعد فقط.
    """
    src = _read("silk_platform/bootstrap.py")
    for needle in ("def seed_problem", '"seed_error"', "validate_policy"):
        assert needle in src, (
            f"silk_platform/bootstrap.py: {needle} محذوف — الجهوزيّة تعود "
            "تُبلِّغ بالفشل بلا سببه، وهو العجزُ الذي وُجدت لإلغائه")
    # الرفضُ المسبق: `seed_problem()` مستدعًى داخل `maybe_seed` نفسه، لا في
    # الجهوزيّة وحدها — فبلا ذلك يعود السجلّ إلى «فشل» بلا اسم متغيّر.
    body = src.split("def maybe_seed(")[1]
    assert "seed_problem()" in body, (
        "silk_platform/bootstrap.py: `maybe_seed` لا يرفض مسبقاً ⇒ رسالةُ "
        "السجلّ تعود بلا اسم المتغيّر المخالف")
    # لا تُطبَع قيمةُ أيّ متغيّر بذر — الأسماءُ فقط (`/health` عامّة).
    assert "os.environ.get(env" in src or "_IDENTITY_ENV" in src, (
        "silk_platform/bootstrap.py: أسماءُ متغيّرات البذر لم تبقَ بياناتٍ "
        "واحدةَ المصدر")
    tests = _read("tests/test_platform_bootstrap.py")
    for name in ("test_a_policy_violating_seed_password_is_named_in_readiness",
                 "test_readiness_never_leaks_the_seed_password_value",
                 "test_a_bad_optional_password_names_that_variable_not_the_admin",
                 "test_the_policy_refusal_logs_the_variable_name_and_not_its_value",
                 "test_the_platform_prefix_leads_to_the_page_not_a_bare_404"):
        assert f"def {name}" in tests, f"قفلُ تشخيصٍ مفقود: {name}"
    # والبادئة تقود إلى الصفحة — المالك فتح `/platform` فرأى 404 بصيغة JSON.
    # **يُفحَص تسجيلُ المسار نفسه** لا اسمُ الدالّة: الفحص الأوّل كان على «def
    # platform_root»، فتعطيلُ المعالج بإعادة تسميته `platform_root_DISABLED`
    # يُبقي النصَّ الفرعيّ حاضراً ⇒ يجتاز الحارسُ مساراً محذوفاً. نفسُ ثقب
    # النصّ الفرعيّ المعروف، أُغلق هنا بفحص المُزخرِف والهدف معاً.
    api_src = _read("silk_platform/api.py")
    assert "@app.get(_PREFIX)\n" in api_src, (
        "silk_platform/api.py: لا مسارَ مسجَّلاً على البادئة المجرّدة ⇒ "
        "`/platform` يعود 404 بصيغة JSON فيتكرّر لبسُ المالك")
    assert 'RedirectResponse("/platform.html"' in api_src, (
        "silk_platform/api.py: البادئة لا تُحوِّل إلى الصفحة فعلاً")


def _guard_image_evidence_decides_prepared_form():
    """LESSONS ٧٨ — حادثة حليب الفراولة: أدلةُ الملصق بلغت نداء التصنيف ثم
    ضاعت في أربع نقاطٍ متراكبة. الحارس السلوكي: (١) فائزُ كلود عبر-البند
    يتصدّر القائمة المعروضة (العنصر [0]) ولا يُدفَن بفرز محور بند الاسم؛
    (٢) الاقتراحاتُ الخام والمرفوضُ بنيوياً يُسجَّلان عند INFO؛ (٣) الـprompt
    يُعلن أولويةَ أدلة الملصق على الاسم ولا يُجبِر نموذجاً بعينه بلا قرار
    مشغّل؛ (٤) المخزنُ الموجَّه صراحةً إلى SQLite لا يختطف DATABASE_URL."""
    import json as _json
    import logging as _logging
    import tempfile as _tempfile
    import unittest.mock as _mock

    import silk_hs_classifier as hsc
    import silk_store

    hints = ["Sugars 11g", "Strawberry (حليب بالفراولة)", "Tetra pack carton"]
    seen, records = {}, []

    def _capture(system, user, **kw):
        seen["user"] = user
        seen.update(kw)
        return _json.dumps({"candidates": [
            {"hs6": "220299",
             "description_ar": "مشروبات غير كحولية — مشروب حليب milk منكّه",
             "reason_ar": "شكلٌ محضَّر للشرب بحسب أدلة الملصق",
             "confidence": 0.9},
            {"hs6": "999999", "description_ar": "فصل لا وجود له",
             "reason_ar": "لاختبار الرفض البنيوي", "confidence": 0.5}]})

    handler = _logging.Handler()
    handler.emit = lambda rec: records.append(rec.getMessage())
    lg = _logging.getLogger("silk.hs_classifier")
    old_level = lg.level
    lg.addHandler(handler)
    lg.setLevel(_logging.INFO)
    try:
        with _tempfile.TemporaryDirectory() as d, \
             _mock.patch.object(silk_store, "_db_path",
                                return_value=os.path.join(d, "store.db")), \
             _mock.patch.dict(os.environ,
                              {"SILK_HS_CLASSIFIER": "1",
                               "SILK_HS_CLASSIFY_MODEL": ""}), \
             _mock.patch("silk_ai_judge.available", return_value=True), \
             _mock.patch("silk_ai_judge._call", side_effect=_capture), \
             _mock.patch("silk_usage.try_reserve_paid_calls",
                         return_value=True), \
             _mock.patch("silk_usage.try_reserve_usd", return_value=True):
            r = hsc.classify_general("milk حليب", ingredients=hints,
                                     allow_claude=True)
    finally:
        lg.removeHandler(handler)
        lg.setLevel(old_level)

    codes = [c["hs6"] for c in r["candidates"]]
    assert r["candidates"][0]["hs6"] == "220299", (
        f"فائز عبر-البند لم يتصدّر العرض: {codes}")
    blob = "\n".join(records)
    assert "hs llm proposed" in blob and "220299" in blob, blob
    assert "rejected by structural gate" in blob and "999999" in blob, blob
    assert "بيّنةٌ أقوى من الاسم" in seen["user"], "إعلان الأولوية غائب"
    assert seen.get("model") is None, (
        f"نموذجٌ مُجبَرٌ بلا قرار مشغّل: {seen.get('model')!r}")

    # (٤) المخزن: التوجيه الصريح يسبق DATABASE_URL؛ وبلا توجيهٍ يبقى العقد.
    saved = {k: os.environ.get(k)
             for k in ("DATABASE_URL", "SILK_DATA_DIR", "SILK_STORE_DB")}
    try:
        for k in saved:
            os.environ.pop(k, None)
        os.environ["DATABASE_URL"] = "postgresql://u:x@h.invalid:5432/db"
        assert silk_store._is_postgres() is True
        os.environ["SILK_DATA_DIR"] = "/tmp/anywhere"
        assert silk_store._is_postgres() is False, (
            "SILK_DATA_DIR الصريح لم يفز على DATABASE_URL")
    finally:
        for k, v in saved.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v


_LESSONS = {
    1: _needles("docs/LIVE_PROOF_RUNBOOK.md", "لا يُشغَّل هيرمتياً"),
    2: _needles("silk_render.py", "_deep_research_view"),
    3: _guard_docx501_row3,          # docx-501 (١)
    4: _needles("api.py", "SILK_REQUIRE_PERSISTENT_DATA_DIR",
                "SILK_DATA_DIR غير مضبوط"),
    5: _needles("silk_storage.py", "def create_research_run",
                "def load_mission_checkpoints"),
    6: _needles("silk_llm_runtime.py", "_JSON_PARSE_FAILURE_GAP"),
    7: _needles("silk_data_layer.py", "_WB_INDICATOR_SOURCE"),
    8: _needles("silk_data_layer.py", "class DataPoint"),
    9: _absent("web/index.html", 'id="snapBtn"'),
    10: _needles("docs/AUDIT_STATUS.md", "قراءة فقط", "غير موجود"),
    11: _guard_docx501_row11,         # docx-501 (٣)
    12: _needles("silk_render.py", "_reconcile_mission_limits", "_first_clause"),
    13: _guard_docx501_row13,         # docx-501 (٢، الفشل الحيّ)
    14: _needles("silk_quality_gate.py", "_check_confidentiality_leaks",
                 "_check_style"),
    15: _needles("tools/live_shape_server.py", "class LiveShapeServer",
                 "def seed_db"),
    16: _needles("silk_ai_judge.py", "_WRITER_MAX_TOKENS", "_MAX_TOKENS_CEILING",
                 "max_tokens=_MAX_TOKENS_CEILING"),
    17: _guard_datapoint_repr_flexible,  # هجوم المشرف — ريبر DataPoint المرن
    18: _guard_vendor_name_leak,         # بلاغ UK — تسريب اسم مزوّد للعميل
    19: _guard_export_format_contract,   # بلاغ المُشرِف — زرّ PDF كان ينزّل docx
    20: _guard_world_tier2_no_fabrication,  # الميزة أ — لا تلفيق فئة-٢/تفجّر ميزانية
    21: _guard_intake_no_silent_guess,      # الميزة ب — لا اختلاق منتج من صورة
    22: _guard_out_of_coverage_thin_study,  # الميزة أ — سوق خارج التغطية لا دراسة هزيلة
    23: _guard_unresolved_hs_silent_spend,  # Wave 1 — الفيتوتشيني: لا إنفاق برمز HS مجهول
    24: _guard_hardcoded_product_rule,      # Wave 1 — الحارسان قاعدتان مبنيّتان على البيانات
    25: _guard_wrong_direction_study,       # Wave 1.5 A — أشقّاء «الدراسة بالاتجاه الخاطئ»
    26: _guard_silent_external_failure,     # Wave 1.5 C — لا فشلٌ صامت لخدمةٍ خارجية
    27: _guard_readiness_before_spend,      # Wave 1.5 D — كلُّ تدهورٍ قبل الحجز
    28: _guard_leads_table_hygiene,         # Wave 2 — نقاء جدول الروابط (جغرافيا/نثر/حشو)
    29: _guard_report_arabic_shape_a4,      # Wave 2 — «سلك» متّصلة + A4
    30: _guard_client_template_no_hardcoded_product,  # Wave 2 — لا منتج مثبَّت في القوالب
    31: _guard_analyze_persist_canonical_db,   # /analyze — التخزين للقاعدة القانونية لا قرصٍ نسبيّ فانٍ
    32: _guard_report_quality_upgrade,         # ترقية جودة التقرير — إصلاحُ المحرّك لا تحرير التقرير
    33: _guard_parse_provenance_not_prose,     # التقادُم من المصدر لا النثر (قرار المالك)
    34: _guard_new_source_contracts,           # دمج مصادر جديدة — نفس العقود (فجوة/ops/مخزَّن/محكوم/نظيف)
    35: _guard_hs_gate_shared_choke_point_fail_safe,  # تقرير الكويت — بوّابة HS فشل-آمن + نقطة اختناق مشتركة
    36: _guard_cross_market_checkpoint_leak,          # تقرير الكويت — تسرّب يمن↔كويت عبر نقاط تفتيش بعثات
    37: _guard_golden_contract_test_exists_and_covers_both_paths,  # الاختبار الذهبي — كل العقود، كلا المسارين
    38: _guard_watchdog_owner_only_no_client_contamination,  # الحارس — مراقبةٌ للمالك حصراً، صفر تلوّث للعميل
    39: _guard_general_hs_classifier_no_lookup_table_ceiling,  # المصنّف العام — جدول البحث تلميحٌ ابتدائي لا حاكمٌ نهائي
    40: _guard_ui_tier_consumption_single_choke_point,  # UI-ONLY FIX — نقطة اختناق tier واحدة، لا مسار ثانٍ يثق بـhs6 خامًا
    41: _guard_active_resolution_beats_rejected_and_short_root_collision,  # ONE FIX — المصادَق يتصدّر على المرفوض، لا تصادف جذرٍ قصير
    42: _guard_dza_quality_gate_six_findings,  # تحليل #1 DZA — ست نتائج فشل بوّابة الجودة معاً على تشغيلة واحدة
    43: _guard_hs_classifier_valve_fail_safe_default,  # المُصنِّف العام — صمّامٌ فشل-آمن مفعَّل افتراضياً لا مُطفأ
    44: _guard_verdict_tone_recognizes_arabic_labels,  # Master Prompt Part 2 §B — _verdict_tone تتعرّف على التسمية العربية أيضاً
    45: _guard_price_fix_scoped_to_table_window,  # دالة إصلاح عملة السعر مقيَّدة بنافذة الجدول لا كامل المستند
    46: _guard_quality_gate_is_client_export_delivery_condition,  # حزمة v2.1 — بوابة الجودة شرط تسليم + عائلة فحوصات كاتب/عرض
    47: _guard_wp1_verdict_determinism,        # WP-1 — حتمية الحكم ومصدره الواحد
    48: _guard_wp2_no_raw_internal_output,     # WP-2 — لا مخرَج داخلي خام للعميل
    49: _guard_wp3_evidence_integrity,         # WP-3 — نزاهة الأدلة والمصالحة
    50: _guard_wp4_gaps_consistency,           # WP-4 — اتساق الفجوات مع الختام
    51: _guard_wp5_rtl_bracket_isolation,      # WP-5 — عزل أقواس RTL + فحص PDF
    52: _guard_wp6_injector_adversarial_locks,  # WP-6 — أقفال الحاقنات العدائية
    53: _guard_wp7_delivery_gate_hardening,    # WP-7 — تصليب بوابة التسليم
    54: _guard_zero_confidence_finding_declared_gap,  # بند بثقة 0.0 => فجوة معلنة لا بند (خرق حارس المراقبة الحي)
    56: _guard_coverage_gate_year_fallback,    # تدقيق v2 الموجة ١ — سُلَّم سنوات بوّابة التغطية
    57: _guard_sanitizer_obfuscation_variants,  # الموجة ١ — ست صيغ تشويش المشرف
    58: _needles("CLAUDE.md", "/code-review",   # المراجعة الذاتية قبل فتح/وسم أي PR جاهزًا (Yemen stale-tag)
                 "self-review catches what hermetic tests structurally cannot"),
    55: _needles("tests/conftest.py", "def _hermetic_env_guard"),  # عزل SILK_HERMETIC لكل اختبار — لا تسرّب لافتة «نموذج توضيحي»
    59: _guard_wave2_med_hardening,   # الموجة ٢ — خمسة إصلاحات MED من تدقيق v2
    60: _guard_composite_source_id_attribution,   # بلاغ قطر HF1 — إسنادٌ ذرّيّ لا مركّب
    61: _guard_renderer_truncation_and_empty_parens,  # بلاغ قطر HF2 — لا بترٌ داخل رقم/قوسٌ فارغ
    62: _guard_cross_source_plausibility,         # بلاغ قطر HF3 — حارسُ معقوليةٍ عبر المصادر
    63: _guard_bloc_list_single_source,           # DEF-2 — عضويةُ الكتلة من مصدرٍ واحد (EU27 كاملة)
    64: _guard_g41_domestic_production,           # DEF-1/G4.1 — مرتكزُ الإنتاج المحليّ (سوقٌ مُنتِجة لا تُوسَم)
    65: _guard_ask_what_the_product_answers,      # بلاغ المُشرِف — قِسِ الرقمَ قبل أن تسأل عنه
    66: _guard_band_boundary_strictness_and_second_axis,  # صرامةُ الحدّ + المحورُ الثاني
    67: _guard_dimension_terms_not_frozen_in_code,        # مصطلحُ بُعدٍ مجمَّد (عودةُ ٣٠)
    68: _guard_multi_axis_heading_confident_wrong_code,    # F2 — محورٌ ثانٍ غيرُ رقميّ
    69: _guard_client_operator_document_divergence,        # F3 — تباعدُ مُسلَّمَي العميل/المشغّل
    70: _guard_attribute_resolver_flag_off_by_default,     # D1 — صمّامٌ مُطفأٌ افتراضياً
    71: _guard_cross_basis_edge_refusal,                   # D2 — أساسُ النسبة قربَ الحافّة
    72: _guard_dialog_band_text_from_the_official_reference_only,  # E2 — نصُّ الحدّ من المرجع حصراً
    73: _guard_dialog_axis_siblings_never_partial,         # E3 — لا مجموعةَ محورٍ جزئية
    74: _guard_dialog_prose_carries_no_product_brand_or_country,  # E4 — لا صدى منتج/علامة/دولة
    75: _guard_gate_passes_synthetic_but_silent_on_real,  # تحليل ٧ — بوّابة مرّت التركيبيّ ثم صمتت على الحقيقيّ (قفلان لكلّ بوّابة)
    76: _guard_seat_lock_is_load_bearing,  # PR-2 — قفل المقعد وحارسه المُميِّز
    77: _guard_readiness_names_the_offending_variable,  # #197 — تشخيصٌ بلا سبب
    78: _guard_image_evidence_decides_prepared_form,  # حليب الفراولة — أدلة الصورة تحسم
}

_TRAPS = [
    ("mock_passes_real_fails", "Mock-passes / real-fails",
     _needles("silk_render.py", "_deep_research_view")),
    ("markets_empty_misroute", "misroutes exporters",
     _guard_trap_markets_empty),
    ("two_sanitizers", "Two different sanitizers",
     _guard_trap_two_sanitizers),
    ("redaction_mangling", "Redaction mangling",
     _guard_trap_redaction_mangling),
    ("parallel_cache_window", "Parallel missions and the cache window",
     _guard_trap_parallel_cache_window),
    ("cap_counts_not_dollars", "Cap counted operations, not dollars",
     _needles("silk_usage.py", "def try_reserve_usd")),
    ("styled_not_wired", "Styled-but-never-wired UI affordance",
     _needles("web/index.html", 'data-id="',
              '$("#histList").addEventListener("click"')),
    ("silent_noop_family", "silent no-op has three forms",
     _needles("web/index.html", "function openStoredAnalysis",
              # حزمة الإغلاق، البند ٤: سلسلة /markets في بناء نيّة الدردشة
              # اكتسبت .catch عربياً (كانت ترفض صامتةً فتُعلّق مؤشّر الانتظار).
              "تعذّر تحميل قائمة الأسواق")),
    ("view_after_persist", "view attached AFTER persist",
     _guard_trap_view_after_persist),
    ("strip_plumbing_three_leaks", "leaked three raw forms",
     _guard_trap_strip_plumbing_three_leaks),
    ("orphan_reservation_leak", "Orphaned runs leak their USD reservation",
     lambda: (
         _needles("silk_storage.py", "def reap_orphan_research_runs",
                  "reconcile_usd", "SILK_ORPHAN_STALE_MINUTES")(),
         _needles("api.py", "reap_orphan_research_runs")(),
         _needles("silk_collectors.py", "reap_orphan_research_runs")())),
]


# ── حارس واحد لكل حادثة (تُوسَّع برمجياً لتقارير pytest واضحة) ────────────────

@pytest.mark.parametrize("row", sorted(_LESSONS), ids=[f"lessons-{n}" for n in sorted(_LESSONS)])
def test_lessons_incident_guard_holds(row):
    """كل صفّ في docs/LESSONS.md له حارس حيّ يُفشِل على عودة عائلته."""
    _LESSONS[row]()


@pytest.mark.parametrize("slug,match,check", _TRAPS,
                         ids=[t[0] for t in _TRAPS])
def test_operations_trap_guard_holds(slug, match, check):
    """كل فخّ في silk-operations §2 (THE TRAPS) له حارس حيّ."""
    check()


# ── الاختبار الشامل: التغطية كاملة ضدّ كلا السِّجلّين ─────────────────────────

def _lessons_rows_in_ledger() -> set[int]:
    ledger = _read("docs/LESSONS.md")
    return {int(m.group(1))
            for m in re.finditer(r"^\|\s*(\d+)\s*\|", ledger, re.M)}


def _trap_rows_in_skill() -> list[str]:
    """صفوف بيانات جدول §2 (THE TRAPS) — الخلية الأولى (اسم الفخّ العريض)."""
    skill = _read(".claude/skills/silk-operations/SKILL.md")
    m = re.search(r"## 2\. THE TRAPS(.*?)\n## 3\.", skill, re.S)
    assert m, "قسم §2 THE TRAPS غير موجود في مهارة silk-operations"
    rows = []
    for line in m.group(1).splitlines():
        line = line.strip()
        if not line.startswith("| **"):        # صفوف البيانات فقط
            continue
        first_cell = line.split("|")[1].strip()
        rows.append(first_cell)
    return rows


def test_meta_registry_covers_every_known_incident():
    """التغطية الشاملة: كل صفّ LESSONS وكل فخّ §2 مُسجَّل هنا بحارس، ولا مدخلة
    يتيمة بلا صفّ مقابل. حادثة جديدة تسقط في أي سِجلّ بلا حارس تُحمِّر هنا."""
    # (أ) LESSONS: مفاتيح السجلّ = أرقام الصفوف بالضبط، متتابعة ١..N.
    ledger_rows = _lessons_rows_in_ledger()
    assert ledger_rows == set(range(1, max(ledger_rows) + 1)), (
        f"أرقام صفوف LESSONS غير متتابعة: {sorted(ledger_rows)}")
    registry_rows = set(_LESSONS)
    assert registry_rows == ledger_rows, (
        f"صفوف LESSONS بلا حارس في السجلّ: {sorted(ledger_rows - registry_rows)}؛ "
        f"حُرّاس بلا صفّ: {sorted(registry_rows - ledger_rows)}")

    # (ب) TRAPS: كل صفّ فخّ في §2 يطابقه حارس واحد بالضبط عبر إبرة `match`،
    # وكل حارس فخّ يطابق صفّاً واحداً على الأقل (لا يتيم).
    trap_rows = _trap_rows_in_skill()
    assert trap_rows, "لم تُقرَأ صفوف فخاخ من المهارة"
    for row in trap_rows:
        matched = [slug for slug, match, _ in _TRAPS if match in row]
        assert len(matched) == 1, (
            f"صفّ الفخّ «{row[:60]}…» يطابقه {len(matched)} حُرّاس (المتوقّع ١): "
            f"{matched}")
    for slug, match, _ in _TRAPS:
        assert any(match in row for row in trap_rows), (
            f"حارس الفخّ «{slug}» (match={match!r}) لا يطابق أيّ صفّ في §2 — "
            "يتيم؛ حدِّث السجلّ أو المهارة")


def test_meta_docx501_trio_all_have_behavioral_guards():
    """الحوادث الثلاث لعطل docx-501 (LESSONS ٣/١١/١٣) لها حُرّاس **سلوكية**
    (تبني/تنقّي فعلياً)، لا مجرّد وجود رمز — أمر المُشرِف الصريح."""
    behavioral = {3: _guard_docx501_row3, 11: _guard_docx501_row11,
                  13: _guard_docx501_row13}
    for row, guard in behavioral.items():
        assert _LESSONS[row] is guard, (
            f"صفّ docx-501 رقم {row} ليس مربوطاً بحارسه السلوكي")
        guard()  # يجب أن يمرّ فعلياً الآن
