"""Deep-research (Top-5) report: engine seam, worker task, and endpoints (ADR-0007).

The paid, key-gated deep-research report is DISTINCT from the always-available
executive report. These tests lock the fail-closed contract end to end WITHOUT
ever calling a real paid API — the engine mission runner / synthesis / the
product seam are all mocked:

* keyless ⇒ ``engine.run_deep_research`` returns ``None`` with ZERO paid calls;
* keyless worker ⇒ a declared-gap "pending API key" docx, no fabricated narrative;
* keyed worker (seam mocked) ⇒ a combined docx with the Top-5 names + source lines;
* trigger/poll/download are tenant-scoped and fail closed before generation.
"""

from __future__ import annotations

import types
import uuid

import pytest

from app.config import get_settings
from tests.test_report_docx import _seed_funnel


def _set_key(monkeypatch, value: str) -> None:
    monkeypatch.setenv("ANTHROPIC_API_KEY", value)
    get_settings.cache_clear()


@pytest.fixture(autouse=True)
def _restore_settings():
    yield
    get_settings.cache_clear()


def _fake_market(iso3="NGA", name_en="Nigeria", name_ar="نيجيريا"):
    return types.SimpleNamespace(iso3=iso3, iso2="NG", m49="566", name_en=name_en, name_ar=name_ar)


def _fake_datapoint(value, source, note):
    return types.SimpleNamespace(
        value=value, source=source, confidence=0.8, note=note, data_year=2022
    )


def _fake_report(summary, findings, failed=False):
    return types.SimpleNamespace(summary=summary, findings=findings, failed=failed)


# -- engine seam: fail-closed + scoped -------------------------------------


def test_run_deep_research_gated_makes_zero_paid_calls(monkeypatch):
    """No ANTHROPIC_API_KEY ⇒ None and the engine mission runner is NEVER invoked."""
    import silk_missions

    from app.services import engine

    _set_key(monkeypatch, "")

    called = {"n": 0}

    def _spy(*a, **k):
        called["n"] += 1
        raise AssertionError("deep_research must not run without an API key")

    monkeypatch.setattr(silk_missions, "deep_research", _spy)

    out = engine.run_deep_research("Nigeria", "dates", "080410")

    assert out is None
    assert called["n"] == 0


def test_run_deep_research_keyed_runs_inside_deepen_and_budget_scope(monkeypatch):
    """With a key: runs the engine inside the deepen + budget scope and normalizes."""
    import silk_market_resolver
    import silk_missions
    import silk_synthesis

    from app.services import engine
    from app.services.api_budget import current_budget

    _set_key(monkeypatch, "sk-test")

    seen = {"deepen": None, "budget": None}

    def _fake_deep_research(market, product="", hs_code=None, **k):
        # The paid engine runs ONLY inside the re-armed deepen scope (I5) and the
        # per-report budget scope (decision #5).
        seen["deepen"] = engine.deepen_active()
        seen["budget"] = current_budget() is not None
        return {
            "reports": {
                "trade": _fake_report(
                    "trade flows found",
                    [_fake_datapoint(1000.0, "UN Comtrade", "total imports")],
                ),
                "risk_news": _fake_report("no sources", [], failed=True),
            },
            "trace_id": "trace-xyz",
        }

    monkeypatch.setattr(silk_market_resolver, "resolve_market", lambda q: (_fake_market(), []))
    monkeypatch.setattr(silk_missions, "deep_research", _fake_deep_research)
    monkeypatch.setattr(
        silk_synthesis,
        "synthesize",
        lambda reports, **k: {"verdict": "PRELIMINARY GO", "confidence": 0.5},
    )

    out = engine.run_deep_research("Nigeria", "dates", "080410")

    assert seen["deepen"] is True
    assert seen["budget"] is True
    assert out is not None
    assert out["market"]["name_en"] == "Nigeria"
    assert out["trace_id"] == "trace-xyz"
    # The successful mission is a section with a sourced finding …
    trade = next(s for s in out["sections"] if s["key"] == "trade")
    assert trade["findings"][0]["source"] == "UN Comtrade"
    # … the failed mission becomes a declared-gap line (I1), never a fabricated one.
    assert any("risk" in g.lower() or "no sources" in g for g in out["gaps"])
    assert out["verdict"] == "PRELIMINARY GO"


def test_run_deep_research_unresolved_market_is_a_declared_gap(monkeypatch):
    """A market that cannot be resolved ⇒ None (never a guess), zero engine calls."""
    import silk_market_resolver
    import silk_missions

    from app.services import engine

    _set_key(monkeypatch, "sk-test")
    monkeypatch.setattr(silk_market_resolver, "resolve_market", lambda q: (None, ["Nigeria?"]))

    def _spy(*a, **k):
        raise AssertionError("must not run when the market is unresolved")

    monkeypatch.setattr(silk_missions, "deep_research", _spy)

    assert engine.run_deep_research("Nowhereland", "dates", None) is None


# -- worker task -----------------------------------------------------------


def _docx_text(data: bytes) -> str:
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_render_research_report_keyless_is_a_declared_gap(
    db, factory, product, monkeypatch, tmp_path
):
    """Keyless ⇒ 'pending API key' declared-gap docx, no fabricated narrative."""
    pytest.importorskip("docx")
    from app.services import storage as storage_mod
    from app.workers.tasks import render_research_report

    _set_key(monkeypatch, "")
    _seed_funnel(db, product)
    store = storage_mod.LocalStorage(str(tmp_path))
    monkeypatch.setattr(storage_mod, "get_storage", lambda *a, **k: store)

    out = render_research_report.delay(str(product.id)).get()

    key = f"reports/research/{product.id}.docx"
    assert out == {"product_id": str(product.id), "status": "gated", "report_key": key}

    data = store.get_bytes(key)
    assert data is not None and data[:2] == b"PK"
    text = _docx_text(data)
    # The honest gated line is present; a real verdict/narrative is NOT invented.
    assert "pending API key" in text
    assert "Preliminary verdict" not in text
    # The Top-5 target markets are still named (from the persisted funnel).
    assert "DEU" in text and "NLD" in text and "IND" in text

    db.expire_all()
    from app.models import Product

    refreshed = db.get(Product, product.id)
    assert refreshed.research_status == "gated"
    assert refreshed.research_report_key == key


def test_render_research_report_keyed_renders_combined_docx(
    db, factory, product, monkeypatch, tmp_path
):
    """Keyed, engine seam MOCKED ⇒ combined docx with Top-5 names + source labels."""
    pytest.importorskip("docx")
    from app.services import engine
    from app.services import storage as storage_mod
    from app.workers.tasks import render_research_report

    _set_key(monkeypatch, "sk-test")
    _seed_funnel(db, product)
    store = storage_mod.LocalStorage(str(tmp_path))
    monkeypatch.setattr(storage_mod, "get_storage", lambda *a, **k: store)

    def _fake_run(market_name, product_name, hs_code):
        return {
            "market": {"iso3": market_name, "iso2": "", "name_en": market_name, "name_ar": ""},
            "verdict": "PRELIMINARY GO",
            "verdict_confidence": 0.6,
            "verdict_reasoning": "solid demand",
            "sections": [
                {
                    "key": "trade",
                    "label": "وكيل التجارة",
                    "summary": "trade flows",
                    "failed": False,
                    "findings": [
                        {
                            "value": 1234.0,
                            "source": "UN Comtrade",
                            "note": "total imports",
                            "confidence": 0.8,
                            "data_year": 2022,
                        }
                    ],
                }
            ],
            "gaps": [],
        }

    monkeypatch.setattr(engine, "run_deep_research", _fake_run)

    out = render_research_report.delay(str(product.id)).get()

    key = f"reports/research/{product.id}.docx"
    assert out == {"product_id": str(product.id), "status": "ready", "report_key": key}

    text = _docx_text(store.get_bytes(key))
    # Every Top-5 market from the funnel is a section …
    assert "DEU" in text and "NLD" in text and "IND" in text
    # … each figure carries its engine source label (per-figure provenance, I1).
    assert "UN Comtrade" in text
    assert "source" in text
    assert "PRELIMINARY GO" in text


# -- endpoints -------------------------------------------------------------


def test_research_endpoints_are_tenant_scoped_and_fail_closed(
    client, auth_headers, db, factory, product, monkeypatch, tmp_path
):
    pytest.importorskip("docx")
    from app.services import storage as storage_mod

    _set_key(monkeypatch, "")  # keyless ⇒ eager task reaches the gated terminal state
    _seed_funnel(db, product)
    store = storage_mod.LocalStorage(str(tmp_path))
    monkeypatch.setattr(storage_mod, "get_storage", lambda *a, **k: store)

    base = f"/api/v1/products/{product.id}/report/research"

    # Unauthenticated is rejected.
    assert client.post(base).status_code == 401
    assert client.get(f"{base}.docx").status_code == 401

    # Download BEFORE generation fails closed (409, never a fabricated file).
    pre = client.get(f"{base}.docx", headers=auth_headers)
    assert pre.status_code == 409

    # Trigger (eager) → reaches the gated terminal state.
    res = client.post(base, headers=auth_headers)
    assert res.status_code == 202
    assert res.json()["status"] == "gated"

    # Poll.
    status = client.get(f"{base}/status", headers=auth_headers)
    assert status.status_code == 200
    body = status.json()
    assert body["status"] == "gated" and body["ready"] is True

    # Download now returns the declared-gap docx.
    dl = client.get(f"{base}.docx", headers=auth_headers)
    assert dl.status_code == 200
    assert dl.content[:2] == b"PK"

    # Unknown product is a 404 (tenant isolation via get_owned_product).
    missing = client.post(f"/api/v1/products/{uuid.uuid4()}/report/research", headers=auth_headers)
    assert missing.status_code == 404
