"""أدوات اختبار مشتركة — shared test helpers (M0).

يوفّر `block_network` القانوني الواحد بدل النسخ المكرَّرة في ملفات الموجات
(الأثر التاريخي يُنظَّف في M9). الاختبارات الجديدة تستورد من هنا حصراً.
Canonical network guard for hermetic tests; new tests import from here only.
"""
import contextlib
import os
import socket
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# الاختبارات المحكمة تبني التطبيق بلا مفتاح خدمة عمداً؛ الإغلاق المسبق الجديد
# (تدقيق 2026-08-07 C6) يتطلب الاختيار الصريح لوضع التطوير المفتوح.
# The hermetic suite builds the app keyless on purpose; the C6 fail-closed
# guard requires this explicit open-dev opt-in (production never sets it).
os.environ.setdefault("SILK_ALLOW_OPEN_DEV", "1")


@contextlib.contextmanager
def block_network():
    """اقطع الشبكة مؤقتاً — make outbound sockets fail so 'no data' paths hold.

    بلاغ حي (تسريب تسلسل اختبارات CI): جلسة requests المشتركة الدائمة
    (silk_data_layer._session — تجميع اتصالات keep-alive للأداء الإنتاجي)
    قد تحمل اتصالاً TCP حياً فعلياً تركه نداء سابق غير محظور في نفس عملية
    pytest (تشغيل تسلسلي واحد لكل ملفات tests/). إعادة استعمال اتصال
    مجمَّع قائم لا يستدعي socket.socket() من جديد، فيتجاوز الحجب أدناه
    صامتاً ويُرجع بيانات حقيقية رغم دخول هذا السياق — ظهر هذا حين أضاف
    ملف اختبار جديد بضعة نداءات فأزاح ترتيب التنفيذ فكشف اتصالاً مجمَّعاً
    كان يبقى خاملاً غير مستغَل سابقاً. إغلاق تجمّعات الاتصال المعروفة عند
    كل دخول يمنع نجاة اتصال حيّ لاختبار يُفترض به حجب كامل — Session.close()
    يُغلق التجمّع الحالي فقط لا الكائن نفسه، فيُعاد فتح اتصال جديد طبيعياً
    خارج هذا السياق حين تُستأنف الشبكة.
    """
    real = socket.socket

    def _no_net(*a, **k):  # noqa: ANN002, ANN003
        # صياغة بلا كلمة hermetic عمداً: حارس تقارير الإنتاج يرفض أي أثر يحمل
        # الكلمة (إصلاح مراجعة Stage 5) — قطع الشبكة حالة تشغيل صادقة لا بديل
        # بيانات، فلا يجوز أن تسمّم ملاحظاتُه تقريراً مشتقاً في اختبار.
        raise OSError("network disabled for offline test")

    try:
        import silk_data_layer
        silk_data_layer._session.close()
    except Exception:  # noqa: BLE001 — أفضل جهد؛ الحجب الأساسي (socket) نافذ بدونه
        pass

    socket.socket = _no_net
    try:
        yield
    finally:
        socket.socket = real


def docx_all_text(path: str) -> str:
    """كل نص مستند Word — فقرات + خلايا جداول (مراجعة المشروع: بعض أقسام
    render_docx صارت جداولاً حقيقية بدل نقاط سردية؛ `doc.paragraphs` وحدها
    لا تصل خلايا الجداول، فتفوّت اختباراتٌ محتوًى انتقل إليها بلا انحدار فعلي).
    """
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


import tempfile

import pytest


@pytest.fixture(autouse=True)
def _isolated_fact_store(monkeypatch):
    """عزل مخزن الحقائق لكل اختبار — كتابة M2 العابرة دفّأت المخزن الافتراضي
    فتسرّبت حقائق حقيقية بين الاختبارات (اكتُشف عبر test_engine_localprice_layer_offline
    بعد تشغيلات تدقيق Stage 1). Every test gets its own store unless it overrides."""
    monkeypatch.setenv("SILK_STORE_DB",
                       os.path.join(tempfile.mkdtemp(), "store.db"))
    # قاعدة المنصّة (المستأجرون/المصادقة/المحافظ) تُعزَل هنا أيضاً: api.py الجذر
    # يركّب /platform داخل create_app()، فأيّ اختبار قديم يمسّ مساراً تحتها كان
    # سيهيّئ `data/platform.db` في شجرة العمل ويسرّب حالة بين الاختبارات — نفس
    # عائلة الحادثة التي وُلد لها هذا التثبيت. Isolate the 5th store too.
    monkeypatch.setenv("SILK_PLATFORM_DB",
                       os.path.join(tempfile.mkdtemp(), "platform.db"))
    # 1b: عطّل مباعدة النداءات في الاختبارات — الشبكة مقطوعة أصلاً، والمباعدة
    # 250ms × مئات النداءات الفاشلة كانت ستبطئ الحزمة بلا فائدة.
    monkeypatch.setenv("SILK_HTTP_MIN_GAP_MS", "0")
    # نافذة كومتريد الخاصة (بلاغ 429، افتراضي 1100ms) تُصفَّر أيضاً — بلا
    # هذا نامت الحزمة الهيرمتية ~ساعتين فعلياً (1.1ث × مئات نداءات كومتريد
    # المقطوعة الشبكة) — اكتُشف حياً عند إضافة النافذة.
    monkeypatch.setenv("SILK_COMTRADE_MIN_GAP_MS", "0")
    # الموجة ٦ (V5): عزل ملفات التتبّع أيضاً — بلا هذا، اختبارات /research
    # الحقيقية (TestClient) تكتب data/traces/*.jsonl فعلياً على القرص.
    monkeypatch.setenv("SILK_TRACE_DIR", tempfile.mkdtemp())
    # عزل عدّاد data_economics بين الاختبارات — contextvar بلا حدود عملية
    # مستقلة (pytest يُشغّل كل الاختبارات على نفس الخيط)، فاختبار سابق ترك
    # عدّاداً بأرقام عالية كان سيُفعِّل سقف silk_llm_runtime._run_loop
    # الكلي زوراً في اختبار لاحق لا علاقة له (انحدار اكتُشف فعلياً، الموجة
    # ٦). الإنتاج غير متأثر: كل طلب /research يستدعي begin_data_counter()
    # صراحة قبل أي استخدام.
    import silk_context
    silk_context._data_counter.set(None)
    # لا حاجة لتصفير خنق الدخول يدوياً: حالته صارت في قاعدة المنصّة (المعزولة
    # لكل اختبار بالسطر أعلاه) لا في ذاكرة الوحدة — فالعزل يأتي مجّاناً.
    # Login-throttle state lives in the (per-test isolated) platform DB.
    #
    # عامل عمل كلمات المرور مُخفَّض **في الاختبارات فقط**: قياس فعلي — تجزئة
    # bcrypt بعامل ١٢ = ~٢٧٧ms وتحقّق = ~٢٧٣ms، فحزمة المنصّة قفزت ٢٥ث→١٣٣ث
    # واتجاهها يسوء مع كل اختبار جديد. الإنتاج غير متأثّر بثلاث حمايات
    # (الافتراضي ١٢، وقيمة تالفة ⇒ ١٢، وحارس الإقلاع يرفض أقلّ من ١٢ في الإنتاج)،
    # واختبار عامل ١٢ يمسح هذا المتغيّر ويدفع تكلفة تجزئة حقيقية واحدة ليُثبِته.
    # Test-only work-factor reduction; production is protected three ways.
    monkeypatch.setenv("SILK_PLATFORM_BCRYPT_ROUNDS", "4")
    monkeypatch.setenv("SILK_PLATFORM_SCRYPT_N", "1024")


@pytest.fixture(autouse=True)
def _isolated_circuit_breaker():
    """عزل قاطع الدائرة المشترك لكل اختبار — وكيل التعريفات صار يسجّل أعطاله
    في `silk_circuit.http_breaker` (موجة المنصة ٣)، فاختبارٌ يُفشل الشبكة
    عمداً خمسَ مرات كان يفتح القاطع لكل اختبارٍ لاحقٍ في نفس التشغيلة
    (test_wits_400 يرى ملاحظة «قاطع مفتوح» بدل ملاحظة 400 النظيفة). كل
    اختبار يبدأ وينتهي بقاطعٍ صافٍ."""
    import silk_circuit

    silk_circuit.http_breaker.reset()
    yield
    silk_circuit.http_breaker.reset()


@pytest.fixture(autouse=True)
def _hermetic_env_guard():
    """عزل علم `SILK_HERMETIC` لكل اختبار — اختبارات كانت تضبطه خاماً
    (`os.environ[...] = "1"`) بلا استرجاع فيتسرّب لكل اختبار لاحق: build_view
    يسم test_run زوراً فتظهر لافتة «نموذج توضيحي» في PDF عميلٍ إنتاجي المسار
    (اكتُشف بإعادة إنتاج مباشرة: القفل البصري
    test_visual_pdf_lock_production_entrypoint_bare_no_split_no_leaks يسقط على
    «⚠» فقط حين تسبقه اختبارات التصدير في نفس الحزمة وأدوات PDF مثبَّتة).
    كل اختبار يبدأ بلا العلم، وأي ضبطٍ خام داخله يُمسَح بعده حتماً."""
    old = os.environ.pop("SILK_HERMETIC", None)
    yield
    if old is None:
        os.environ.pop("SILK_HERMETIC", None)
    else:
        os.environ["SILK_HERMETIC"] = old


# ── طبقة الدخان الحية (opt-in) — gated live-integration lane ─────────────────
# اختبارات موسومة `live` تضرب الشبكة الحقيقية (مصادر مجانية بلا مفتاح فقط —
# لا حرق أرصدة). تُتخطّى دائماً في CI الافتراضي (`pytest tests/ -q`) وتعمل
# فقط حين SILK_RUN_LIVE=1 (مسار workflow_dispatch يدوي، راجع
# .github/workflows/live-smoke.yml). هذا يبدأ إغلاق أكبر فجوة اختبار: غياب
# أي اختبار تكامل حي (كل شيء آخر يقطع الشبكة).

def pytest_configure(config):
    config.addinivalue_line(
        "markers",
        "live: real-network integration smoke test — skipped unless "
        "SILK_RUN_LIVE=1 (opt-in lane, never runs on default CI).")
    # رُتبتا ٢–٣ (خادم حقيقي + متصفّح حقيقي) — تُقلِعان uvicorn فعلياً (وrung 3
    # تشغّل chromium)، أبطأ من الحزمة الهرمتية وخارج ضمانتها «بلا شبكة/بلا
    # عملية خارجية». تُجمَع في وظيفة CI المخصّصة `e2e-live-shape` حصراً
    # (SILK_RUN_E2E=1)، فتبقى `pytest tests/ -q` الافتراضية هرمتية سريعة.
    config.addinivalue_line(
        "markers",
        "e2e: real-server / real-browser rung — boots uvicorn (rung 2) and "
        "may drive chromium (rung 3); skipped unless SILK_RUN_E2E=1 "
        "(the e2e-live-shape CI job — see .github/workflows/e2e-live-shape.yml).")


def pytest_collection_modifyitems(config, items):
    live_on = os.environ.get("SILK_RUN_LIVE") == "1"
    e2e_on = os.environ.get("SILK_RUN_E2E") == "1"
    skip_live = pytest.mark.skip(
        reason="live-network test; set SILK_RUN_LIVE=1 to run "
               "(opt-in lane — see .github/workflows/live-smoke.yml)")
    skip_e2e = pytest.mark.skip(
        reason="real-server/browser rung; set SILK_RUN_E2E=1 to run "
               "(the e2e-live-shape CI job).")
    for item in items:
        if "live" in item.keywords and not live_on:
            item.add_marker(skip_live)
        if "e2e" in item.keywords and not e2e_on:
            item.add_marker(skip_e2e)
