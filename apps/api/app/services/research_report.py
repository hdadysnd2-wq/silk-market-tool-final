"""Combined factory-facing deep-research report (ADR-0007), rendered product-side.

The engine (``silk_missions.deep_research``, via ``engine.run_deep_research``) is
the single research brain: 12 Claude missions per market, each returning sourced
``DataPoint`` findings. This module is the BODY half — it composes those
per-market findings into one combined docx for the product's Top-5 markets, the
same product-side composition pattern the executive report uses.

Two honesty rules carry through (I1):

* every figure prints its engine ``source`` label — a finding without a source is
  structurally impossible here because the engine attaches one to each DataPoint;
* a missing figure / failed mission is a DECLARED GAP line, never a fabricated
  value; and with no ``ANTHROPIC_API_KEY`` the whole document is the declared-gap
  "deep research pending API key" report (:func:`render_pending_docx`), never a
  synthesized narrative.

Reads ONLY the deep-research output (never the funnel/analyze template), so the
deep-research report stays distinct from the executive one.
"""

from __future__ import annotations

from sqlalchemy.orm import Session

from app.models import Market, Product
from app.providers.countries import iso3_to_iso2
from app.services.report_view import TOP_MARKETS, _analysis_rankings, _latest_analysis

_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

#: Bilingual report chrome — Arabic-first (the docx is Arabic-primary, like the
#: engine reports), with an English mirror so an English-speaking factory reads it.
_TITLE = "تقرير البحث العميق — أفضل ٥ أسواق · Deep Research Report — Top 5 Markets"
_PENDING_LINE = (
    "البحث العميق بانتظار مفتاح واجهة البرمجة (ANTHROPIC_API_KEY) — deep research pending API key"
)
_NO_MARKETS_LINE = (
    "لا توجد أسواق مرشّحة بعد لهذا المنتج — شغّل تحليل القمع العالمي أولاً · "
    "No shortlisted markets yet — run the world funnel first."
)
_SOURCE_LABEL = "المصدر · source"
_GAPS_HEADING = "حدود هذا البحث · Limits of this research"
_VERDICT_LABEL = "الحكم الأولي · Preliminary verdict"


def top5_markets(db: Session, product: Product) -> list[tuple[str, str]]:
    """The product's persisted Top-5 markets as ``(iso3, display_name)`` pairs.

    Reads the latest analysis's world-funnel ``CountryRanking`` rows (rank order),
    resolving each ISO3 to a real market name where the platform has one. Empty
    when the product has no analysis yet (a declared absence, I1).
    """
    analysis = _latest_analysis(db, product)
    rankings = _analysis_rankings(db, analysis)
    out: list[tuple[str, str]] = []
    for ranking in rankings[:TOP_MARKETS]:
        iso3 = ranking.importer_iso3
        iso2 = iso3_to_iso2(iso3)
        market = db.get(Market, iso2) if iso2 else None
        out.append((iso3, (market.name_en if market else None) or iso3))
    return out


def _new_document():
    from docx import Document

    return Document()


def _add_header(doc, product_name: str, hs_code: str | None) -> None:
    doc.add_heading(_TITLE, level=0)
    doc.add_paragraph(f"{product_name}" + (f"  ·  HS {hs_code}" if hs_code else ""))


def _add_market_section(doc, report: dict) -> None:
    """Render one market's deep-research findings with a source line per figure."""
    market = report.get("market") or {}
    name_en = market.get("name_en") or market.get("iso3") or ""
    name_ar = market.get("name_ar") or ""
    heading = name_en + (f" · {name_ar}" if name_ar else "")
    doc.add_heading(heading, level=1)

    verdict = (report.get("verdict") or "").strip()
    if verdict:
        line = f"{_VERDICT_LABEL}: {verdict}"
        conf = report.get("verdict_confidence")
        if conf is not None:
            line += f"  ({conf})"
        doc.add_paragraph(line)
        reasoning = (report.get("verdict_reasoning") or "").strip()
        if reasoning:
            doc.add_paragraph(reasoning)

    for section in report.get("sections") or []:
        if section.get("failed"):
            # A failed mission's gap is collected into the limits section below;
            # do not print an empty findings block here.
            continue
        doc.add_heading(section.get("label") or section.get("key") or "", level=2)
        summary = (section.get("summary") or "").strip()
        if summary:
            doc.add_paragraph(summary)
        for finding in section.get("findings") or []:
            value = finding.get("value")
            note = (finding.get("note") or "").strip()
            source = (finding.get("source") or "").strip()
            if value is None and not note:
                continue
            label = note or "—"
            shown = "—" if value is None else str(value)
            para = doc.add_paragraph(style=None)
            para.add_run(f"• {label}: {shown}")
            # Per-figure source line — never dropped (I1). Even a gap value keeps
            # the honest source-or-declared-gap provenance.
            src = source or "فجوة معلنة · declared gap"
            para.add_run(f"   [{_SOURCE_LABEL}: {src}]")

    gaps = report.get("gaps") or []
    if gaps:
        doc.add_heading(_GAPS_HEADING, level=2)
        for gap in gaps:
            doc.add_paragraph(f"— {gap}")


def render_combined_docx(
    product_name: str, hs_code: str | None, per_market_reports: list[dict]
) -> bytes:
    """Render the combined Top-5 deep-research docx and return its bytes.

    ``per_market_reports`` is the list of normalized dicts from
    ``engine.run_deep_research`` (one per market). Each market becomes a section
    with its verdict, sourced mission findings, and declared-gap limits. Raises
    ``RuntimeError`` if python-docx is unavailable (mirrors the engine renderers).
    """
    import tempfile
    from pathlib import Path

    try:
        doc = _new_document()
    except ImportError as exc:  # python-docx missing
        raise RuntimeError("Word export is unavailable (python-docx not installed)") from exc

    _add_header(doc, product_name, hs_code)
    if not per_market_reports:
        doc.add_paragraph(_NO_MARKETS_LINE)
    for report in per_market_reports:
        _add_market_section(doc, report)

    tmp = Path(tempfile.mkdtemp(prefix="silk_research_")) / "research.docx"
    try:
        doc.save(str(tmp))
        return tmp.read_bytes()
    finally:
        import shutil

        shutil.rmtree(tmp.parent, ignore_errors=True)


def render_pending_docx(
    product_name: str, hs_code: str | None, markets: list[tuple[str, str]]
) -> bytes:
    """The fail-closed declared-gap report shown when no ``ANTHROPIC_API_KEY`` (C3/I1).

    Names the Top-5 markets that WOULD be researched and states plainly that deep
    research is pending an API key — never a fabricated narrative. Returns docx
    bytes so the download path is coherent even in the gated state.
    """
    import shutil
    import tempfile
    from pathlib import Path

    try:
        doc = _new_document()
    except ImportError as exc:
        raise RuntimeError("Word export is unavailable (python-docx not installed)") from exc

    _add_header(doc, product_name, hs_code)
    doc.add_paragraph(_PENDING_LINE)
    if markets:
        doc.add_heading("الأسواق المستهدفة · Target markets", level=1)
        for iso3, name in markets:
            doc.add_paragraph(f"— {name} ({iso3})")
    else:
        doc.add_paragraph(_NO_MARKETS_LINE)

    tmp = Path(tempfile.mkdtemp(prefix="silk_research_pending_")) / "research.docx"
    try:
        doc.save(str(tmp))
        return tmp.read_bytes()
    finally:
        shutil.rmtree(tmp.parent, ignore_errors=True)
