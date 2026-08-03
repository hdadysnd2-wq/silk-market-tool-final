"""قفلُ التوصيل — القياسُ يعمل على **كلا** مسارَي الدخول عبر HTTP حقيقي.

> **لماذا ملفٌّ منفصل عن `test_hs_attribute_autoresolve.py`؟** ذاك يقفل
> **المنطق** (قراءةُ النطاق، الاختيار، عدمُ الاختلاق)؛ هذا يقفل **التوصيل**:
> أنّ `/analyze` و`/research` كليهما يمرّان بنقطة القياس قبل الحوار، وأنّ
> الرمزَ المحسوم يصل النتيجةَ موسوماً بمصدره. الدرسان ٣٥/٣٧: «إصلاحٌ مُثبَتٌ
> على مسارٍ واحد نصفُ إصلاح» — الحادثةُ الأصلية عادت مرّتين لهذا السبب
> بالضبط، فلا يُقبَل قفلٌ يفحص مساراً واحداً.

هرمتي: الشبكة محجوبة على مستوى `requests` (TestClient يحتاج مقابسه)، ولا
مفتاح كلود ولا مفتاح بحث — فمسارُ الويب يتدهور لفجوةٍ معلنة وحدها، ومسارُ
الصورة يعمل بلا أيّ نداء (السماتُ تصل في جسم الطلب).

Run: python3 -m pytest tests/test_hs_attribute_gate_wiring.py -q
"""
from __future__ import annotations

import contextlib
import os
import sys
from unittest import mock

import pytest

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)


# الصمّامُ **مُطفأٌ افتراضياً** (D1) — هذه الملفّاتُ تختبر الميزةَ نفسَها،
# فتُفعّلها صراحةً. اختبارُ الافتراض نفسِه يعيش في
# `test_flag_is_off_by_default_and_needs_explicit_opt_in` ولا يستعمل هذه.
@pytest.fixture(autouse=True)
def _enable_attribute_resolver(monkeypatch):
    monkeypatch.setenv("SILK_HS_ATTRIBUTE_RESOLVE", "1")



@contextlib.contextmanager
def _env(**vals):
    old = {k: os.environ.get(k) for k in vals}
    try:
        for k, v in vals.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
        yield
    finally:
        for k, v in old.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v


def _block_net():
    return mock.patch("requests.sessions.Session.request",
                      side_effect=OSError("network disabled for offline test"))


def _client():
    import api
    from fastapi.testclient import TestClient
    return TestClient(api.create_app())


def _heading_codes(prefix: str) -> list[str]:
    """بنودُ ترويسةٍ لها نطاقاتٌ رقمية فعلية — من المرجع الرسميّ لا من نموذج."""
    import silk_hs_attributes as attrs
    from silk_hs_resolver import load_hs_reference
    return [c for c in sorted(load_hs_reference())
            if c.startswith(prefix) and attrs.band_of(c)]


# ══════════════ ١) نقطةُ الاختناق مشتركةٌ فعلاً (لا نسخةٌ لكلّ مسار) ═════════

def test_preflight_resolve_wraps_the_gate_and_returns_provenance():
    """`preflight_resolve` = البوّابةُ نفسُها + قياس؛ ثلاثيَّتُها صريحة."""
    import silk_hs_confirm as C
    with _env(SILK_HS_CONFIRM_GATE=None):
        # رمزٌ مؤكَّدٌ دلالياً => لا حجب ولا حسمٌ آلي (السلوك السابق حرفياً).
        code, prov, block = C.preflight_resolve("تمور", "080410")
        assert (code, prov, block) == ("080410", None, None)


def test_gate_block_now_carries_the_measurement_probe_not_a_bare_refusal():
    """الحوارُ لم يعد رفضاً عارياً: ٤٢٢ يحمل ما جُرِّب وما نقص."""
    import silk_hs_confirm as C
    with _env(SILK_HS_CONFIRM_GATE=None), _block_net():
        _c, _p, block = C.preflight_resolve("زبدة الفول السوداني", "040510")
    assert block is not None and block["error"] == "hs_confirmation_needed"
    assert "attribute_probe" in block, "٤٢٢ بلا تقرير قياس — رفضٌ عارٍ"


# ══════════════ ٢) الحسمُ بالصورة يعمل على كِلا المسارين عبر HTTP ════════════

def _label_attrs(codes):
    import silk_hs_attributes as attrs
    disc = attrs.discriminator([{"hs6": c} for c in codes])
    assert disc, "لم يُكتشف مُميِّزٌ رقميّ لهذه الترويسة"
    # قيمةٌ تقع داخل نطاقٍ **وحيد** — تُحسَب من النطاقات نفسها لا تُكتَب صلباً.
    target = next(b for b in disc["bands"]
                  if b["lo"] is not None and b["hi"] is not None)
    value = (target["lo"] + target["hi"]) / 2.0
    return disc, target["hs6"], [
        {"name": disc["label_ar"], "value": value, "unit": disc["unit"]}]


def test_research_gate_resolves_from_label_instead_of_asking(monkeypatch):
    """`/research`: قياسُ البطاقة يحسم البند فلا يُعرَض الحوارُ إطلاقاً."""
    codes = _heading_codes("0401")
    disc, expected, attrs_payload = _label_attrs(codes)
    import silk_hs_confirm as C
    seen: dict = {}

    def _fake_resolve(product, hs_code, *a, **k):
        seen["label_attributes"] = k.get("label_attributes")
        import silk_hs_attributes as A
        rep = A.resolve_by_attribute(
            product, [{"hs6": c} for c in codes],
            label_attributes=k.get("label_attributes"), allow_web=False)
        return (rep["hs6"], rep, None) if rep["hs6"] else (hs_code, None, None)

    monkeypatch.setattr(C, "preflight_resolve", _fake_resolve)
    # التحقّق أنّ المعالجَ يمرّر السمات فعلاً — لا حقلٌ يُستقبَل ويُهمَل.
    with _env(SILK_API_KEY=None, ANTHROPIC_API_KEY=None), _block_net():
        _client().post("/research", json={
            "product": "منتجٌ من هذه الترويسة", "market": "Netherlands",
            "label_attributes": attrs_payload, "persist": False})
    assert seen.get("label_attributes") == attrs_payload, (
        "/research لا يمرّر قياسات البطاقة لنقطة القياس")
    assert _fake_resolve("منتجٌ من هذه الترويسة", None,
                         label_attributes=attrs_payload)[0] == expected


def test_analyze_request_carries_label_attributes_to_the_choke_point(monkeypatch):
    """`/analyze`: نفسُ الحقل ونفسُ نقطة القياس — لا إصلاحَ على مسارٍ واحد."""
    codes = _heading_codes("0401")
    _disc, _expected, attrs_payload = _label_attrs(codes)
    import silk_hs_confirm as C
    seen: dict = {}

    def _fake_resolve(product, hs_code, *a, **k):
        seen["label_attributes"] = k.get("label_attributes")
        return hs_code, None, None

    monkeypatch.setattr(C, "preflight_resolve", _fake_resolve)
    with _env(SILK_API_KEY=None, ANTHROPIC_API_KEY=None), _block_net():
        _client().post("/analyze", json={
            "product": "منتجٌ من هذه الترويسة", "markets": ["NLD"],
            "label_attributes": attrs_payload, "persist": False})
    assert seen.get("label_attributes") == attrs_payload, (
        "/analyze لا يمرّر قياسات البطاقة لنقطة القياس")


# ══════════════ ٣) الرمزُ المحسومُ يصل التقريرَ موسوماً بمصدره ═══════════════

@pytest.mark.parametrize("resolved_from,needle", [
    ("image", "الرمز محدَّد من صورة العبوة"),
    ("web", "الرمز محدَّد من مصدر ويب"),
])
def test_view_discloses_how_an_auto_resolved_code_was_determined(
        resolved_from, needle):
    """لا يُعرَض رمزٌ حُسِم آلياً بلا ذكرِ دليله — على مسار /analyze أيضاً."""
    import silk_render
    codes = _heading_codes("0401")
    view = silk_render.build_view({
        "product": "منتجٌ من هذه الترويسة", "hs_code": codes[1],
        "year": 2023, "markets": [],
        "hs_provenance": {"hs6": codes[1], "resolved_from": resolved_from,
                          "attribute": "fat", "label_ar": "نسبة الدهن",
                          "value": 3.5, "unit": "%",
                          "source_url": "https://example.test/label",
                          "confidence": 0.5},
    })
    assert view["hs_provenance"], "العرضُ لا يحمل مصدر الرمز"
    assert view["hs_provenance"]["resolved_from"] == resolved_from
    assert any(needle in ln for ln in view["limits"]), (
        "سطرُ الإفصاح غائبٌ عن حدود التقرير")
    if resolved_from == "web":
        assert any("https://example.test/label" in ln for ln in view["limits"])


def test_view_has_no_provenance_line_when_the_code_was_not_auto_resolved():
    """بلا حسمٍ آليّ لا سطرَ إفصاحٍ إطلاقاً — لا حقلَ صامتاً يُفسَّر خطأً."""
    import silk_render
    view = silk_render.build_view({"product": "تمور", "hs_code": "080410",
                                   "year": 2023, "markets": []})
    assert view["hs_provenance"] is None
    assert not any("الرمز محدَّد من" in ln for ln in view["limits"])


def test_measured_code_is_not_revalidated_against_the_lexical_resolver():
    """منعُ تناقضِ الحدود مع المتن (عائلة اللائحة ١٢): رمزٌ حُسِم بقياسٍ لا
    يُقارَن بمُحلِّلٍ **لفظيّ** — وهو بالضبط ما عجز عن التمييز فاستُدعي
    القياسُ بدلاً منه. سطرا «المُحلِّل يعيد رمزاً آخر» و«الرمز محدَّد من صورة
    العبوة» لا يجوز أن يظهرا معاً."""
    src = open(os.path.join(_ROOT, "api.py"), encoding="utf-8").read()
    guard = 'if not (isinstance(hs_provenance, dict) and hs_provenance.get("hs6")):'
    assert guard in src, "المصالحة اللفظية تعمل على رمزٍ مقيس — تناقضٌ محتوم"
    body = src.split(guard, 1)[1][:600]
    assert "revalidate(" in body, (
        "حارسُ التناقض لا يغلّف `revalidate` فعلياً")


# ══════════════ ٤) سماتُ البطاقة تُطهَّر ولا تُختلَق ═════════════════════════

def test_intake_drops_attributes_without_a_real_number():
    """سمةٌ بلا رقمٍ صالح تُسقَط بالكامل — لا صفرٌ مُقحَم ولا نصٌّ يُقرأ رقماً."""
    from silk_product_intake import _sanitize_attributes
    out = _sanitize_attributes([
        {"name": "نسبة الدهن", "value": "3,5", "unit": "%"},   # فاصلةٌ عشرية
        {"name": "الوزن", "value": "غير محدَّد", "unit": "g"},  # بلا رقم
        {"name": "", "value": 5, "unit": "%"},                  # بلا اسم
        {"name": "الحجم\n\x00", "value": 1, "unit": "L"},       # محارف تحكّم
        "ليست قاموساً",
    ])
    assert out == [{"name": "نسبة الدهن", "value": 3.5, "unit": "%"},
                   {"name": "الحجم", "value": 1.0, "unit": "L"}]


# ══════════════ G5 — لا انحدارَ في حقولِ الرؤية القائمة ══════════════════════
#
# **نطاقٌ مُعلَنٌ صراحةً:** هذا يقيس طبقةَ **التحليل** (نفسُ ردِّ النموذج =>
# نفسُ الحقول قبل/بعد إضافة `attributes`). لا يقيس **سلوكَ النموذج** — هل
# تُضعِف السمةُ الجديدة جودةَ استخلاص الاسم/الفئة؟ ذلك يتطلّب نداءَ رؤيةٍ
# حقيقياً بمفتاح (G8) وهو غيرُ متاحٍ هنا؛ مذكورٌ في التقرير بدلوه الصحيح.

# مجموعةُ تجاربَ ثابتة: ردودُ نموذجٍ بالشكل **القديم** (بلا `attributes`)
# وبالشكل **الجديد** (معها) لنفس المنتجات.
_VISION_FIXTURES = [
    ('{"product_name_ar":"عصير برتقال","product_name_en":"Orange juice",'
     '"category_hint":"مشروبات","ingredients":["برتقال","سكر"],'
     '"readable":true,"confidence":0.91}',
     '{"product_name_ar":"عصير برتقال","product_name_en":"Orange juice",'
     '"category_hint":"مشروبات","ingredients":["برتقال","سكر"],'
     '"attributes":[{"name":"درجة بريكس","value":11,"unit":""}],'
     '"readable":true,"confidence":0.91}'),
    ('```json\n{"product_name_ar":"","product_name_en":"","category_hint":"",'
     '"ingredients":[],"readable":false,"confidence":0.2}\n```',
     '```json\n{"product_name_ar":"","product_name_en":"","category_hint":"",'
     '"ingredients":[],"attributes":[],"readable":false,"confidence":0.2}\n```'),
    ('{"product_name_ar":"زيت زيتون بكر","product_name_en":"",'
     '"category_hint":"زيوت","ingredients":[],"readable":true,'
     '"confidence":0.77}',
     '{"product_name_ar":"زيت زيتون بكر","product_name_en":"",'
     '"category_hint":"زيوت","ingredients":[],'
     '"attributes":[{"name":"الحجم","value":500,"unit":"ml"}],'
     '"readable":true,"confidence":0.77}'),
]

_PREEXISTING_FIELDS = ("product_name_ar", "product_name_en", "category_hint",
                       "ingredients", "confidence")


@pytest.mark.parametrize("before_raw,after_raw", _VISION_FIXTURES,
                         ids=["readable", "unreadable", "no-en-name"])
def test_added_attributes_field_does_not_change_preexisting_extraction(
        before_raw, after_raw):
    """إضافةُ حقلٍ للنداء المقيس **لا تُغيّر** أيّاً من الحقول القائمة."""
    import silk_product_intake as intake
    tiny = "iVBORw0KGgo="            # ترويسةُ PNG صالحة — يكفي للتحقّق

    def _run(raw):
        with mock.patch.object(intake, "_vision_extract", return_value=raw):
            return intake.intake_image(
                __import__("base64").b64encode(
                    b"\x89PNG\r\n\x1a\n" + b"\x00" * 64).decode(),
                "image/png", "product", allow_vision=True)

    before, after = _run(before_raw), _run(after_raw)
    assert before["ok"] == after["ok"] and before["status"] == after["status"]
    assert before["product_name"] == after["product_name"]
    b_x, a_x = before.get("extraction") or {}, after.get("extraction") or {}
    for field in _PREEXISTING_FIELDS:
        assert b_x.get(field) == a_x.get(field), (
            f"حقلٌ قائمٌ تغيّر بإضافة السمات: {field} "
            f"({b_x.get(field)!r} -> {a_x.get(field)!r})")


def test_old_shape_response_still_yields_empty_attributes_not_a_crash():
    """ردٌّ بالشكل القديم (بلا `attributes`) يمرّ بقائمةٍ فارغة — لا انهيار
    ولا سمةٌ مختلَقة (نموذجٌ لم يُحدَّث بعد، أو ذاكرةٌ قديمة)."""
    import silk_product_intake as intake
    assert intake._sanitize_attributes(None) == []
    assert intake.intake_name("تمر سكري")["extraction"]["attributes"] == []


# ══════════════ سلسلةُ الإفصاح في مستندٍ **مُصيَّرٍ فعلاً** ═══════════════════
#
# هذا القفلُ وُلد من عطلٍ حقيقيّ التقطه فحصُ التصيير لا اختبارُ الوحدة:
# `view["limits"]` كان يحمل السطر (فالاختبارُ أخضر)، بينما **مستندُ العميل** —
# المُسلَّم الفعليّ — يبني أقسامَه من `deep_research` فخرج السطرُ منه تماماً.
# القاعدة: ادّعاءُ «التقريرُ يعرض المصدر» يُثبَت بفتح المستند، لا بقراءة العرض.

@pytest.mark.parametrize("renderer,label", [
    ("render_client_docx", "العميل"),
    ("render_docx", "المشغّل"),
])
def test_provenance_string_appears_in_the_actually_rendered_document(
        renderer, label, tmp_path):
    import silk_render
    import silk_reports
    pytest.importorskip("docx")
    sys.path.insert(0, os.path.join(_ROOT, "tools"))
    from canonical_netherlands import netherlands_research_blob
    from docx import Document

    blob = netherlands_research_blob()
    blob["hs_provenance"] = {
        "hs6": "040120", "resolved_from": "web", "attribute": "fat",
        "label_ar": "نسبة الدهن", "value": 3.5, "unit": "%",
        "source_url": "https://example-retailer.test/label", "confidence": 0.5}
    with _env(SILK_HERMETIC="1"):
        view = silk_render.build_view(blob)
        path = getattr(silk_reports, renderer)(
            view, str(tmp_path / f"{renderer}.docx"))
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    assert "الرمز محدَّد من مصدر ويب" in text, (
        f"مستند {label}: سطرُ الإفصاح غائبٌ عن المستند المُصيَّر فعلاً")
    assert "example-retailer.test/label" in text, (
        f"مستند {label}: الرابطُ المُستشهَد به غائب")


def test_no_provenance_sentence_when_the_code_was_not_measured(tmp_path):
    """بلا حسمٍ آليّ لا جملةَ إفصاحٍ في أيّ مستند — لا نصٌّ مُقحَم."""
    import silk_render
    import silk_reports
    pytest.importorskip("docx")
    sys.path.insert(0, os.path.join(_ROOT, "tools"))
    from canonical_netherlands import netherlands_research_blob
    from docx import Document
    with _env(SILK_HERMETIC="1"):
        view = silk_render.build_view(netherlands_research_blob())
        path = silk_reports.render_client_docx(view, str(tmp_path / "c.docx"))
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert "الرمز محدَّد" not in text
