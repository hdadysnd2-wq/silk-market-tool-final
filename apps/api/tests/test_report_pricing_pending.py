"""J3 — "pricing pending data source" in the executive report.

When NO observed-price source is configured (the registry returns the C3
``GatedPriceProvider``: production settings, no ``LOCALPRICE_API_KEY``, no demo
opt-in), a market with no persisted prices must say the section is *pending a
data source* — EN "pricing pending data source" / AR "الأسعار بانتظار مصدر
بيانات" — instead of the generic searched-and-empty state. Prices that ARE
persisted render exactly as today, and a configured (non-gated) slot keeps the
existing generic declared-gap line. Rides the same seam as
``test_executive_report.py``: ``build_executive_result`` →
``silk_render.build_view`` → ``silk_reports.render_executive_docx``.
"""

from __future__ import annotations

import os
import tempfile

import pytest
from sqlalchemy import select

from app.config import Settings
from app.models import MarketSnapshot
from app.services.report_view import PRICING_PENDING_LINE, build_executive_result
from tests.test_report_docx import _seed_funnel, _snapshot_with_saudi

#: The generic empty-prices contract line the pending note must replace.
_GENERIC_GAP = "سعر غير مرصود"

_PRICE_ROW = {
    "competitor": "Rival GmbH",
    "price": 9.5,
    "currency": "EUR",
    "store": "REWE",
    "url": "https://example.de/p/1",
    "source": "localprice",
    "confidence": 0.7,
    "retrieved_at": "2025-01-01",
}


@pytest.fixture
def gated_prices(monkeypatch):
    """Make the registry select the C3 gated price provider, the real way.

    Production settings with no ``LOCALPRICE_API_KEY`` and no ``ALLOW_MOCK_DATA``
    — the registry's own selection logic then returns ``GatedPriceProvider``
    (locked by ``test_no_mock_data_in_prod.py``); nothing about the gate itself
    is stubbed.
    """
    import app.providers.registry as registry

    settings = Settings(
        environment="production",
        secret_key="x" * 40,
        token_encryption_key="t" * 44,
        trusted_proxy_count=1,
        api_base_url="https://api.silk.example",
        localprice_api_key="",
        allow_mock_data=False,
    )
    monkeypatch.setattr(registry, "get_settings", lambda: settings)
    assert registry.get_price_provider().name == "gated-prices"


def _executive_docx_text(db, product) -> str:
    """Render through the product seam and flatten every paragraph/cell."""
    import silk_reports
    from docx import Document
    from silk_render import build_view

    result = build_executive_result(db, product)
    view = build_view(result)
    path = os.path.join(tempfile.mkdtemp(prefix="silk_j3_"), "executive.docx")
    silk_reports.render_executive_docx(view, path)

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


# -- payload ---------------------------------------------------------------


def test_gated_slot_marks_every_priceless_market_pending(db, factory, product, gated_prices):
    _seed_funnel(db, product)

    markets = build_executive_result(db, product)["executive"]["markets"]

    assert markets, "the funnel top markets feed the executive section"
    for m in markets:
        assert m["prices"] == []
        assert m["prices_note"] == PRICING_PENDING_LINE
    # Both languages are present in the one line the client sees.
    assert "pricing pending data source" in PRICING_PENDING_LINE
    assert "الأسعار بانتظار مصدر بيانات" in PRICING_PENDING_LINE


def test_persisted_prices_still_render_and_suppress_the_pending_note(
    db, factory, product, gated_prices
):
    _seed_funnel(db, product)
    _snapshot_with_saudi(db, product.hs_code)
    snapshot = db.scalar(
        select(MarketSnapshot).where(
            MarketSnapshot.hs_code == product.hs_code,
            MarketSnapshot.market_iso2 == "DE",
        )
    )
    snapshot.observed_prices = [_PRICE_ROW]
    db.commit()

    markets = build_executive_result(db, product)["executive"]["markets"]

    de = next(m for m in markets if m["iso3"] == "DEU")
    assert de["prices"] == [_PRICE_ROW]  # stored prices pass through AS STORED
    assert de["prices_note"] == ""
    # Markets without persisted prices stay pending under the gated slot.
    others = [m for m in markets if m["iso3"] != "DEU"]
    assert others and all(m["prices_note"] == PRICING_PENDING_LINE for m in others)


def test_configured_slot_keeps_the_generic_declared_gap(db, factory, product):
    # Default test settings = local → the mock (non-gated) provider is active:
    # a source IS configured, so nothing is "pending a data source".
    _seed_funnel(db, product)

    markets = build_executive_result(db, product)["executive"]["markets"]

    assert markets
    assert all(m["prices_note"] == "" for m in markets)


# -- rendered docx ---------------------------------------------------------


def test_docx_carries_the_pending_line_when_gated(db, factory, product, gated_prices):
    pytest.importorskip("docx")
    _seed_funnel(db, product)

    text = _executive_docx_text(db, product)

    assert "الأسعار بانتظار مصدر بيانات" in text
    assert "pricing pending data source" in text
    # The pending line REPLACES the generic empty state — never both.
    assert _GENERIC_GAP not in text


def test_docx_keeps_generic_gap_and_real_prices_when_not_gated(db, factory, product):
    pytest.importorskip("docx")
    _seed_funnel(db, product)
    _snapshot_with_saudi(db, product.hs_code)
    snapshot = db.scalar(
        select(MarketSnapshot).where(
            MarketSnapshot.hs_code == product.hs_code,
            MarketSnapshot.market_iso2 == "DE",
        )
    )
    snapshot.observed_prices = [_PRICE_ROW]
    db.commit()

    text = _executive_docx_text(db, product)

    # Persisted prices render as today; priceless markets keep the existing
    # generic contract line; the pending-source line never appears.
    assert "Rival GmbH" in text
    assert _GENERIC_GAP in text
    assert "بانتظار مصدر بيانات" not in text
    assert "pricing pending data source" not in text
