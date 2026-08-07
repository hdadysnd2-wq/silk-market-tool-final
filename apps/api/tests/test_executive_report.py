"""Executive Multi-Market Report: view-model, endpoint, and render task.

The executive report rides the same engine seam as the full docx export
(decision #7): ``build_executive_result`` → ``silk_render.build_view`` →
``silk_reports.render_executive_docx``. Its data contract is provenance-first:
scores fall back Stage-2 → Stage-1, snapshot prices/competitors pass through AS
STORED with their provenance fields, and a product with no analyses yields an
empty executive section rendered as a declared-gap report — nothing is
fabricated (I1).
"""

from __future__ import annotations

import pytest
from sqlalchemy import select

from app.models import (
    Analysis,
    Buyer,
    BuyerSource,
    CountryRanking,
    MarketSnapshot,
    ProductBuyerMatch,
    WorldTrade,
)
from app.services.ranking import run_product_world_analysis
from app.services.report_view import build_executive_result
from tests.conftest import make_buyer_with_contact
from tests.test_report_docx import _seed_funnel, _snapshot_with_saudi

_PRICE_ROW = {
    "competitor": "Rival GmbH",
    "price": 9.5,
    "currency": "EUR",
    "store": "REWE",
    "url": "https://example.de/p/1",
    "source": "localprice",
    "confidence": 0.7,
    "retrieved_at": "2025-01-01",
}


def _ensure_executive_renderer(monkeypatch):
    """Use the engine's real ``render_executive_docx``; stub only if absent.

    The engine half ships ``render_executive_docx(view, path)``. Until it lands,
    a minimal python-docx stub with the same signature keeps the product-side
    contract testable without touching ``packages/``.
    """
    import silk_reports

    if hasattr(silk_reports, "render_executive_docx"):
        return

    def _stub(view, path):
        from docx import Document

        doc = Document()
        doc.add_heading("Executive Multi-Market Report (test stub)", level=1)
        doc.add_paragraph(f"markets: {len((view or {}).get('markets') or [])}")
        doc.save(path)
        return path

    monkeypatch.setattr(silk_reports, "render_executive_docx", _stub, raising=False)


def _latest_analysis(db, product) -> Analysis:
    return db.scalar(
        select(Analysis)
        .where(Analysis.product_id == product.id)
        .order_by(Analysis.created_at.desc())
    )


def _ranking(db, analysis, iso3: str) -> CountryRanking:
    return db.scalar(
        select(CountryRanking).where(
            CountryRanking.analysis_id == analysis.id,
            CountryRanking.importer_iso3 == iso3,
        )
    )


def _seed_stage2_on_deu(db, analysis) -> CountryRanking:
    """Persist a Stage-2 score + enrichment on the DEU ranking row."""
    de = _ranking(db, analysis, "DEU")
    de.stage2_score = 0.81
    de.enrichment = {
        "applied_tariff_pct": None,
        "ppp_gni_per_capita": None,
        "source": "mock",
        "note": "",
        "score_confidence": 0.66,
        "score_components": {
            "market_size": {
                "value": 700.0,
                "source": "world_trade",
                "confidence": 0.9,
                "note": "total imports HS392010 2022",
            },
        },
        "score_model": "silk_market_ranker",
    }
    db.commit()
    return de


def _seed_buyers_in_de(db, product, count_extra: int = 5) -> Buyer:
    """One contact-bearing top buyer + ``count_extra`` bare lower-ranked buyers."""
    top_buyer, _contact = make_buyer_with_contact(
        db, market_iso2="DE", email="buyer@acme.example.de", name="Acme Importers"
    )
    db.add(
        ProductBuyerMatch(
            product_id=product.id,
            buyer_id=top_buyer.id,
            market_iso2="DE",
            relevance_score=90,
        )
    )
    for i in range(count_extra):
        buyer = Buyer(
            name=f"Bare Importer {i}",
            normalized_name=f"bare importer {i}",
            country_iso2="DE",
            source=BuyerSource.maps,
            source_confidence=0.4,
            legal_review_required=True,
        )
        db.add(buyer)
        db.flush()
        db.add(
            ProductBuyerMatch(
                product_id=product.id,
                buyer_id=buyer.id,
                market_iso2="DE",
                relevance_score=50 - i,
            )
        )
    db.commit()
    return top_buyer


# -- build_executive_result ------------------------------------------------


def test_executive_result_scores_rationale_and_snapshot_rows(db, factory, product):
    _seed_funnel(db, product)
    _snapshot_with_saudi(db, product.hs_code)
    analysis = _latest_analysis(db, product)
    _seed_stage2_on_deu(db, analysis)
    snapshot = db.scalar(
        select(MarketSnapshot).where(
            MarketSnapshot.hs_code == product.hs_code,
            MarketSnapshot.market_iso2 == "DE",
        )
    )
    snapshot.observed_prices = [_PRICE_ROW]
    db.commit()
    _seed_buyers_in_de(db, product)

    result = build_executive_result(db, product)

    # Base engine shape is intact — build_view consumes it unchanged.
    assert result["hs_code"] == product.hs_code
    assert result["markets"], "the funnel top-5 still feed the engine markets"

    executive = result["executive"]
    screening = executive["screening"]
    assert screening["total_screened"] == 3
    assert screening["analysis_status"] == "ranked"
    assert isinstance(screening["analysis_at"], str)

    de = next(m for m in executive["markets"] if m["iso3"] == "DEU")
    assert de["iso2"] == "DE"
    # Stage-2 score + confidence, with the persisted rationale components.
    assert de["score"] == pytest.approx(0.81)
    assert de["score_confidence"] == pytest.approx(0.66)
    assert de["rationale_components"] == {
        "market_size": {
            "value": 700.0,
            "source": "world_trade",
            "confidence": 0.9,
            "note": "total imports HS392010 2022",
            "retrieved_at": None,  # not stored by Stage 2 → declared absent (I1)
        }
    }
    # Snapshot rows pass through AS STORED, provenance fields included.
    assert de["prices"] == [_PRICE_ROW]
    assert de["competitors"] == snapshot.top_exporters
    assert {"exporter_name", "share_pct", "value_usd"} <= set(de["competitors"][0])

    # No Stage 2 on NLD → honest fallback to the Stage-1 screen score (I1).
    nld_ranking = _ranking(db, analysis, "NLD")
    nld = next(m for m in executive["markets"] if m["iso3"] == "NLD")
    assert nld["score"] == pytest.approx(float(nld_ranking.screen_score))
    assert nld["score_confidence"] is None
    assert nld["rationale_components"] == {}
    # I9 — the transit-hub tag stays visible on the executive row.
    assert nld["transit_hub"] is True
    assert any("transit" in t.lower() for t in nld["tags"])
    assert nld["buyers"] == []

    # No snapshot for India → empty lists, never invented rows.
    ind = next(m for m in executive["markets"] if m["iso3"] == "IND")
    assert ind["prices"] == [] and ind["competitors"] == []


def test_executive_buyers_carry_provenance_and_cap_at_five(db, factory, product):
    _seed_funnel(db, product)
    _seed_buyers_in_de(db, product, count_extra=5)  # 6 matches total in DE

    executive = build_executive_result(db, product)["executive"]
    de = next(m for m in executive["markets"] if m["iso3"] == "DEU")

    assert len(de["buyers"]) == 5  # cap
    top = de["buyers"][0]
    assert top == {
        "name": "Acme Importers",
        "source": "customs",
        "confidence": pytest.approx(0.85),
        "relevance_score": 90,
        "contacts": 1,
        "legal_review_required": False,
        # C6 — a manually-seeded customs buyer carries no mock/fixture provider
        # name, so it is honestly NOT flagged demo. The mock-discovery path that
        # DOES flag demo is covered by test_provenance_honesty.py.
        "is_demo": False,
        "provider": None,
    }
    # Ordered by relevance; the bare maps-tier buyers keep their review flag
    # and their real (zero) contact count.
    scores = [b["relevance_score"] for b in de["buyers"]]
    assert scores == sorted(scores, reverse=True)
    assert de["buyers"][1]["contacts"] == 0
    assert de["buyers"][1]["source"] == "maps"
    assert de["buyers"][1]["legal_review_required"] is True


def test_executive_markets_cap_at_top_five(db, factory, product):
    isos = ["DEU", "NLD", "IND", "FRA", "ITA", "ESP", "POL"]
    for i, iso3 in enumerate(isos):
        db.add(
            WorldTrade(
                hs6=product.hs_code,
                importer_iso3=iso3,
                year=2022,
                import_usd=1000.0 - i * 50,
                is_transit_hub=False,
                is_mirror=False,
                source="UN Comtrade",
            )
        )
    db.commit()
    run_product_world_analysis(db, product)
    db.commit()

    result = build_executive_result(db, product)
    analysis = _latest_analysis(db, product)
    persisted = db.scalars(
        select(CountryRanking).where(CountryRanking.analysis_id == analysis.id)
    ).all()
    assert len(persisted) == len(isos), "the funnel persisted the whole shortlist"
    assert len(result["executive"]["markets"]) == 5


def test_executive_result_with_zero_analyses_declares_the_gap(db, factory, product):
    result = build_executive_result(db, product)

    assert result["markets"] == []
    assert result["executive"]["markets"] == []
    assert result["executive"]["screening"] == {
        "total_screened": None,
        "analysis_status": "none",
        "analysis_at": None,
    }


# -- endpoint --------------------------------------------------------------


def test_executive_endpoint_returns_a_word_file(
    client, auth_headers, db, factory, product, monkeypatch
):
    _ensure_executive_renderer(monkeypatch)
    _seed_funnel(db, product)

    res = client.get(f"/api/v1/products/{product.id}/report/executive", headers=auth_headers)

    assert res.status_code == 200
    assert "openxmlformats" in res.headers["content-type"]
    assert res.content[:2] == b"PK"  # a .docx is a zip archive


def test_executive_endpoint_renders_with_zero_analyses(
    client, auth_headers, db, factory, product, monkeypatch
):
    _ensure_executive_renderer(monkeypatch)

    res = client.get(f"/api/v1/products/{product.id}/report/executive", headers=auth_headers)

    assert res.status_code == 200  # declared-gap report, never a 500
    assert res.content[:2] == b"PK"


def test_executive_endpoint_requires_auth_and_404s_unknown_product(
    client, auth_headers, db, factory, product, monkeypatch
):
    _ensure_executive_renderer(monkeypatch)

    assert client.get(f"/api/v1/products/{product.id}/report/executive").status_code == 401

    import uuid

    res = client.get(f"/api/v1/products/{uuid.uuid4()}/report/executive", headers=auth_headers)
    assert res.status_code == 404


# -- render task -----------------------------------------------------------


def test_render_executive_report_task_stores_the_docx(db, factory, product, monkeypatch, tmp_path):
    from app.services import storage as storage_mod
    from app.workers.tasks import render_executive_report

    _ensure_executive_renderer(monkeypatch)
    _seed_funnel(db, product)
    store = storage_mod.LocalStorage(str(tmp_path))
    monkeypatch.setattr(storage_mod, "get_storage", lambda *a, **k: store)

    out = render_executive_report.delay(str(product.id)).get()

    key = f"reports/executive/{product.id}.docx"
    assert out == {"product_id": str(product.id), "report_key": key}
    data = storage_mod.get_storage().get_bytes(key)
    assert data is not None and data[:2] == b"PK"
